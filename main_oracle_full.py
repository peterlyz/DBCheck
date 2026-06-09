# -*- coding: utf-8 -*-
#
# Copyright (c) 2025-2026 fiyo (Jack Ge) <sdfiyon@gmail.com>
#
# This file is part of DBCheck, an open-source database health inspection tool.
# DBCheck is released under the MIT License with Attribution Requirements.
# See LICENSE for full license text.
#

"""
DBCheck - Oracle 全面巡检工具（增强版）
=======================================
基于 OS 层 + 数据库层
支持: 10g / 11g / 12c / 18c / 19c / 21c / 23c
作者: Jack Ge
"""

import sys
import os

# frozen 模式下路径处理
if getattr(sys, 'frozen', False):
    sys.path.insert(0, sys._MEIPASS)

import warnings
warnings.filterwarnings("ignore", category=UserWarning, message="pkg_resources is deprecated")

from version import __version__ as VER
import time
import re
import datetime
import json

import io
import argparse
import platform
import getpass

try:
    import oracledb
    _HAS_ORACLE = True
except ImportError:
    _HAS_ORACLE = False

def _get_bundled_instant_client_dir():
    """
    返回 DBCheck 内置的 Oracle Instant Client 路径。
    兼容打包后（sys._MEIPASS）和源码运行模式。
    """
    _system = platform.system().lower()
    if _system == 'windows':
        _subdir = 'windows_x64'
        _marker = 'oci.dll'
    elif _system == 'linux':
        _subdir = 'linux_x64'
        _marker = 'libclntsh.so'
    elif _system == 'darwin':
        _subdir = 'darwin_x64'
        _marker = 'libclntsh.dylib'
    else:
        return None

    # frozen 打包后优先用 _MEIPASS
    _base = getattr(sys, '_MEIPASS', None)
    if _base:
        _client_dir = os.path.join(_base, 'oracle_client', _subdir)
        if os.path.isdir(_client_dir) and os.path.isfile(os.path.join(_client_dir, _marker)):
            return _client_dir

    # 源码运行模式：基于当前文件所在目录
    _base = os.path.dirname(os.path.abspath(__file__))
    _client_dir = os.path.join(_base, 'oracle_client', _subdir)
    if os.path.isdir(_client_dir) and os.path.isfile(os.path.join(_client_dir, _marker)):
        return _client_dir

    return None


def _find_oracle_client_lib_dir():
    """
    自动检测 Oracle Instant Client 库目录（Windows 为主）。
    优先检测 DBCheck 内置路径，其次检测用户环境。
    返回 lib_dir 字符串，找不到返回 None。
    """
    import glob as _glob

    # 0. 优先检测 DBCheck 内置的 Instant Client
    bundled_dir = _get_bundled_instant_client_dir()
    if bundled_dir:
        return bundled_dir

    # 1. ORACLE_HOME 环境变量
    oracle_home = os.environ.get('ORACLE_HOME', '')
    if oracle_home:
        lib_dir = os.path.join(oracle_home, 'bin') if os.name == 'nt' else os.path.join(oracle_home, 'lib')
        if os.path.isdir(lib_dir):
            return lib_dir

    # 2. Windows: 搜索常见 Instant Client 安装路径
    if os.name == 'nt':
        search_roots = [r'C:\oracle', r'C:\instantclient',
                        r'D:\oracle', r'D:\instantclient',
                        r'C:\app', r'D:\app']
        for root in search_roots:
            if not os.path.isdir(root):
                continue
            # 按版本号降序排列，优先使用最新版
            matches = sorted(_glob.glob(os.path.join(root, 'instantclient*')), reverse=True)
            for m in matches:
                if os.path.isdir(m) and os.path.isfile(os.path.join(m, 'oci.dll')):
                    return m

    # 3. Linux/macOS: 检查常见路径
    for candidate in ['/usr/lib/oracle/instantclient', '/opt/oracle/instantclient',
                       '/usr/local/lib/instantclient']:
        # 找具体版本目录
        base = candidate
        if os.path.isdir(base):
            subdirs = sorted([d for d in os.listdir(base) if os.path.isdir(os.path.join(base, d))], reverse=True)
            for sd in subdirs:
                full = os.path.join(base, sd)
                if os.path.isfile(os.path.join(full, 'libclntsh.so')) or \
                   os.path.isfile(os.path.join(full, 'libclntsh.dylib')):
                    return full

    return None

def _get_oracle_conn_thunk_first(dsn, user, password, mode=None,
                                  ssh_host=None, ssh_port=22,
                                  ssh_user=None, ssh_password=None,
                                  ssh_key=None):
    """
    连接 Oracle。
    优先尝试 thin mode（python-oracledb 默认），
    Oracle 11g 及以下自动 fallback 到 thick mode（需 Oracle Instant Client）。
    支持 SSH 隧道。
    """
    import oracledb as _odb

    # SSH 隧道
    _tunnel = None
    _real_dsn = dsn
    if ssh_host:
        try:
            from ssh_tunnel import SSHTunnel
            
            # 从 DSN 解析 Oracle 主机和端口
            _dsn_part = dsn.replace('//', '').split('/')[0]
            _ora_host, _ora_port = (_dsn_part.split(':') + ['1521'])[:2]
            
            # 创建 SSH 隧道
            _tunnel = SSHTunnel(
                ssh_host=ssh_host,
                ssh_port=int(ssh_port),
                ssh_user=ssh_user,
                ssh_password=ssh_password,
                ssh_key=ssh_key,
                remote_host=_ora_host,
                remote_port=int(_ora_port)
            )
            _tunnel.__enter__()
            _local_port = _tunnel.local_port
            _real_dsn = dsn.replace(f'{_ora_host}:{_ora_port}', f'localhost:{_local_port}')
            print(f"  🔗 SSH 隧道: localhost:{_local_port} → {_ora_host}:{_ora_port}")
        except Exception as e:
            print(f"  ⚠️  SSH 隧道失败: {e}, 改为直连")
            _tunnel = None

    def _do_connect():
        if mode is not None:
            return _odb.connect(user=user, password=password, dsn=_real_dsn, mode=mode)
        else:
            return _odb.connect(user=user, password=password, dsn=_real_dsn)

    # ── 第一次尝试：thin mode ──
    try:
        return _do_connect(), _tunnel
    except Exception as e:
        err_str = str(e)
        # thin mode 失败：DPY-3010（11g 不支持）/ DPY-3016（缺少 cryptography 包）→ fallback 到 thick mode
        if 'DPY-3010' not in err_str and 'DPY-3016' not in err_str:
            raise
        # thin mode 不支持 Oracle 11g 及以下 / 缺少依赖包 → fallback 到 thick mode
        print("  ⚠️  thin mode 不可用（%s），切换到 thick mode..." % ('Oracle 11g 及以下' if 'DPY-3010' in err_str else '缺少 cryptography 依赖'))

    # ── 启用 thick mode ──
    try:
        # 先尝试不传 lib_dir（让 oracledb 自动从 PATH / 模块目录查找）
        _odb.init_oracle_client()
        print("  ✅ thick mode 已启用（自动检测 Oracle Client）")
    except Exception:
        # 自动检测 Instant Client 安装路径
        lib_dir = _find_oracle_client_lib_dir()
        if lib_dir:
            print(f"  🔍 检测到 Oracle Client: {lib_dir}")
            try:
                _odb.init_oracle_client(lib_dir=lib_dir)
                print("  ✅ thick mode 已启用")
            except Exception as init_err:
                raise Exception(
                    f'thick mode 初始化失败，Oracle Client 库加载异常。\n'
                    f'库目录: {lib_dir}\n'
                    f'错误: {init_err}\n\n'
                    f'可能原因：\n'
                    f'  1. Instant Client 版本与 python-oracledb 不兼容（建议 19.x 或 23.x）\n'
                    f'  2. 缺少 Visual C++ Redistributable（Windows 需安装 VC++ 2015-2022）\n'
                    f'     https://aka.ms/vs/17/release/vc_redist.x64.exe\n'
                    f'  3. 32/64 位不匹配（Python 64-bit 需 Instant Client 64-bit）'
                ) from init_err
        else:
            _platform = platform.system()
            if _platform == 'Windows':
                _help = (
                    '📥 安装步骤（Windows）：\n'
                    '  1. 下载 Oracle Instant Client Basic Package (64-bit)：\n'
                    '     https://www.oracle.com/database/technologies/instant-client/winx64-64-downloads.html\n'
                    '  2. 解压到 C:\\oracle\\instantclient_23_5（版本号可能不同）\n'
                    '  3. 将 C:\\oracle\\instantclient_23_5 添加到系统 PATH：\n'
                    '     "系统属性" → "环境变量" → Path → 新建 → 粘贴路径\n'
                    '  4. 重新打开命令行窗口后重试巡检\n'
                    '  或者设置环境变量 ORACLE_HOME=C:\\oracle\\instantclient_23_5'
                )
            elif _platform == 'Darwin':
                _help = (
                    '📥 安装步骤（macOS）：\n'
                    '  1. 下载 Oracle Instant Client Basic Package (ARM64 或 x64)：\n'
                    '     https://www.oracle.com/database/technologies/instant-client/macos-arm64-downloads.html\n'
                    '  2. 解压到 ~/Downloads/instantclient_23_3\n'
                    '  3. 设置环境变量后重试巡检'
                )
            else:
                _help = (
                    '📥 安装步骤（Linux）：\n'
                    '  1. 下载 Oracle Instant Client Basic Package (x64)：\n'
                    '     https://www.oracle.com/database/technologies/instant-client/linux-x86-64-downloads.html\n'
                    '  2. 安装到 /usr/lib/oracle/instantclient\n'
                    '  3. 运行 ldconfig 后重试巡检'
                )
            raise Exception(
                f'未检测到 Oracle Instant Client，无法连接 Oracle 11g。\n\n{_help}'
            )

    # ── 第二次尝试：thick mode 重连 ──
    try:
        return _do_connect(), _tunnel
    except Exception as e:
        raise Exception(f'Oracle thick mode 连接失败: {e}') from e

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

# ── ANSI 颜色 ──────────────────────────────────────────────────────────────
CYAN   = "\033[96m"
GREEN  = "\033[92m"
YELLOW = "\033[93m"
MAGENTA= "\033[95m"
BOLD   = "\033[1m"
DIM    = "\033[2m"
RESET  = "\033[0m"
RED    = "\033[91m"

def _enable_ansi():
    try:
        import ctypes
        if os.name == "nt":
            ctypes.windll.kernel32.SetConsoleMode(
                ctypes.windll.kernel32.GetStdHandle(-11), 7)
    except Exception:
        pass

_enable_ansi()

# ── SSH 系统信息采集器（复用 MySQL 的 RemoteSystemInfoCollector）─────────────
try:
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from main_mysql import RemoteSystemInfoCollector
    _HAS_SSH = True
except Exception:
    _HAS_SSH = False

# ── 巡检报告保存 ────────────────────────────────────────────────────────────
try:
    from save_doc_context import SaveDocContext
    _HAS_SAVE = True
except ImportError:
    _HAS_SAVE = False

# ═══════════════════════════════════════════════════════════════════════════
#                    巡检数据采集 — OS 层（SSH / 本地）
# ═══════════════════════════════════════════════════════════════════════════

def get_db_version_and_major(conn):
    """同时获取版本字符串和主版本号，返回 (version_str, ver_major)"""
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT version FROM v$instance")
            row = cur.fetchone()
            if row:
                version_str = row[0]
                m = re.match(r'(\d+)', version_str)
                ver_major = m.group(1) if m else ""
                return version_str, ver_major
    except Exception:
        pass
    return "", ""

# 巡检配置管理章节标题 → 代码检查项名称映射
# DB中的 chapter_title_zh 与 get_checks_for_version() 中的 name 不完全一致，需要映射
ORACLE_CHAPTER_TO_CHECK = {
    # 直接匹配（DB标题 == 代码名称）
    '实例信息': '实例信息',
    'Top SQL': 'Top SQL',
    '无效对象': '无效对象',
    'Data Guard': 'Data Guard',
    'RAC+ASM': 'RAC+ASM',
    # DB章节标题与代码名称不一致的映射
    '健康状态概览': None,          # 无对应检查项（概览由报告生成时汇总）
    '数据库空间使用': '表空间',
    '会话与连接检查': '实例信息',  # 连接信息包含在实例检查中
    '锁等待检查': '阻塞会话',
    '参数与配置检查': '关键参数',
    'RMAN 备份状态': '备份信息',
    '归档日志状态': 'Redo日志',    # 归档与Redo日志相关
    '数据库文件状态': '控制文件',
    '重做日志状态': 'Redo日志',
    'UNDO 表空间': 'Undo信息',
    'AWR 快照状态': 'AWR快照',
    'ADDM 发现': None,             # 无直接对应（ADDM已合并到AWR）
    'TOP SQL 分析': 'Top SQL',
    '表统计信息状态': None,         # Oracle完整版无独立表统计检查项
    '索引统计信息状态': None,       # Oracle完整版无独立索引统计检查项
    '数据库链接': None,             # Oracle完整版无独立db link检查项
    '作业/调度器状态': '作业调度',
    '用户与安全审计': '用户安全',
    '补丁级别': '版本/补丁',
    'Data Guard 状态': 'Data Guard',
    # get_checks_for_version 中有但DB模板没有的章节（保持原样不参与映射）
    # '数据库信息', '长SQL', '性能指标', '死锁检测', '长事务',
    # '闪回/回收站', 'Alert日志'
}

def get_enabled_oracle_chapters(template_id):
    """根据 template_id 从 inspection.db 获取启用的章节标题列表，
    并通过 ORACLE_CHAPTER_TO_CHECK 映射为代码检查项名称"""
    try:
        import sqlite3
        _db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'inspection.db')
        if not os.path.exists(_db_path) or template_id is None:
            return None
        _conn = sqlite3.connect(_db_path)
        _cur = _conn.cursor()
        _cur.execute('''
            SELECT DISTINCT chapter_title_zh
            FROM inspection_chapter ch
            WHERE ch.template_id = ? AND ch.enabled = 1
            ORDER BY ch.sort_order
        ''', (template_id,))
        rows = _cur.fetchall()
        _conn.close()
        if rows:
            # 将DB章节标题映射为代码检查项名称，过滤掉无对应项的（None）
            check_names = set()
            for (title,) in rows:
                check_name = ORACLE_CHAPTER_TO_CHECK.get(title)
                if check_name:
                    check_names.add(check_name)
            return list(check_names) if check_names else None
        return None
    except Exception:
        return None

def load_oracle_chapter_structure(template_id):
    """从 inspection.db 加载 Oracle 章节结构（包含查询SQL），
    返回章节列表，每个章节包含 queries 列表
    格式: [{chapter_number, chapter_title_zh, queries: [{query_key, query_sql, query_description_zh}]}]"""
    try:
        import sqlite3
        _db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'inspection.db')
        if not os.path.exists(_db_path) or template_id is None:
            return []
        _conn = sqlite3.connect(_db_path)
        _cur = _conn.cursor()
        _cur.execute(
            "SELECT id, chapter_number, chapter_title_zh, sort_order, enabled "
            "FROM inspection_chapter WHERE template_id=? ORDER BY sort_order",
            (template_id,))
        chapters = []
        for ch_row in _cur.fetchall():
            ch_id, ch_num, ch_title, ch_sort, ch_enabled = ch_row
            if not ch_enabled:
                continue
            _cur.execute(
                "SELECT query_key, query_sql, query_description_zh, enabled, sort_order "
                "FROM inspection_query WHERE chapter_id=? ORDER BY sort_order",
                (ch_id,))
            queries = []
            for q_row in _cur.fetchall():
                q_key, q_sql, q_desc, q_enabled, q_sort = q_row
                if not q_enabled or not q_sql:
                    continue
                # 清洗SQL: 去除多余空白和末尾分号（Oracle oracledb 不接受分号，否则报 ORA-00933）
                q_sql_clean = ' '.join(q_sql.split()).rstrip(';')
                queries.append({
                    'query_key': q_key,
                    'query_sql': q_sql_clean,
                    'query_description_zh': q_desc or q_key,
                })
            if queries:
                chapters.append({
                    'chapter_number': ch_num,
                    'chapter_title_zh': ch_title,
                    'queries': queries,
                })
        _conn.close()
        return chapters
    except Exception as e:
        print(f"[WARN] 加载 Oracle 章节结构失败: {e}")
        return []

def execute_oracle_chapter_queries(conn, chapters):
    """执行所有章节的查询，返回 chapter_results
    格式: [{chapter_number, chapter_title_zh, queries_results: [{query_key, query_description_zh, columns, data}]}]"""
    chapter_results = []
    for chapter in chapters:
        query_results = []
        for query in chapter['queries']:
            try:
                cur = conn.cursor()
                cur.execute(query['query_sql'])
                columns = [col[0] for col in cur.description] if cur.description else []
                data = cur.fetchall()
                cur.close()
                query_results.append({
                    'query_key': query['query_key'],
                    'query_description_zh': query['query_description_zh'],
                    'columns': columns,
                    'data': data,
                })
            except Exception as e:
                query_results.append({
                    'query_key': query['query_key'],
                    'query_description_zh': query['query_description_zh'],
                    'columns': [],
                    'data': [],
                    'error': str(e),
                })
        chapter_results.append({
            'chapter_number': chapter['chapter_number'],
            'chapter_title_zh': chapter['chapter_title_zh'],
            'queries_results': query_results,
        })
    return chapter_results

def get_checks_for_version(ver_major, chapter_filter=None):
    """根据主版本号返回对应的巡检函数列表

    10g → dbcheck10g.sql：WMSYS.WM_CONCAT 替代 listagg
    11g → dbcheck11g.sql：标准 listagg
    12c+ → dbcheck12c.sql：CDB/PDB 支持、gv$crs_resource_v2 等

    chapter_filter: 可选的章节标题列表，如果指定则只返回匹配的巡检项
    """
    # ── 10g ──────────────────────────────────────────────────────────────────
    if ver_major == "10":
        checks = [
            ("实例信息",      oracle_check_instance),
            ("数据库信息",    oracle_check_database_v10),
            ("版本/补丁",     oracle_check_version_and_patches),
            ("表空间",        oracle_check_tablespace),
            ("Redo日志",      oracle_check_redolog),
            ("控制文件",      oracle_check_controlfile),
            ("SGA/PGA内存",   oracle_check_sga_pga),
            ("关键参数",      oracle_check_params),
            ("Undo信息",      oracle_check_undo),
            ("长SQL",         oracle_check_long_sql),
            ("性能指标",      oracle_check_performance),
            ("Top SQL",       oracle_check_top_sql),
            ("无效对象",      oracle_check_invalid_objects),
            ("阻塞会话",      oracle_check_blocking),
            ("死锁检测",      oracle_check_deadlock),
            ("长事务",         oracle_check_long_trx),
            ("用户安全",      oracle_check_users),
            # 备份/DataGuard/RAC/AWR 在 10g 通常不可用，已在对应函数内做版本适配
            ("备份信息",      oracle_check_backup),
            ("闪回/回收站",   oracle_check_flashback),
            ("Data Guard",    oracle_check_dataguard),
            ("RAC+ASM",       oracle_check_rac),
            # 10g 无 AWR，用 statspack 替代（函数内部已处理）
            ("AWR快照",       oracle_check_awr),
            ("作业调度",      oracle_check_jobs),
            ("Alert日志",     oracle_check_alert),
        ]
        if chapter_filter:
            checks = [(name, fn) for name, fn in checks if name in chapter_filter]
        return checks

    # ── 11g ──────────────────────────────────────────────────────────────────
    if ver_major == "11":
        checks = [
            ("实例信息",      oracle_check_instance),
            ("数据库信息",    oracle_check_database_v11),
            ("版本/补丁",     oracle_check_version_and_patches),
            ("表空间",        oracle_check_tablespace),
            ("Redo日志",      oracle_check_redolog),
            ("控制文件",      oracle_check_controlfile),
            ("SGA/PGA内存",   oracle_check_sga_pga),
            ("关键参数",      oracle_check_params),
            ("Undo信息",      oracle_check_undo),
            ("长SQL",         oracle_check_long_sql),
            ("性能指标",      oracle_check_performance),
            ("Top SQL",       oracle_check_top_sql),
            ("无效对象",      oracle_check_invalid_objects),
            ("阻塞会话",      oracle_check_blocking),
            ("死锁检测",      oracle_check_deadlock),
            ("长事务",         oracle_check_long_trx),
            ("用户安全",      oracle_check_users),
            ("备份信息",      oracle_check_backup),
            ("闪回/回收站",   oracle_check_flashback),
            ("Data Guard",    oracle_check_dataguard),
            ("RAC+ASM",       oracle_check_rac),
            ("AWR快照",       oracle_check_awr),
            ("作业调度",      oracle_check_jobs),
            ("Alert日志",     oracle_check_alert),
        ]
        if chapter_filter:
            checks = [(name, fn) for name, fn in checks if name in chapter_filter]
        return checks

    # ── 12c 及以上（基准）──────────────────────────────────────────────────────
    # 12c SQL 作为基准；19c 及以上出错的查询项，使用 v19 兼容版本覆盖
    checks = [
        ("实例信息",      oracle_check_instance),
        ("数据库信息",    oracle_check_database_v12plus),
        ("版本/补丁",     oracle_check_version_and_patches),
        ("表空间",        oracle_check_tablespace),
        ("Redo日志",      oracle_check_redolog),
        ("控制文件",      oracle_check_controlfile),
        ("SGA/PGA内存",   oracle_check_sga_pga),
        ("关键参数",      oracle_check_params),
        ("Undo信息",      oracle_check_undo),
        ("长SQL",         oracle_check_long_sql),
        ("性能指标",      oracle_check_performance),
        ("Top SQL",       oracle_check_top_sql),
        ("无效对象",      oracle_check_invalid_objects),
        ("阻塞会话",      oracle_check_blocking),
        ("死锁检测",      oracle_check_deadlock),
        ("长事务",         oracle_check_long_trx),
        ("用户安全",      oracle_check_users),
        ("备份信息",      oracle_check_backup),
        ("闪回/回收站",   oracle_check_flashback),
        ("Data Guard",    oracle_check_dataguard),
        ("RAC+ASM",       oracle_check_rac),
        ("AWR快照",       oracle_check_awr),
        ("作业调度",      oracle_check_jobs),
        ("Alert日志",     oracle_check_alert),
    ]

    # 19c 及以上：对出错项应用 v19 兼容版（保留 12c 基准不变）
    if ver_major and int(ver_major) >= 19:
        for i, (name, _fn) in enumerate(checks):
            if name == "数据库信息":
                checks[i] = ("数据库信息", oracle_check_database_v19)
            elif name == "表空间":
                checks[i] = ("表空间", oracle_check_tablespace_v19)
            elif name == "Redo日志":
                checks[i] = ("Redo日志", oracle_check_redolog_v19)
            elif name == "Top SQL":
                checks[i] = ("Top SQL", oracle_check_top_sql_v19)
            elif name == "备份信息":
                checks[i] = ("备份信息", oracle_check_backup_v19)
            elif name == "闪回/回收站":
                checks[i] = ("闪回/回收站", oracle_check_flashback_v19)
            elif name == "Data Guard":
                checks[i] = ("Data Guard", oracle_check_dataguard_v19)
            elif name == "AWR快照":
                checks[i] = ("AWR快照", oracle_check_awr_v19)
            elif name == "Alert日志":
                checks[i] = ("Alert日志", oracle_check_alert_v19)

    # 如果指定了章节过滤，只保留匹配的巡检项
    if chapter_filter:
        checks = [(name, fn) for name, fn in checks if name in chapter_filter]

    return checks

class OSCollector:
    """OS 层信息采集（通过 SSH 或本地命令）"""

    def __init__(self, ssh_conn=None):
        self.ssh = ssh_conn  # paramiko SSHClient 或 None（本地）

    def run_cmd(self, cmd):
        """通过 SSH 或本地执行命令"""
        if self.ssh:
            try:
                stdin, stdout, stderr = self.ssh.exec_command(cmd)
                return stdout.read().decode('utf-8', errors='ignore')
            except Exception:
                return ""
        else:
            import subprocess
            try:
                return subprocess.check_output(cmd, shell=True, stderr=subprocess.DEVNULL, timeout=30).decode('utf-8', errors='ignore')
            except Exception:
                return ""

    def collect(self):
        """收集 OS 层全部信息"""
        uname = platform.system()
        is_linux = (uname == "Linux")

        data = {}

        # 基础信息
        data['hostname'] = platform.node()
        data['os_version'] = self.run_cmd("cat /etc/*release 2>/dev/null | head -1").strip()
        data['kernel'] = platform.release()
        data['uptime'] = self.run_cmd("uptime 2>/dev/null | head -n1").strip()

        # CPU
        data['cpu_model'] = self.run_cmd("awk -F': ' '/model name/ {print $2; exit}' /proc/cpuinfo 2>/dev/null").strip()
        data['cpu_count'] = self.run_cmd("nproc 2>/dev/null").strip() or self.run_cmd("grep -c processor /proc/cpuinfo 2>/dev/null").strip()
        cpu_idle = self.run_cmd("vmstat 1 2 | awk 'NR==4 {print 100-$15}'").strip()
        data['cpu_usage_pct'] = cpu_idle

        # 内存
        mem_total = self.run_cmd("free -m | awk '/Mem:/ {print $2}'").strip()
        mem_used   = self.run_cmd("free -m | awk '/Mem:/ {print $3}'").strip()
        data['mem_total_mb']  = mem_total
        data['mem_used_mb']   = mem_used
        data['mem_usage_pct'] = self.run_cmd("free -m | awk '/Mem:/ {print $3/$2*100}'").strip()

        # Swap
        swap_total = self.run_cmd("free -m | awk '/Swap:/ {print $2}'").strip()
        swap_used   = self.run_cmd("free -m | awk '/Swap:/ {print $3}'").strip()
        data['swap_total_mb'] = swap_total
        data['swap_used_mb']  = swap_used

        # 负载
        data['load_average'] = self.run_cmd("uptime | awk -F': ' '{print $2}'").strip()

        # 磁盘
        data['disk_usage'] = self.run_cmd("df -Ph 2>/dev/null | grep -v 'tmpfs\\|devtmpfs\\|overlay\\|shm'").strip()

        # /etc/hosts
        data['hosts'] = self.run_cmd("sed '1,2d' /etc/hosts 2>/dev/null | grep -v '^$'").strip()

        # sysctl 参数
        data['sysctl'] = self.run_cmd(
            "grep -E 'kernel.shmall|kernel.shmmax|kernel.sem|kernel.shmmni|fs.aio-max-nr|fs.file-max|vm.swappiness|vm.nr_hugepages' "
            "/etc/sysctl.conf 2>/dev/null"
        ).strip()

        # limits.conf
        data['limits'] = self.run_cmd("grep -v '^#\\|^$' /etc/security/limits.conf 2>/dev/null").strip()

        # HugePages
        data['hugepages'] = self.run_cmd(
            "awk '/MemTotal|HugePages_Total|HugePages_Free/ {print $1\":\"$2}' /proc/meminfo 2>/dev/null"
        ).strip()

        # Transparent HugePages
        thp = self.run_cmd("cat /sys/kernel/mm/transparent_hugepage/enabled 2>/dev/null").strip()
        data['thp'] = thp

        # crontab
        data['crontab'] = self.run_cmd("crontab -l 2>/dev/null").strip()

        # 网络
        data['network'] = self.run_cmd("ip addr show 2>/dev/null | grep 'inet '").strip()

        # /etc/passwd（数据库用户检查用）
        data['oracle_users'] = self.run_cmd(
            "grep -E '^(oracle|grid|root):' /etc/passwd 2>/dev/null"
        ).strip()

        return data

# ═══════════════════════════════════════════════════════════════════════════
#                    巡检数据采集 — Oracle 数据库层
# ═══════════════════════════════════════════════════════════════════════════

def oracle_check_instance(conn):
    """实例基本信息"""
    results = {}
    cur = conn.cursor()
    try:
        # 实例信息
        cur.execute("""
            SELECT INST_ID, INSTANCE_NAME, HOST_NAME, VERSION,
                   STARTUP_TIME, STATUS, PARALLEL, LOG_MODE,
                   DATABASE_ROLE, OPEN_MODE
            FROM gv$instance
        """)
        results['instance'] = cur.fetchall()
    except Exception as e:
        results['instance_error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_database(conn):
    """数据库基本信息"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT DBID, NAME, DATABASE_ROLE, CREATED, LOG_MODE, OPEN_MODE,
                   CDB, PLUGGABLE_NDB
            FROM v$database
        """)
        results['database'] = cur.fetchall()

        # 全局名
        cur.execute("SELECT global_name FROM global_name")
        results['global_name'] = cur.fetchone()

        # 字符集（sys.props$ 在 SYSDBA 下可能无权限，改为 nls_database_parameters）
        try:
            cur.execute("""
                SELECT parameter, value
                FROM nls_database_parameters
                WHERE parameter IN ('NLS_CHARACTERSET', 'NLS_NCHAR_CHARACTERSET')
            """)
            rows = cur.fetchall()
            results['charset'] = tuple(r[1] for r in rows) if rows else ('', '')
        except Exception as e:
            results['charset'] = ('', '')

        # 块大小
        cur.execute("SELECT value FROM v$parameter WHERE name='db_block_size'")
        r = cur.fetchone()
        results['block_size'] = r[0] if r else ''

        # SGA/PGA
        for param in ['sga_max_size', 'sga_target', 'pga_aggregate_target',
                      'memory_max_target', 'memory_target']:
            cur.execute(f"SELECT value FROM v$parameter WHERE name='{param}'")
            r = cur.fetchone()
            results[param] = r[0] if r else ''

        # SPFILE
        cur.execute("SELECT value FROM v$parameter WHERE name='spfile'")
        r = cur.fetchone()
        results['spfile'] = r[0] if r else ''

        # OMF
        cur.execute("SELECT value FROM v$parameter WHERE name='db_create_file_dest'")
        r = cur.fetchone()
        results['omf'] = r[0] if r else ''

        # 归档模式
        cur.execute("SELECT log_mode FROM v$database")
        r = cur.fetchone()
        results['log_mode'] = r[0] if r else ''

        # force logging
        cur.execute("SELECT force_logging FROM v$database")
        r = cur.fetchone()
        results['force_logging'] = r[0] if r else ''

        #  flashback
        cur.execute("SELECT flashback_on FROM v$database")
        r = cur.fetchone()
        results['flashback_on'] = r[0] if r else ''

        # 创建时间
        cur.execute("SELECT TO_CHAR(CREATED, 'YYYY-MM-DD HH24:MI:SS') FROM v$database")
        r = cur.fetchone()
        results['created'] = r[0] if r else ''

        # 启动时间
        cur.execute("SELECT TO_CHAR(STARTUP_TIME, 'YYYY-MM-DD HH24:MI:SS') FROM v$instance")
        r = cur.fetchone()
        results['startup_time'] = r[0] if r else ''

    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

# ═══════════════════════════════════════════════════════════════════════════
#            版本专用数据库层巡检（对应 dbcheck10g/11g/12c.sql）
# ═══════════════════════════════════════════════════════════════════════════

def _base_db_check(conn, cdb_col="", pdb_col=""):
    """三个版本公用的基础查询部分，避免代码重复"""
    results = {}
    cur = conn.cursor()
    try:
        if cdb_col and pdb_col:
            # 12c+：CDB/PDB 架构
            cur.execute(f"""
                SELECT DBID, NAME, DATABASE_ROLE, CREATED, LOG_MODE, OPEN_MODE,
                       {cdb_col} CDB, {pdb_col} PLUGGABLE_DB
                FROM v$database
            """)
        else:
            # 10g/11g：无 CDB 概念
            cur.execute("""
                SELECT DBID, NAME, DATABASE_ROLE, CREATED, LOG_MODE, OPEN_MODE
                FROM v$database
            """)
        results['database'] = cur.fetchall()

        cur.execute("SELECT global_name FROM global_name")
        results['global_name'] = cur.fetchone()

        # 字符集
        try:
            cur.execute("""
                SELECT parameter, value
                FROM nls_database_parameters
                WHERE parameter IN ('NLS_CHARACTERSET', 'NLS_NCHAR_CHARACTERSET')
            """)
            rows = cur.fetchall()
            results['charset'] = tuple(r[1] for r in rows) if rows else ('', '')
        except Exception:
            results['charset'] = ('', '')

        for param in ['db_block_size', 'sga_max_size', 'sga_target',
                      'pga_aggregate_target', 'memory_max_target', 'memory_target']:
            cur.execute(f"SELECT value FROM v$parameter WHERE name='{param}'")
            r = cur.fetchone()
            results[param] = r[0] if r else ''

        cur.execute("SELECT value FROM v$parameter WHERE name='spfile'")
        r = cur.fetchone()
        results['spfile'] = r[0] if r else ''

        cur.execute("SELECT value FROM v$parameter WHERE name='db_create_file_dest'")
        r = cur.fetchone()
        results['omf'] = r[0] if r else ''

        for col in ['log_mode', 'force_logging', 'flashback_on']:
            cur.execute(f"SELECT {col} FROM v$database")
            r = cur.fetchone()
            results[col] = r[0] if r else ''

        cur.execute("SELECT TO_CHAR(CREATED, 'YYYY-MM-DD HH24:MI:SS') FROM v$database")
        r = cur.fetchone()
        results['created'] = r[0] if r else ''

        cur.execute("SELECT TO_CHAR(STARTUP_TIME, 'YYYY-MM-DD HH24:MI:SS') FROM v$instance")
        r = cur.fetchone()
        results['startup_time'] = r[0] if r else ''

    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_database_v10(conn):
    """Oracle 10g 数据库层巡检 — 基于 dbcheck10g.sql

    特点：WMSYS.WM_CONCAT 替代 listagg；无 CDB；无 PLUGGABLE_NDB
    """
    return _base_db_check(conn)

def oracle_check_database_v11(conn):
    """Oracle 11g 数据库层巡检 — 基于 dbcheck11g.sql

    特点：标准 listagg；无 CDB；无 PLUGGABLE_NDB；gv$instance 需额外字段兼容
    """
    results = _base_db_check(conn)
    cur = conn.cursor()
    try:
        # 11g 特有的 listagg（RAC 节点列表）
        # 10g 用 WMSYS.WM_CONCAT，11g 用标准 listagg
        cur.execute("""
            SELECT 'Instances: [' || listagg(instance_name, ', ') within group(order by instance_name) || '] ' as instances
            FROM gv$instance
        """)
        row = cur.fetchone()
        results['rac_instances'] = row[0] if row else ''

        # 11g 回收站（dba_recyclebin）
        try:
            cur.execute("""
                SELECT owner,
                       round(SUM(a.space *
                                 (SELECT value FROM v$parameter WHERE name='db_block_size')) / 1024 / 1024, 2) recyb_size_M,
                       count(1) recyb_cnt
                FROM dba_recyclebin a
                GROUP BY owner
            """)
            results['recyclebin'] = cur.fetchall()
        except Exception:
            pass

    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_database_v12plus(conn):
    """Oracle 12c+ 数据库层巡检 — 基于 dbcheck12c.sql

    特点：CDB/PDB 架构；listagg；cdb_recyclebin；gv$crs_resource_v2
    """
    return _base_db_check(conn, cdb_col="CDB", pdb_col="PLUGGABLE_DB")

def oracle_check_version_and_patches(conn):
    """数据库版本和补丁"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("SELECT banner FROM v$version")
        results['version'] = cur.fetchall()

        # OPatch 补丁
        try:
            cur.execute("SELECT * FROM v$system_patch WHERE patch_id IS NOT NULL")
            results['patches'] = cur.fetchall()
        except Exception:
            # 12c 及以上用这个
            try:
                cur.execute("""
                    SELECT patch_id, patch_type, description, action, action_time
                    FROM dba_registry_sqlpatch
                    ORDER BY action_time DESC
                """)
                results['patches'] = cur.fetchall()
            except Exception:
                results['patches'] = []
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_tablespace(conn):
    """表空间使用率（通用版）：实测列名 dba_free_space.BYTES / dba_temp_free_space.FREE_SPACE"""
    results = {}
    cur = conn.cursor()
    try:
        # 永久表空间：dba_tablespaces LEFT JOIN dba_data_files LEFT JOIN dba_free_space
        # pct_used 考虑可扩展空间：分母用 GREATEST(curr_mb, max_mb)
        cur.execute("""
            SELECT t.tablespace_name,
                   t.status,
                   ROUND(NVL(df.curr_mb,0), 2) curr_mb,
                   ROUND(NVL(df.max_mb,0), 2) max_mb,
                   ROUND(NVL(df.curr_mb,0) - NVL(fs.free_mb,0), 2) used_mb,
                   ROUND(NVL(fs.free_mb,0), 2) free_mb,
                   ROUND((NVL(df.curr_mb,0) - NVL(fs.free_mb,0)) /
                         NULLIF(GREATEST(NVL(df.curr_mb,0), NVL(df.max_mb,0)), 0) * 100, 2) pct_used
            FROM dba_tablespaces t
            LEFT JOIN (SELECT tablespace_name,
                              SUM(bytes/1024/1024) curr_mb,
                              SUM(MAXBYTES/1024/1024) max_mb
                       FROM dba_data_files GROUP BY tablespace_name) df
               ON t.tablespace_name = df.tablespace_name
            LEFT JOIN (SELECT tablespace_name,
                              SUM(bytes/1024/1024) free_mb
                       FROM dba_free_space GROUP BY tablespace_name) fs
               ON t.tablespace_name = fs.tablespace_name
            WHERE t.contents = 'PERMANENT'
            ORDER BY pct_used DESC NULLS LAST
        """)
        results['data_tablespaces'] = cur.fetchall()

        # 临时表空间：dba_tablespaces LEFT JOIN dba_temp_files LEFT JOIN dba_temp_free_space(FREE_SPACE)
        cur.execute("""
            SELECT t.tablespace_name,
                   t.status,
                   ROUND(NVL(tf.curr_mb,0), 2) curr_mb,
                   ROUND(NVL(tf.max_mb,0), 2) max_mb,
                   '-' used_mb,
                   ROUND(NVL(tfs.free_mb,0), 2) free_mb,
                   '-' pct_used
            FROM dba_tablespaces t
            LEFT JOIN (SELECT tablespace_name,
                              SUM(bytes/1024/1024) curr_mb,
                              SUM(MAXBYTES/1024/1024) max_mb
                       FROM dba_temp_files GROUP BY tablespace_name) tf
               ON t.tablespace_name = tf.tablespace_name
            LEFT JOIN (SELECT tablespace_name,
                              SUM(free_space/1024/1024) free_mb
                       FROM dba_temp_free_space GROUP BY tablespace_name) tfs
               ON t.tablespace_name = tfs.tablespace_name
            WHERE t.contents = 'TEMPORARY'
            ORDER BY t.tablespace_name
        """)
        results['temp_tablespaces'] = cur.fetchall()

        # 自动扩展文件
        cur.execute("""
            SELECT tablespace_name, file_name,
                   ROUND(bytes/1024/1024,2) curr_mb,
                   ROUND(MAXBYTES/1024/1024,2) max_mb,
                   AUTOEXTENSIBLE
            FROM dba_data_files
            WHERE AUTOEXTENSIBLE = 'YES'
            ORDER BY tablespace_name
        """)
        results['autoextend_files'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_redolog(conn):
    """Redo 日志检查（11g 兼容）"""
    results = {}
    cur = conn.cursor()
    try:
        # 组成员和大小
        cur.execute("""
            SELECT GROUP#, THREAD#, SEQUENCE#, ROUND(BYTES/1024/1024,2) size_mb,
                   STATUS, MEMBERS, ARCHIVED
            FROM v$log
            ORDER BY THREAD#, GROUP#
        """)
        results['logs'] = cur.fetchall()

        # 日志文件
        cur.execute("""
            SELECT GROUP#, MEMBER, TYPE, STATUS
            FROM v$logfile
            ORDER BY GROUP#
        """)
        results['logfiles'] = cur.fetchall()

        # 最近 Redo 切换频率（11g 兼容：避免 v$loghist 与 v$log 的 GROUP# 别名问题）
        cur.execute("""
            SELECT hl.thread#,
                   COUNT(*) switch_cnt,
                   ROUND(COUNT(*) * (SELECT MAX(bytes) FROM v$log)/1024/1024/1024, 2) total_mb
            FROM v$loghist hl
            WHERE hl.first_time > SYSDATE - 7
            GROUP BY hl.thread#
        """)
        results['redo_switch'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_controlfile(conn):
    """控制文件"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT NAME, STATUS, IS_RECOVERY_DEST_FILE, BLOCK_SIZE, FILE_SIZE_BLKS
            FROM v$controlfile
            ORDER BY STATUS, NAME
        """)
        results['controlfiles'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_invalid_objects(conn):
    """无效对象"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT OWNER, OBJECT_TYPE, COUNT(*) cnt
            FROM dba_objects
            WHERE STATUS = 'INVALID'
            GROUP BY OWNER, OBJECT_TYPE
            ORDER BY cnt DESC
        """)
        results['invalid_by_type'] = cur.fetchall()

        cur.execute("""
            SELECT OWNER, OBJECT_NAME, OBJECT_TYPE, STATUS, LAST_DDL_TIME
            FROM dba_objects
            WHERE STATUS = 'INVALID'
              AND OWNER NOT IN ('SYS','SYSTEM')
            ORDER BY LAST_DDL_TIME DESC
        """)
        results['invalid_detail'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_users(conn):
    """用户安全检查"""
    results = {}
    cur = conn.cursor()
    try:
        # 默认密码用户（常见高危账号）
        default_accounts = ['SCOTT','JONES','ADAMS','CLARK','BLAKE','HR','OE','PM','IX','SH','DIP','ORACLE_OCM','XS$NULL','APPQOSSYS']
        placeholders = ','.join([f":{i}" for i in range(len(default_accounts))])
        cur.execute(f"""
            SELECT username, account_status, lock_date, expiry_date, created
            FROM dba_users
            WHERE username IN ({placeholders})
        """, default_accounts)
        results['default_accounts'] = cur.fetchall()

        # 锁定的用户
        cur.execute("""
            SELECT username, account_status, lock_date, expiry_date
            FROM dba_users
            WHERE account_status NOT IN ('OPEN', 'EXPIRED(GRACE)')
            ORDER BY account_status
        """)
        results['locked_users'] = cur.fetchall()

        # 系统角色
        cur.execute("""
            SELECT granted_role, grantee, admin_option
            FROM dba_role_privs
            WHERE grantee NOT IN ('SYS', 'SYSTEM')
              AND admin_option = 'YES'
            ORDER BY granted_role
        """)
        results['admin_roles'] = cur.fetchall()

        # Profile
        cur.execute("""
            SELECT profile, resource_name, resource_type, LIMIT
            FROM dba_profiles
            ORDER BY profile, resource_name
        """)
        results['profiles'] = cur.fetchall()

        # 密码有效期
        cur.execute("""
            SELECT profile, LIMIT PASSWORD_LIFE_TIME
            FROM dba_profiles
            WHERE resource_name = 'PASSWORD_LIFE_TIME'
            ORDER BY profile
        """)
        results['password_policy'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_top_sql(conn, limit=20):
    """Top SQL（按逻辑读/物理读/executions）"""
    results = {}
    cur = conn.cursor()
    try:
        # 按 Buffer Gets 排序
        cur.execute(f"""
            SELECT * FROM (
                SELECT sql_id, SUBSTR(sql_text,1,80) sql_text,
                       ROUND(buffer_gets/1024/1024,2) buf_mb,
                       ROUND(disk_reads/1024/1024,2) disk_mb,
                       executions, ROUND(elapsed_time/1000000,2) elapsed_sec,
                       ROUND(buffer_gets/DECODE(executions,0,1,executions)) gets_per_exec,
                       module
                FROM v$sql
                WHERE executions > 0
                ORDER BY buffer_gets DESC
            ) WHERE ROWNUM <= {limit}
        """)
        results['top_sql_buffer_gets'] = cur.fetchall()

        # 按磁盘读排序
        cur.execute(f"""
            SELECT * FROM (
                SELECT sql_id, SUBSTR(sql_text,1,80) sql_text,
                       ROUND(disk_reads/1024/1024,2) disk_mb,
                       executions, ROUND(elapsed_time/1000000,2) elapsed_sec,
                       module
                FROM v$sql
                WHERE executions > 0
                ORDER BY disk_reads DESC
            ) WHERE ROWNUM <= {limit}
        """)
        results['top_sql_disk_reads'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_awr(conn):
    """AWR 快照信息"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT instance_number, snap_id, TO_CHAR(begin_interval_time,'YYYY-MM-DD HH24:MI') bt,
                   TO_CHAR(end_interval_time,'YYYY-MM-DD HH24:MI') et,
                   ROUND(EXTRACT(DAY FROM (end_interval_time - begin_interval_time)) * 24
                         + EXTRACT(HOUR FROM (end_interval_time - begin_interval_time))
                         + EXTRACT(MINUTE FROM (end_interval_time - begin_interval_time)) / 60, 2) elapsed_hr,
                   ERROR_COUNT
            FROM dba_hist_snapshot
            WHERE end_interval_time > SYSDATE - 7
            ORDER BY instance_number, snap_id DESC
        """)
        results['awr_snaps'] = cur.fetchall()

        # AWR 设置
        cur.execute("""
            SELECT * FROM dba_hist_wr_control
        """)
        results['awr_settings'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_performance(conn):
    """性能指标（Session / Wait Events / SGA/PGA）"""
    results = {}
    cur = conn.cursor()
    try:
        # 当前会话数
        cur.execute("""
            SELECT status, COUNT(*) FROM v$session GROUP BY status
        """)
        results['session_by_status'] = cur.fetchall()

        # 等待事件 Top10（11g 兼容：避免除零）
        cur.execute("""
            SELECT * FROM (
                SELECT event, total_waits, ROUND(time_waited/100,2) time_sec,
                       ROUND(total_waits/GREATEST(time_waited,0.001)*100,4) wait_pct,
                       wait_class
                FROM v$system_event
                WHERE event NOT IN ('rdbms ipc message','smon timer','pmon','pipe get',
                                    'SQL*Net message from client','class slave wait')
                ORDER BY time_waited DESC
            ) WHERE ROWNUM <= 10
        """)
        results['wait_events'] = cur.fetchall()

        # SGA 组件
        cur.execute("""
            SELECT name, ROUND(bytes/1024/1024,2) size_mb
            FROM v$sgastat
            WHERE pool IS NOT NULL
            ORDER BY bytes DESC
        """)
        results['sga_pools'] = cur.fetchall()

        # PGA Target
        cur.execute("""
            SELECT a.name, ROUND(a.value/1024/1024,2) size_mb
            FROM v$pgastat a
            WHERE a.name IN ('aggregate PGA target parameter',
                             'total PGA allocated',
                             'total PGA inuse',
                             'maximum PGA allocated')
        """)
        results['pga'] = cur.fetchall()

        # 缓冲区命中率
        cur.execute("""
            SELECT name, ROUND((1 - physical_reads / (db_block_gets + consistent_gets)) * 100, 2) hit_pct,
                   db_block_gets, consistent_gets, physical_reads
            FROM v$buffer_pool_statistics
            WHERE db_block_gets + consistent_gets > 0
        """)
        results['buffer_hit'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_dataguard(conn):
    """Data Guard 配置"""
    results = {}
    cur = conn.cursor()
    try:
        # Data Guard 配置
        cur.execute("""
            SELECT GROUP#, TYPE, MEMBER, IS_RECOVERY_DEST_FILE
            FROM v$logfile
            WHERE TYPE = 'STANDBY'
        """)
        results['standby_logs'] = cur.fetchall()

        # 归档目的地
        cur.execute("""
            SELECT dest_id, status, destination, archiver, transmit_mode,
                   archiver, REGISTER
            FROM v$archive_dest
            WHERE destination IS NOT NULL
        """)
        results['archive_dest'] = cur.fetchall()

        # 实时查询（11g 兼容：v$archive_dest_status 无 STANDBY_DB_UNIQUE_NAME 列）
        cur.execute("""
            SELECT database_mode, recovery_mode, protection_mode
            FROM v$archive_dest_status
            WHERE status != 'INACTIVE'
        """)
        results['dg_status'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_backup(conn):
    """RMAN 备份信息（11g 兼容）"""
    results = {}
    cur = conn.cursor()
    try:
        # 最近备份（v$rman_backup_job_details 在 11g 中列可能不同，用 try/except 兼容）
        try:
            cur.execute("""
                SELECT session_key, INPUT_TYPE, STATUS,
                       TO_CHAR(START_TIME,'YYYY-MM-DD HH24:MI') start_t,
                       TO_CHAR(END_TIME,'YYYY-MM-DD HH24:MI') end_t,
                       OUTPUT_DEVICE_TYPE, elapsed_seconds
                FROM v$rman_backup_job_details
                WHERE end_time > SYSDATE - 30
                ORDER BY end_time DESC
            """)
        except Exception:
            # 11g 降级：使用确认存在的列
            cur.execute("""
                SELECT session_key, INPUT_TYPE, STATUS,
                       TO_CHAR(START_TIME,'YYYY-MM-DD HH24:MI') start_t,
                       TO_CHAR(END_TIME,'YYYY-MM-DD HH24:MI') end_t,
                       NULL, ELAPSED_SECONDS
                FROM v$rman_backup_job_details
                WHERE end_time > SYSDATE - 30
                ORDER BY end_time DESC
            """)
        results['rman_jobs'] = cur.fetchall()

        # 备份集（11g v$backup_set 只有 17 列，无 INPUT_BYTES/OUTPUT_BYTES/COMPRESSION_RATIO）
        # 使用 v$backup_set_details 获取字节统计，它 11g 就存在
        try:
            cur.execute("""
                SELECT SET_STAMP, DEVICE_TYPE,
                       ROUND(ORIGINAL_INPUT_BYTES/1024/1024/1024,2) input_gb,
                       ROUND(OUTPUT_BYTES/1024/1024/1024,2) output_gb,
                       COMPRESSION_RATIO
                FROM v$backup_set_details
                WHERE COMPLETION_TIME > SYSDATE - 30
            """)
        except Exception:
            # 再降级：只用 v$backup_set 的基本列
            cur.execute("""
                SELECT SET_STAMP, BACKUP_TYPE,
                       NULL, NULL, NULL
                FROM v$backup_set
                WHERE COMPLETION_TIME > SYSDATE - 30
            """)
        results['backup_sets'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_sga_pga(conn):
    """SGA/PGA 内存信息"""
    results = {}
    cur = conn.cursor()
    try:
        # SGA 动态组件
        cur.execute("""
            SELECT component, current_size/1024/1024 AS curr_mb,
                   min_size/1024/1024 AS min_mb,
                   user_specified_size/1024/1024 AS user_mb
            FROM v$sga_dynamic_components
            WHERE current_size > 0
            ORDER BY current_size DESC
        """)
        results['sga_components'] = cur.fetchall()

        # SGA 总计
        try:
            cur.execute("""
                SELECT SUM(value)/1024/1024 AS sga_total_mb
                FROM v$sga
            """)
            results['sga_total'] = cur.fetchall()
        except Exception:
            results['sga_total'] = []

        # PGA 统计
        cur.execute("""
            SELECT NAME, VALUE/1024/1024 AS value_mb
            FROM v$pgastat
            WHERE NAME IN (
                'total PGA allocated','total PGA inuse',
                'aggregate PGA target parameter','aggregate PGA auto target',
                'maximum PGA allocated','total freeable PGA memory'
            )
        """)
        results['pga_stats'] = cur.fetchall()

        # Memory Target / SGA Target / PGA Aggregate Target
        cur.execute("""
            SELECT NAME, VALUE, DISPLAY_VALUE, ISDEFAULT
            FROM v$parameter
            WHERE NAME IN (
                'memory_target','memory_max_target',
                'sga_target','pga_aggregate_target','sga_max_size'
            )
        """)
        results['memory_params'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_params(conn):
    """关键参数"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT NAME, VALUE, DISPLAY_VALUE, ISDEFAULT,
                   ISSES_MODIFIABLE, ISSYS_MODIFIABLE, DESCRIPTION
            FROM v$parameter
            WHERE NAME IN (
                'processes','sessions','open_cursors','db_block_size',
                'db_file_multiblock_read_count','db_writer_processes',
                'undo_retention','compatible','nls_characterset',
                'nls_nchar_characterset','job_queue_processes',
                'parallel_max_servers','audit_trail','recyclebin',
                'optimizer_mode','cursor_sharing','statistics_level',
                'control_file_record_keep_time','remote_login_passwordfile',
                'resource_manager_plan'
            )
            ORDER BY NAME
        """)
        results['params'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_undo(conn):
    """Undo 表空间信息"""
    results = {}
    cur = conn.cursor()
    try:
        # Undo 表空间基本信息（容错，10g 可能略有差异）
        try:
            cur.execute("""
                SELECT d.undo_tablespace,
                       r.retention,
                       ROUND(NVL(ts.used_bytes/1024/1024,0),2) AS used_mb,
                       ROUND(NVL(ts.tbs_bytes/1024/1024,0),2) AS total_mb,
                       u.exp_blks AS exp_undo_blks,
                       u.unexp_blks AS unexp_undo_blks,
                       u.blk_cnt AS undo_blk_cnt
                FROM (
                    SELECT UPPER(VALUE) AS undo_tablespace
                    FROM v$parameter WHERE NAME='undo_tablespace'
                ) d,
                (
                    SELECT UPPER(VALUE) AS retention
                    FROM v$parameter WHERE NAME='undo_retention'
                ) r,
                (
                    SELECT SUM(df.bytes) AS tbs_bytes,
                           SUM(df.bytes)-NVL(SUM(ff.free_bytes),0) AS used_bytes
                    FROM dba_data_files df
                    LEFT JOIN (
                        SELECT tablespace_name, SUM(bytes) AS free_bytes
                        FROM dba_free_space GROUP BY tablespace_name
                    ) ff ON df.tablespace_name = ff.tablespace_name
                    WHERE df.tablespace_name = (
                        SELECT UPPER(VALUE) FROM v$parameter WHERE NAME='undo_tablespace'
                    )
                ) ts,
                (
                    SELECT COUNT(*) AS exp_blks, 0 AS unexp_blks, COUNT(*) AS blk_cnt
                    FROM v$undostat
                    WHERE begin_time > SYSDATE-1 AND undoblks > 0
                ) u
            """)
            results['undo_info'] = cur.fetchall()
        except Exception as e:
            results['undo_info'] = []

        # Undo 段统计
        try:
            cur.execute("""
                SELECT status, COUNT(*) AS num_segments,
                       SUM(bytes)/1024/1024 AS total_mb
                FROM dba_undo_extents
                GROUP BY status
            """)
            results['undo_segments'] = cur.fetchall()
        except Exception:
            results['undo_segments'] = []
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_long_sql(conn):
    """长时间运行的 SQL（11g 兼容）"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT * FROM (
                SELECT sid, serial#, username, sql_id, opname,
                       sofar, totalwork,
                       ROUND(sofar/GREATEST(totalwork,0.001)*100,1) AS pct_complete,
                       elapsed_seconds, time_remaining
                FROM v$session_longops
                WHERE totalwork > 0 AND sofar < totalwork AND elapsed_seconds > 30
                ORDER BY elapsed_seconds DESC
            ) WHERE ROWNUM <= 10
        """)
        results['long_sql'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_blocking(conn):
    """阻塞会话检测（v$lock + v$session）"""
    results = {}
    cur = conn.cursor()
    try:
        # 阻塞链：查询被阻塞的会话及其阻塞源
        cur.execute("""
            SELECT
                b.sid AS blocked_sid,
                b.serial# AS blocked_serial,
                b.username AS blocked_user,
                b.machine AS blocked_machine,
                b.program AS blocked_program,
                b.sql_id AS blocked_sql_id,
                b.event AS blocked_event,
                b.seconds_in_wait AS sec_in_wait,
                bl.sid AS blocking_sid,
                bl.serial# AS blocking_serial,
                bl.username AS blocking_user,
                bl.machine AS blocking_machine,
                bl.program AS blocking_program,
                bl.sql_id AS blocking_sql_id,
                l.type AS lock_type,
                l.lmode AS lock_mode_held,
                l.request AS lock_mode_requested,
                o.owner || '.' || o.object_name AS locked_object
            FROM v$lock l
            JOIN v$session b ON b.sid = l.sid
            JOIN v$lock l2 ON l2.id1 = l.id1 AND l2.id2 = l.id2 AND l2.sid != l.sid
            JOIN v$session bl ON bl.sid = l2.sid
            LEFT JOIN dba_objects o ON o.object_id = l.id1
            WHERE l.request > 0 AND l2.lmode > 0 AND l2.request = 0
            ORDER BY b.seconds_in_wait DESC
        """)
        results['blocking_chain'] = cur.fetchall()[:20]
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_deadlock(conn):
    """死锁统计检测"""
    results = {}
    cur = conn.cursor()
    try:
        # 死锁统计（系统启动以来的累计值）
        cur.execute("""
            SELECT name, value
            FROM v$sysstat
            WHERE name LIKE '%deadlock%' OR name LIKE '%enqueue deadlock%'
            ORDER BY name
        """)
        results['deadlock_stats'] = cur.fetchall()

        # 当前锁等待 Top10（被阻塞的会话）
        cur.execute("""
            SELECT s.sid, s.serial#, s.username, s.status, s.event,
                   s.seconds_in_wait, s.sql_id,
                   l.type, l.lmode, l.request,
                   o.owner || '.' || o.object_name AS obj_name
            FROM v$session s
            JOIN v$lock l ON s.sid = l.sid
            LEFT JOIN dba_objects o ON o.object_id = l.id1
            WHERE s.wait_class != 'Idle' AND l.request > 0
            ORDER BY s.seconds_in_wait DESC
        """)
        results['lock_waiters'] = cur.fetchall()[:10]
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_long_trx(conn):
    """长事务检测（运行超过60秒的事务）"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT
                s.sid, s.serial#, s.username, s.machine, s.program,
                s.sql_id, s.status, s.event,
                t.start_date AS trx_start,
                ROUND((SYSDATE - t.start_date) * 86400) AS trx_seconds,
                t.used_ublk AS undo_blocks,
                t.used_urec AS undo_records,
                s.sql_exec_start
            FROM v$transaction t
            JOIN v$session s ON s.saddr = t.ses_addr
            WHERE (SYSDATE - t.start_date) * 86400 > 60
            ORDER BY t.start_date
        """)
        results['long_trx'] = cur.fetchall()[:20]
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_rac(conn):
    """RAC 检查"""
    results = {}
    cur = conn.cursor()
    try:
        # 实例列表
        cur.execute("""
            SELECT INST_ID, INSTANCE_NAME, HOST_NAME, STATUS, PARALLEL
            FROM gv$instance
        """)
        results['rac_instances'] = cur.fetchall()

        # CRS 资源状态（11gR2+）
        try:
            cur.execute("""
                SELECT name, TYPE, STATE, TARGET, host_name
                FROM gv$crs_resource
                WHERE TYPE IN ('database','service','listener')
                ORDER BY TYPE, name
            """)
            results['crs_resources'] = cur.fetchall()
        except Exception:
            results['crs_resources'] = []

        # OCR 备份
        try:
            cur.execute("""
                SELECT group_number, name, state, type, total_mb, free_mb,
                       ROUND(free_mb/total_mb*100,2) free_pct
                FROM v$asm_diskgroup
            """)
            results['asm_diskgroups'] = cur.fetchall()
        except Exception:
            results['asm_diskgroups'] = []
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_jobs(conn):
    """Scheduler Jobs / DBMS_JOBS"""
    results = {}
    cur = conn.cursor()
    try:
        # Scheduler Jobs
        cur.execute("""
            SELECT JOB_NAME, STATE, ENABLED, SCHEDULE_TYPE,
                   TO_CHAR(NEXT_RUN_DATE,'YYYY-MM-DD HH24:MI') next_run,
                   RUN_COUNT, FAILURE_COUNT
            FROM dba_scheduler_jobs
            WHERE owner NOT IN ('SYS','SYSTEM')
            ORDER BY owner, job_name
        """)
        results['scheduler_jobs'] = cur.fetchall()

        # 失败的后台作业
        cur.execute("""
            SELECT job, WHAT, LAST_DATE, NEXT_DATE, FAILURES, BROKEN
            FROM dba_jobs
            WHERE failures > 0 OR broken = 'Y'
        """)
        results['failed_jobs'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_flashback(conn):
    """闪回配置（11g 兼容）"""
    results = {}
    cur = conn.cursor()
    try:
        # 闪回配置（11g 兼容：oldest_flashback_time 在闪回未启用时可能报错）
        try:
            cur.execute("""
                SELECT flashback_on,
                       TO_CHAR(oldest_flashback_time,'YYYY-MM-DD HH24:MI') oldest_t
                FROM v$database
            """)
        except Exception:
            # OLDEST_FLASHBACK_TIME 不可用（闪回未启用或版本差异）
            cur.execute("""
                SELECT flashback_on, NULL
                FROM v$database
            """)
        results['flashback'] = cur.fetchall()

        # 回收站（CDB 下用 cdb_recyclebin）
        try:
            cur.execute("""
                SELECT r.owner, r.original_name, r.type,
                       ROUND(r.space * (SELECT value FROM v$parameter WHERE name='db_block_size')/1024/1024,2) mb,
                       r.can_undrop, r.can_purge
                FROM dba_recyclebin r
                ORDER BY mb DESC
            """)
            results['recyclebin'] = cur.fetchall()
        except Exception:
            # CDB 环境可能需要从 PDB 查询
            try:
                cur.execute("""
                    SELECT owner, original_name, type,
                           ROUND(space * (SELECT value FROM v$parameter WHERE name='db_block_size')/1024/1024,2) mb,
                           can_undrop, can_purge
                    FROM cdb_recyclebin
                    ORDER BY mb DESC
                """)
                results['recyclebin'] = cur.fetchall()
            except Exception:
                results['recyclebin'] = []
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_alert(conn, days=7):
    """最近 Alert 日志错误（11g 兼容：v$diag_alert_text 仅 12c+，11g 用 v$database_alert_log 或直接返回）"""
    results = {}
    cur = conn.cursor()
    try:
        # 12c+ 视图
        cur.execute("""
            SELECT TO_CHAR(alert_time,'YYYY-MM-DD HH24:MI:SS') t,
                   SUBSTR(message_text,1,200) message
            FROM v$diag_alert_text
            WHERE alert_time > SYSDATE - :days
              AND (message_text LIKE '%ORA-%' OR message_text LIKE '%ERROR%')
            ORDER BY alert_time DESC
        """, days=days)
        results['alert_errors'] = cur.fetchall()
    except Exception:
        # 11g：v$diag_alert_text 不存在，尝试 v$diag_alert_xml
        try:
            cur.execute("""
                SELECT TO_CHAR(trap_time,'YYYY-MM-DD HH24:MI:SS') t,
                       SUBSTR(message_text,1,200) message
                FROM v$diag_alert_xml
                WHERE trap_time > SYSDATE - :days
                  AND (message_text LIKE '%ORA-%' OR message_text LIKE '%ERROR%')
                ORDER BY trap_time DESC
            """, days=days)
            results['alert_errors'] = cur.fetchall()
        except Exception:
            # 11g 也没有 v$diag_alert_xml，记录提示信息
            results['alert_errors'] = []
            results['note'] = 'Oracle 11g 不支持 v$diag_alert_text/xml，请手动查看 alert_SID.log 文件'
    finally:
        cur.close()
    return results

# ═══════════════════════════════════════════════════════════════════════════
#                    报告生成（HTML）
# ═══════════════════════════════════════════════════════════════════════════

def _html_table(headers, rows, id_="", class_="dbcheck_tbl"):
    """生成 HTML 表格"""
    lines = []
    if id_ or class_:
        lines.append(f'<table id="{id_}" class="{class_}" border="1" cellpadding="4" cellspacing="0">')
    else:
        lines.append('<table border="1" cellpadding="4" cellspacing="0">')
    # 表头
    lines.append('<thead><tr>')
    for h in headers:
        lines.append(f'<th>{h}</th>')
    lines.append('</tr></thead>')
    # 数据行
    lines.append('<tbody>')
    for i, row in enumerate(rows):
        bg = '#FFFFFF' if i % 2 == 0 else '#F5F5F5'
        lines.append(f'<tr style="background:{bg}">')
        for cell in row:
            lines.append(f'<td>{str(cell) if cell is not None else ""}</td>')
        lines.append('</tr>')
    lines.append('</tbody></table>')
    return '\n'.join(lines)

def _html_section(title, content, anchor=""):
    anchor_tag = f'<a name="{anchor}"></a>' if anchor else ''
    return f'''
<h2 id="{anchor}">{anchor_tag}{title}</h2>
<div class="section">{content}</div>
'''

def _make_html_id(title):
    """将章节标题转为合法的 HTML id 属性值"""
    import re
    s = title.replace(' ', '_').replace('（', '_').replace('）', '_')
    s = re.sub(r'[^_\w\u4e00-\u9fff]', '', s)
    return s or 'section'


def build_html_report(db_info, os_data, check_results, db_version,
                      ai_advice='', inspector='', chapter_results=None):
    """构建完整 HTML 巡检报告"""
    from datetime import datetime

    # ── 样式 ────────────────────────────────────────────────────────────────
    css = """
    <style>
    * { box-sizing: border-box; margin: 0; padding: 0; }
    body { font-family: Consolas, 'Courier New', monospace; font-size: 13px;
           background: #f4f4f4; color: #222; padding: 20px; }
    h1 { font-size: 22px; color: #fff; background: #0066cc;
         padding: 16px 24px; border-radius: 6px; margin-bottom: 20px; }
    h2 { font-size: 16px; color: #fff; background: #0066cc;
         padding: 8px 14px; margin: 24px 0 10px; border-radius: 4px; }
    h3 { font-size: 14px; color: #336699; margin: 14px 0 6px; }
    table { width: 100%%; border-collapse: collapse; margin: 8px 0;
             background: #fff; font-size: 12px; }
    th { background: #336699; color: #fff; padding: 8px 10px; text-align: left; }
    td { padding: 6px 10px; border: 1px solid #ddd; vertical-align: top; }
    tr:nth-child(odd) { background: #fff; }
    tr:nth-child(even) { background: #f0f6ff; }
    .summary-grid { display: grid; grid-template-columns: repeat(auto-fill, minmax(280px, 1fr));
                    gap: 12px; margin: 12px 0; }
    .summary-card { background: #fff; border: 1px solid #ccc; border-radius: 6px;
                    padding: 12px 16px; box-shadow: 2px 2px 4px rgba(0,0,0,0.08); }
    .summary-card .label { font-size: 11px; color: #888; margin-bottom: 4px; }
    .summary-card .value { font-size: 16px; font-weight: bold; color: #0066cc; }
    .ok    { color: green; font-weight: bold; }
    .warn  { color: #cc6600; font-weight: bold; }
    .error { color: red; font-weight: bold; }
    .nav { background: #e8f0fe; padding: 10px 16px; border-radius: 4px;
            margin-bottom: 16px; font-size: 12px; }
    .nav a { margin-right: 14px; color: #0066cc; text-decoration: none; }
    .nav a:hover { text-decoration: underline; }
    .section { background: #fff; border: 1px solid #ddd; border-radius: 6px;
               padding: 14px; margin-top: 6px; }
    .footer { text-align: center; color: #888; font-size: 11px; margin-top: 30px; }
    </style>
    """

    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    # ── 汇总卡片 ────────────────────────────────────────────────────────────
    inst = check_results.get('instance', [])
    db   = db_info
    ver  = db_version

    hostname = os_data.get('hostname', '')
    uptime   = os_data.get('uptime', '')
    cpu      = f"{os_data.get('cpu_model','')} × {os_data.get('cpu_count','?')}"
    mem      = f"{os_data.get('mem_total_mb','?')} MB，使用率 {os_data.get('mem_usage_pct','?')}%"

    cards_html = f"""
    <div class="summary-grid">
        <div class="summary-card">
            <div class="label">主机名</div>
            <div class="value">{hostname}</div>
        </div>
        <div class="summary-card">
            <div class="label">数据库版本</div>
            <div class="value">{ver}</div>
        </div>
        <div class="summary-card">
            <div class="label">运行时间</div>
            <div class="value">{uptime}</div>
        </div>
        <div class="summary-card">
            <div class="label">CPU</div>
            <div class="value">{cpu}</div>
        </div>
        <div class="summary-card">
            <div class="label">内存</div>
            <div class="value">{mem}</div>
        </div>
        <div class="summary-card">
            <div class="label">巡检时间</div>
            <div class="value">{now}</div>
        </div>
    </div>
    """

    # ── OS 信息 ─────────────────────────────────────────────────────────────
    os_rows = []
    for k in ['hostname','os_version','kernel','uptime','cpu_model','cpu_count',
              'mem_total_mb','mem_used_mb','mem_usage_pct','swap_total_mb','swap_used_mb',
              'load_average','disk_usage','hugepages','thp']:
        v = os_data.get(k, 'N/A')
        os_rows.append((k, v))
    os_section = _html_section('🖥  OS 主机信息', _html_table(['项目','内容'], os_rows), 'os_info')

    # ── 统一章节渲染 ──────────────────────────────────────────────────────
    has_template = chapter_results is not None and len(chapter_results) > 0
    unified_chapters = build_unified_chapters(
        check_results, chapter_results, has_template, os_data=os_data, db_info=db_info
    )

    # 动态导航栏
    nav_links = []
    for ch in unified_chapters:
        ch_id = _make_html_id(ch['chapter_title'])
        nav_links.append(f'<a href="#{ch_id}">{ch["chapter_title"]}</a>')
    nav_links.append('<a href="#ai_diagnosis">AI诊断</a>')
    nav = f'<div class="nav">{"  ".join(nav_links)}</div>'

    # 统一章节 HTML 渲染（跳过 OS 章节，已在上方单独渲染）
    unified_html_parts = []
    for ch in unified_chapters:
        if ch['chapter_title'] == 'OS 主机信息':
            continue
        ch_id = _make_html_id(ch['chapter_title'])
        ch_html = f'<h2 id="{ch_id}"><a name="{ch_id}"></a>{ch["chapter_title"]}</h2><div class="section">'
        for subs in ch.get('subsections', []):
            if subs.get('title'):
                ch_html += f'<h3>{subs["title"]}</h3>'
            if subs.get('columns') and subs.get('data'):
                ch_html += _html_table(subs['columns'], subs['data'])
            elif subs.get('columns'):
                ch_html += '<p>无数据</p>'
        ch_html += '</div>'
        unified_html_parts.append(ch_html)
    unified_html = '\n'.join(unified_html_parts)

    # ── AI 诊断 ─────────────────────────────────────────────────────────────
    ai_section = ''
    if ai_advice:
        ai_lines = []
        for line in ai_advice.split('\n'):
            if line.startswith('# '):
                ai_lines.append(f'<h3>{line[2:]}</h3>')
            elif line.startswith('- ') or line.startswith('* '):
                ai_lines.append(f'<li>{line[2:]}</li>')
            elif re.match(r'^\d+\.', line):
                ai_lines.append(f'<li>{line}</li>')
            elif line.strip():
                ai_lines.append(f'<p>{line}</p>')
        if ai_lines:
            ai_section = _html_section('🤖  AI 诊断建议', '<br>'.join(ai_lines), 'ai_diagnosis')
    else:
        ai_section = _html_section('🤖  AI 诊断建议', '<p style="color:#888;">AI 诊断未启用或无可用建议。请在 dbc_config.json 中配置 Ollama 后重新巡检以获取 AI 诊断。</p>', 'ai_diagnosis')

    # ── 组合 ─────────────────────────────────────────────────────────────────
    body = (
        f'<h1>DBCheck Oracle 全面巡检报告 | {ver} | {now}</h1>'
        + (f'<p style="text-align:center;color:#0066cc;">巡检人: {inspector}</p>' if inspector else '')
        + nav
        + cards_html
        + os_section
        + unified_html
        + ai_section
        + f'<div class="footer">DBCheck Oracle 巡检工具 {VER} | 报告生成时间 {now}</div>'
    )

    return f"""<!DOCTYPE html>
<html lang="zh">
<head><meta charset="utf-8"><title>DBCheck Oracle 巡检报告</title>{css}</head>
<body>{body}</body>
</html>"""

# ═══════════════════════════════════════════════════════════════════════════
#                    报告生成（Word）
# ═══════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color):
    """设置单元格背景色"""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def _docx_table(doc, headers, rows, header_bg='336699'):
    """生成 Word 表格（带表头背景色）"""
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 表头
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            _set_cell_bg(hdr_cells[i], header_bg)
        except Exception:
            pass
    # 数据行（容错：列数不足时填充空白，多于 header 时截断）
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        for ci in range(len(headers)):
            cell_val = row[ci] if ci < len(row) else None
            if cell_val is not None:
                cells[ci].text = str(cell_val)
    return tbl

# ═══════════════════════════════════════════════════════════════════════════
#            统一章节数据结构转换器
# ═══════════════════════════════════════════════════════════════════════════

# check_results 中每个 (章节, sub_key) → 列名列表的映射表
_ORACLE_COLUMNS_MAP = {
    ('实例信息', 'instance'): ['Inst_ID', 'Inst_Name', 'Host', 'Version', 'Startup', 'Status', 'Parallel', 'Log_Mode', 'Role', 'Open_Mode'],
    ('版本/补丁', 'version'): ['BANNER'],
    ('表空间', 'data_tablespaces'): ['Tablespace', 'Status', 'Type', 'Logging', 'Max_MB', 'Cur_MB', 'Used_MB', 'Pct'],
    ('表空间', 'temp_tablespaces'): ['Tablespace', 'Status', 'Max_MB', 'Cur_MB', 'Used_MB', 'Pct'],
    ('Redo日志', 'logs'): ['Group#', 'Thread#', 'Sequence#', 'Size_MB', 'Status', 'Members', 'Archived'],
    ('控制文件', 'controlfiles'): ['Name', 'Status', 'Reuse', 'Block_Size', 'Blocks'],
    ('SGA/PGA内存', 'sga_components'): ['Component', 'Current_MB', 'Min_MB', 'User_MB'],
    ('SGA/PGA内存', 'sga_total'): ['SGA_Total'],
    ('SGA/PGA内存', 'pga_stats'): ['Name', 'Value'],
    ('SGA/PGA内存', 'memory_params'): ['Name', 'Value', 'Display_Value', 'Default', 'IsSES_Mod', 'IsSYS_Mod', 'Description'],
    ('关键参数', 'params'): ['Name', 'Value', 'Display_Value', 'Default', 'IsSES_Mod', 'IsSYS_Mod', 'Description'],
    ('Undo信息', 'undo_info'): ['Undo_Tablespace', 'Retention_Min', 'Used_MB', 'Total_MB', 'Committed', 'Uncommitted', 'Total_Blocks'],
    ('Undo信息', 'undo_segments'): ['Status', 'Count', 'Size'],
    ('长SQL', 'long_sql'): ['SID', 'Serial#', 'Username', 'SQL_ID', 'Operation', 'Done', 'Total_Work', 'Pct_Done', 'Elapsed', 'Remaining'],
    ('性能指标', 'session_by_status'): ['Status', 'Count'],
    ('性能指标', 'wait_events'): ['Event', 'Total_Waits', 'Time_S', 'Pct', 'Class'],
    ('性能指标', 'buffer_hit'): ['Pool', 'Hit_Pct', 'Block_Gets', 'Consistent_Gets', 'Physical_Reads'],
    ('Top SQL', 'top_sql_buffer_gets'): ['SQL_ID', 'SQL_Text', 'Buf_MB', 'Disk_MB', 'Executions', 'Elapsed_S', 'Gets_Per_Exec', 'Module'],
    ('Top SQL', 'top_sql_disk_reads'): ['SQL_ID', 'SQL_Text', 'Buf_MB', 'Disk_MB', 'Executions', 'Elapsed_S', 'Reads_Per_Exec', 'Module'],
    ('无效对象', 'invalid_by_type'): ['Owner', 'Object_Type', 'Count'],
    ('无效对象', 'invalid_detail'): ['Owner', 'Object_Name', 'Object_Type', 'Status', 'Last_DDL_Time'],
    ('用户安全', 'default_accounts'): ['Username', 'Status', 'Lock_Date', 'Expiry_Date', 'Created'],
    ('用户安全', 'locked_users'): ['Username', 'Status', 'Lock_Date', 'Expiry_Date'],
    ('用户安全', 'admin_roles'): ['Role', 'Grantee', 'Admin_Option'],
    ('备份信息', 'rman_jobs'): ['Session_Key', 'Type', 'Status', 'Start_Time', 'End_Time', 'Size_GB', 'Duration_Min'],
    ('闪回/回收站', 'flashback'): ['Flashback_On', 'Retention_Target', 'Oldest_FLASHBACK', 'Retention_Min'],
    ('闪回/回收站', 'recyclebin'): ['Owner', 'Original_Name', 'Type', 'Space_MB', 'Purgeable', 'Clearable'],
    ('Data Guard', 'dg_status'): ['Database_Mode', 'Recovery_Mode', 'Protection_Mode', 'Standby'],
    ('Data Guard', 'archive_dest'): ['Dest_ID', 'Status', 'Destination', 'Archiver', 'Transmit_Mode'],
    ('RAC+ASM', 'rac_instances'): ['Inst#', 'Inst_Name', 'Host', 'Status', 'Parallel'],
    ('RAC+ASM', 'crs_resources'): ['Name', 'Type', 'State', 'Target', 'Server'],
    ('RAC+ASM', 'asm_diskgroups'): ['Group_Number', 'Name', 'State', 'Type', 'Total_MB', 'Free_MB', 'Pct_Free'],
    ('AWR快照', 'awr_snaps'): ['Inst#', 'Snap_ID', 'Begin_Time', 'End_Time', 'Elapsed_HR', 'Errors'],
    ('作业调度', 'scheduler_jobs'): ['Job_Name', 'State', 'Enabled', 'Schedule_Type', 'Next_Run', 'Run_Count', 'Failure_Count'],
    ('作业调度', 'failed_jobs'): ['Job#', 'What', 'Last_Date', 'Next_Date', 'Failures', 'Broken'],
    ('Alert日志', 'alert_errors'): ['Time', 'Message'],
}

# 子标题映射：(章节, sub_key) → 子标题
_ORACLE_SUBTITLE_MAP = {
    ('表空间', 'data_tablespaces'): '永久表空间',
    ('表空间', 'temp_tablespaces'): '临时表空间',
    ('SGA/PGA内存', 'sga_components'): 'SGA 组件',
    ('SGA/PGA内存', 'sga_total'): 'SGA 总计',
    ('SGA/PGA内存', 'pga_stats'): 'PGA 统计',
    ('SGA/PGA内存', 'memory_params'): '内存参数',
    ('Undo信息', 'undo_info'): 'Undo 信息',
    ('Undo信息', 'undo_segments'): 'Undo 段统计',
    ('性能指标', 'session_by_status'): '会话状态',
    ('性能指标', 'wait_events'): 'Top10 等待事件',
    ('性能指标', 'buffer_hit'): '缓冲区命中率',
    ('Top SQL', 'top_sql_buffer_gets'): '按 Buffer Gets',
    ('Top SQL', 'top_sql_disk_reads'): '按磁盘读',
    ('用户安全', 'default_accounts'): '默认账户（高危）',
    ('用户安全', 'locked_users'): '锁定/过期用户',
    ('用户安全', 'admin_roles'): '带管理权限的角色',
    ('闪回/回收站', 'flashback'): '闪回配置',
    ('闪回/回收站', 'recyclebin'): '回收站',
    ('Data Guard', 'dg_status'): 'DG 状态',
    ('Data Guard', 'archive_dest'): '归档目的地',
    ('RAC+ASM', 'rac_instances'): '实例列表',
    ('RAC+ASM', 'crs_resources'): 'CRS 资源',
    ('RAC+ASM', 'asm_diskgroups'): 'ASM 磁盘组',
    ('作业调度', 'scheduler_jobs'): 'Scheduler Jobs',
    ('作业调度', 'failed_jobs'): '失败的后台作业',
}

# 特殊行转换函数：(章节, sub_key) → transform_fn(data_rows)
_ORACLE_ROW_TRANSFORMS = {
    ('版本/补丁', 'version'): lambda rows: [[r[0]] for r in rows],
    ('SGA/PGA内存', 'sga_total'): lambda rows: [[r[0]] for r in rows],
}


def build_unified_chapters(check_results, chapter_results, has_template, os_data=None, db_info=None):
    """将 check_results 或 chapter_results 转换为统一的章节结构。

    返回格式:
    [
        {chapter_title: str, subsections: [{title: str, columns: [str], data: [[str]]}]},
        ...
    ]

    排序规则：
    - 模板模式：严格按 chapter_results 的 sort_order 排序（由 SQL ORDER BY 保证）
    - 硬编码模式：OS 信息固定在前，其余按 check_results 顺序

    os_data: OS 信息 dict（硬编码模式下用于构建 OS 章节）
    db_info: 数据库基本信息 dict（当前未使用，保留兼容性）
    """
    unified = []
    os_data = os_data or {}
    db_info = db_info or {}

    # ── 模板驱动模式：严格按配置 sort_order 顺序 ───────────────────
    if has_template and chapter_results:
        for ch in chapter_results:
            subsections = []
            for qr in ch.get('queries_results', []):
                subs_title = qr.get('query_description_zh', qr.get('query_key', ''))
                columns = [str(c) for c in qr.get('columns', [])]
                data = [[str(d) if d is not None else '' for d in row] for row in qr.get('data', [])]
                error = qr.get('error')
                subsections.append({'title': subs_title, 'columns': columns, 'data': data, 'error': error})
            unified.append({
                'chapter_title': ch['chapter_title_zh'],
                'chapter_number': ch.get('chapter_number', 0),
                'subsections': subsections
            })
    else:
        # ── 硬编码回退模式：OS 信息在前 + check_results ──────────
        os_keys = ['hostname', 'os_version', 'kernel', 'uptime', 'cpu_model', 'cpu_count',
                   'mem_total_mb', 'mem_used_mb', 'mem_usage_pct', 'swap_total_mb', 'swap_used_mb',
                   'load_average', 'hugepages', 'thp']
        os_data_list = [[k, str(os_data.get(k, 'N/A'))] for k in os_keys]
        unified.append({'chapter_title': 'OS 主机信息', 'chapter_number': 0,
                        'subsections': [{'title': '', 'columns': ['项目', '内容'], 'data': os_data_list}]})

        for ch_name, ch_data in check_results.items():
            if not ch_data or not isinstance(ch_data, dict):
                continue
            subsections = []
            for sub_key, raw_rows in ch_data.items():
                if not raw_rows:
                    continue
                col_key = (ch_name, sub_key)
                columns = _ORACLE_COLUMNS_MAP.get(col_key, [])
                transform_fn = _ORACLE_ROW_TRANSFORMS.get(col_key)
                if transform_fn:
                    data = transform_fn(raw_rows)
                else:
                    data = raw_rows
                data = [[str(d) if d is not None else '' for d in row] for row in data]
                subs_title = _ORACLE_SUBTITLE_MAP.get(col_key, sub_key)
                subsections.append({'title': subs_title, 'columns': columns, 'data': data})
            if subsections:
                unified.append({'chapter_title': ch_name, 'chapter_number': 0, 'subsections': subsections})

    return unified


def extract_risks_from_unified(unified_chapters):
    """从统一章节数据中提取风险项，返回 (risk_items, summary_items) 列表。
    用于模板驱动模式下替代基于 check_results 硬编码键的风险评估。
    """
    risk_items = []
    summary_items = []

    for ch in unified_chapters:
        ch_title = ch['chapter_title']
        for subs in ch.get('subsections', []):
            columns = subs.get('columns', [])
            data = subs.get('data', [])
            if not data:
                continue

            # 表空间风险：通过列名中包含 'Pct' 或 '%' 且数据列数>=8 来识别
            pct_col_idx = None
            name_col_idx = None
            for i, col in enumerate(columns):
                col_s = str(col).lower()
                if 'pct' in col_s or '%' in col_s or '使用率' in col_s or 'used_pct' in col_s:
                    pct_col_idx = i
                if 'name' in col_s or '表空间' in col_s or 'tablespace' in col_s:
                    name_col_idx = i

            if pct_col_idx is not None:
                for row in data:
                    if len(row) > pct_col_idx:
                        try:
                            pct_val = float(row[pct_col_idx]) if str(row[pct_col_idx]).replace('.', '').replace('-', '').isdigit() else 0
                            row_name = row[name_col_idx] if name_col_idx is not None and len(row) > name_col_idx else 'N/A'
                            if pct_val > 90:
                                risk_items.append(('tablespace', row_name, 'high', f'表空间使用率 {pct_val:.1f}%'))
                            elif pct_val > 80:
                                risk_items.append(('tablespace', row_name, 'mid', f'表空间使用率 {pct_val:.1f}%'))
                        except (ValueError, TypeError):
                            pass

            # 无效对象风险：列名包含 'count' 且数据 > 0
            for i, col in enumerate(columns):
                col_s = str(col).lower()
                if ('count' in col_s or '数量' in col_s) and 'invalid' in ch_title.lower() + str(subs.get('title', '')).lower():
                    for row in data:
                        if len(row) > i:
                            try:
                                cnt = int(row[i])
                                if cnt > 0:
                                    row_type = row[1] if len(row) > 1 else 'N/A'
                                    risk_items.append(('invalid_obj', row_type, 'mid', f'无效对象数量: {cnt}'))
                            except (ValueError, TypeError):
                                pass

            # Alert 日志风险：章节标题包含 'alert' 或 '告警'
            if ('alert' in ch_title.lower() or '告警' in ch_title) and data:
                risk_items.append(('alert', ch_title, 'high', f'发现 {len(data)} 条告警'))

    return risk_items, summary_items


# ═══════════════════════════════════════════════════════════════════════════
#                    报告生成（Word）
# ═══════════════════════════════════════════════════════════════════════════

def build_word_report(db_info, os_data, check_results, db_version, ai_advice='', inspector='', lang='zh', desensitize=False, config_baseline_result=None, index_health_result=None, host='', health_status='', chapter_results=None):
    """构建完整 Word 巡检报告（纯 python-docx，无模板依赖）"""
    if not _HAS_DOCX:
        return None

    def _t(key):
        try:
            from i18n import t
            return t(key, lang)
        except Exception:
            return key

    # ── 脱敏处理（IP / 端口 / 用户名 / 服务名 / 主机名）───────────
    if desensitize:
        try:
            from desensitize import apply_desensitization
            _ds = apply_desensitization
            db_info = _ds({'db_info': db_info})['db_info']
            os_data  = _ds({'system_info': os_data})['system_info']
        except Exception:
            pass

    doc = Document()

    # ── 页面设置 ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

    # ── 封面 ────────────────────────────────────────────────────────────────
    # Logo 图片
    logo_path = os.path.join(os.path.dirname(__file__), 'dbcheck_logo.png')
    if os.path.exists(logo_path):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture(logo_path, width=Cm(3.5))

    # 报告标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(_t('report.oracle_title'))
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 75, 135)
    title_run.font.name = '微软雅黑'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 英文副标题
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run('Database Health Inspection Report')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    subtitle_run.font.italic = True
    subtitle_run.font.name = 'Times New Roman'

    doc.add_paragraph()  # 空行

    # 封面信息表格
    report_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    boot_time = os_data.get('uptime', 'Unknown')
    hs = health_status if health_status else 'Unknown'
    cover_labels = ['服务器地址', '实例启动时间', '巡检结果', '巡检人员', '报告生成时间'] if lang == 'zh' else ['Server Address', 'Instance Start Time', 'Status', 'Inspector', 'Report Time']
    cover_data = [
        f"{host}" if host else "N/A",
        boot_time,
        hs,
        inspector if inspector else 'Jack',
        report_time
    ]
    tbl = doc.add_table(rows=len(cover_labels), cols=2, style='Table Grid')
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(zip(cover_labels, cover_data)):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = str(value)
        for cell in tbl.rows[i].cells:
            cell.paragraphs[0].runs[0].font.name = '微软雅黑'
            cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
            if cell == tbl.rows[i].cells[0]:
                _set_cell_bg(cell, '336699')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # ── 章节序号计数器（配置章节 + 固定附录共用，序号连续）───────
    _section_num = 0
    _subsection_num = 0

    def _add_section(title):
        nonlocal _section_num, _subsection_num
        _section_num += 1
        _subsection_num = 0
        numbered_title = f"{_section_num} {title}"
        h = doc.add_heading(numbered_title, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)   # 深蓝，一级标题
            run.font.size = Pt(14)

    def _add_subsection(title):
        """二级子标题（自动编号：X.Y）"""
        nonlocal _subsection_num
        _subsection_num += 1
        numbered_title = f"{_section_num}.{_subsection_num} {title}"
        h = doc.add_heading(numbered_title, level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 102, 204)  # 蓝色，二级标题
            run.font.size = Pt(12)

    def _add_inline_runs(p, text):
        """向段落中添加文本_run，支持 **加粗** 和 `行内代码` 格式"""
        import re
        # 匹配顺序：**加粗**、`行内代码`、其他普通文本
        # 用占位符保护代码段避免双重匹配
        segments = []
        last = 0
        for m in re.finditer(r'\*\*(.+?)\*\*|`([^`]+)`', text):
            if m.start() > last:
                segments.append(('plain', text[last:m.start()]))
            if m.group(0).startswith('**'):
                segments.append(('bold', m.group(1)))
            else:
                segments.append(('code', m.group(2)))
            last = m.end()
        if last < len(text):
            segments.append(('plain', text[last:]))

        for seg_type, seg_text in segments:
            run = p.add_run(seg_text)
            run.font.size = Pt(10.5)
            if seg_type == 'bold':
                run.bold = True
            elif seg_type == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 102, 204)

    def _render_ai_advice(doc, text):
        """将 Markdown 格式的 AI 诊断建议渲染为 Word 格式"""
        lines = text.split('\n')
        in_code_block = False
        code_buf = []
        h2_counter = 0   # 用于 ## 标题序号
        h3_counter = 0   # 用于 ### 小节序号（在同一 ## 内递增）
        prev_was_content = False  # 上一行是否渲染了内容（排除空行和纯分隔线）

        for raw in lines:
            stripped = raw.strip()

            # 代码块开始/结束
            if stripped.startswith('```'):
                if in_code_block:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.6)
                    for cl in code_buf:
                        r = p.add_run(cl + '\n')
                        r.font.name = 'Courier New'
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(0, 128, 0)
                    code_buf = []
                    in_code_block = False
                    prev_was_content = True
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_buf.append(stripped)
                continue

            # 空行：只打一个间距，不渲染独立空段落
            if not stripped:
                prev_was_content = False
                continue

            # ── 标题处理（>0 个 # + 空格 + 标题文字）─────────────
            heading_match = re.match(r'^(#{1,3})\s+(.*)', stripped)
            if heading_match:
                hashes, title_text = heading_match.groups()
                h_count = len(hashes)

                if h_count == 1:          # # 一级标题（章）
                    h = doc.add_heading(title_text, level=2)
                    for run in h.runs:
                        run.font.color.rgb = RGBColor(0, 102, 204)
                        run.font.size = Pt(12)
                    h2_counter = 0
                    h3_counter = 0

                elif h_count == 2:        # ## 二级标题（24.1/24.2）
                    h2_counter += 1
                    h3_counter = 0
                    h = doc.add_heading(f"24.{h2_counter} {title_text}", level=2)
                    for run in h.runs:
                        run.font.color.rgb = RGBColor(0, 102, 204)
                        run.font.size = Pt(12)

                elif h_count == 3:        # ### 四级标题 → 加粗普通段落
                    p = doc.add_paragraph()
                    run = p.add_run(title_text)
                    run.font.bold = True
                    run.font.size = Pt(10.5)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(2)
                    prev_was_content = True
                    continue

                prev_was_content = False   # 标题不计入内容行
                continue

            # 水平线 → 分隔段落
            if re.match(r'^[-*_]{3,}$', stripped):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                prev_was_content = True
                continue

            # 引用块
            if stripped.startswith('> '):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                _add_inline_runs(p, stripped[2:])
                for run in p.runs:
                    run.font.color.rgb = RGBColor(96, 96, 96)
                    run.font.italic = True
                prev_was_content = True
                continue

            # 无序列表
            if stripped.startswith('- ') or stripped.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                _add_inline_runs(p, stripped[2:])
                prev_was_content = True
                continue

            # 有序列表（数字+.）
            m = re.match(r'^(\d+)\.\s*(.*)', stripped)
            if m:
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                _add_inline_runs(p, m.group(2))
                prev_was_content = True
                continue

            # 普通段落（含行内格式）
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            _add_inline_runs(p, stripped)
            prev_was_content = True

        doc.add_paragraph()

    def _add_kv_table(data, cols=2):
        rows = list(data)
        headers = [_t('report.tbl_col_key'), _t('report.tbl_col_val')]
        _docx_table(doc, headers, rows)
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    #  统一章节渲染：基于 build_unified_chapters() 输出
    #  支持模板驱动（chapter_results）和硬编码回退（check_results）
    # ═══════════════════════════════════════════════════════════════════════════
    has_template = chapter_results is not None and len(chapter_results) > 0
    unified_chapters = build_unified_chapters(
        check_results, chapter_results, has_template, os_data=os_data, db_info=db_info
    )

    def _render_chapter_table(doc, columns, data, error=None):
        """辅助函数：渲染一个数据表格，数据为空时显示提示，查询失败时显示错误"""
        if columns and data:
            _docx_table(doc, columns, data)
            doc.add_paragraph()
        elif columns:
            p = doc.add_paragraph(_t('report.no_data'))
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(128, 128, 128)
            doc.add_paragraph()
        elif error:
            p = doc.add_paragraph(f"查询失败: {error}")
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(204, 0, 0)
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(_t('report.no_data'))
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(128, 128, 128)
            doc.add_paragraph()

    # ── 统一渲染所有章节（严格按照配置 sort_order 顺序）─────────────────
    for ch in unified_chapters:
        ch_title = ch['chapter_title']
        subsections = ch.get('subsections', [])
        if not subsections:
            continue

        _add_section(ch_title)

        for subs in subsections:
            subs_title = subs.get('title', '')
            if subs_title:
                _add_subsection(subs_title)
            columns = subs.get('columns', [])
            data = subs.get('data', [])
            error = subs.get('error')
            _render_chapter_table(doc, columns, data, error=error)

    # ── 风险与建议 ─────────────────────────────────────────────────────────
    _add_section(_t('report.oracle_sec_risks'))
    risk_items = []

    if has_template:
        # ── 模板驱动模式：从统一章节数据中提取风险 ──────────────────────────
        unified_risks, _ = extract_risks_from_unified(unified_chapters)
        sev_map = {'high': (_t('report.risk_high'), _t('report.severity_high')),
                   'mid': (_t('report.risk_mid'), _t('report.severity_mid')),
                   'low': (_t('report.risk_low'), _t('report.severity_low'))}
        for rtype, rname, sev, desc in unified_risks:
            sev_label, sev_display = sev_map.get(sev, sev_map['low'])
            risk_items.append({
                'col1': f'{rtype}/{rname}', 'col2': sev_label,
                'col3': desc, 'col4': sev_display, 'col5': _t('report.risk_dba'), 'fix_sql': ''
            })
    else:
        # ── 硬编码回退模式：从 check_results 提取风险（保持原有逻辑）───────
        # 从表空间数据中提取高使用率风险
        ts = check_results.get('表空间', {})
        for row in ts.get('data_tablespaces', []):
            if len(row) >= 8:
                try:
                    used_pct = float(row[7]) if row[7] != '-' else 0
                    if used_pct > 90:
                        risk_items.append({
                            'col1': _t('report.risk_tablespace').format(name=row[0]), 'col2': _t('report.risk_high'),
                            'col3': _t('report.risk_ts_high').format(pct=f'{used_pct:.1f}'),
                            'col4': _t('report.severity_high'), 'col5': _t('report.risk_dba'), 'fix_sql': _t('report.risk_fix_ts')
                        })
                    elif used_pct > 80:
                        risk_items.append({
                            'col1': _t('report.risk_tablespace').format(name=row[0]), 'col2': _t('report.risk_mid'),
                            'col3': _t('report.risk_ts_mid').format(pct=f'{used_pct:.1f}'),
                            'col4': _t('report.severity_mid'), 'col5': _t('report.risk_dba'), 'fix_sql': ''
                        })
                except (ValueError, TypeError):
                    pass
        # 从无效对象数据中提取风险
        io = check_results.get('无效对象', {})
        for row in io.get('invalid_by_type', []):
            if len(row) >= 3:
                cnt = row[2] if isinstance(row[2], int) else 0
                if cnt > 0:
                    risk_items.append({
                        'col1': _t('report.risk_invalid_obj').format(type=row[1]), 'col2': _t('report.risk_mid'),
                        'col3': _t('report.risk_invalid_desc').format(cnt=cnt, type=row[1]),
                        'col4': _t('report.severity_mid'), 'col5': _t('report.risk_dba'), 'fix_sql': f'SELECT * FROM {row[0]}.{row[1]} WHERE status=\'INVALID\';'
                    })
        # 从锁定用户中提取风险
        users = check_results.get('用户安全', {})
        locked = users.get('locked_users', [])
        if locked:
            risk_items.append({
                'col1': _t('report.risk_locked'), 'col2': _t('report.risk_mid'),
                'col3': _t('report.risk_locked_desc').format(n=len(locked)),
                'col4': _t('report.severity_mid'), 'col5': _t('report.risk_dba'), 'fix_sql': f"-- {_t('report.risk_fix_locked')}: SELECT username, lock_date FROM dba_users WHERE account_status LIKE '%LOCKED%';"
            })
        # 从Alert日志错误中提取风险
        alert = check_results.get('Alert日志', {})
        err_count = len(alert.get('alert_errors', []))
        if err_count > 0:
            risk_items.append({
                'col1': _t('report.risk_alert'), 'col2': _t('report.risk_high'),
                'col3': _t('report.risk_alert_desc').format(n=err_count),
                'col4': _t('report.severity_high'), 'col5': _t('report.risk_dba'), 'fix_sql': f"-- {_t('report.risk_fix_alert')}"
            })

    if risk_items:
        # 23.1 问题明细
        _add_subsection(_t('report.oracle_risk_detail_chapter'))
        risk_table_data = [(str(i+1), x['col1'], x['col2'], x['col3'], x['col4'], x['col5']) for i, x in enumerate(risk_items)]
        _docx_table(doc, [_t('report.col_seq'), _t('report.col_risk_item'), _t('report.col_level'), _t('report.col_desc'), _t('report.col_severity'), _t('report.col_owner')], risk_table_data)
        doc.add_paragraph()
        # 23.2 修复SQL速查
        fix_sqls = [(x['col1'], x['fix_sql']) for x in risk_items if x['fix_sql']]
        if fix_sqls:
            _add_subsection(_t('report.oracle_risk_fix_chapter'))
            for fname, sql in fix_sqls:
                p = doc.add_paragraph()
                p.add_run(f'【{fname}】').bold = True
                doc.add_paragraph(sql, style='List Bullet')
            doc.add_paragraph()
    else:
        # 即使无高风险，也汇总各项巡检结论
        _add_subsection(_t('report.oracle_risk_summary_chapter'))
        summary_items = []
        if has_template:
            # 模板模式：从统一章节中统计汇总信息
            for ch in unified_chapters:
                ch_title = ch['chapter_title']
                total_rows = sum(len(subs.get('data', [])) for subs in ch.get('subsections', []))
                if total_rows > 0:
                    summary_items.append(f'{ch_title}: 共 {total_rows} 条记录')
        else:
            # 硬编码回退模式
            ts = check_results.get('表空间', {})
            if ts.get('data_tablespaces'):
                total_ts = len(ts['data_tablespaces'])
                high_ts = sum(1 for r in ts['data_tablespaces'] if len(r) >= 8 and str(r[7]).replace('.','').isdigit() and float(r[7]) > 80)
                summary_items.append(_t('report.summary_ts').format(total=total_ts, high=high_ts))
            perf = check_results.get('性能指标', {})
            if perf.get('wait_events'):
                top_wait = perf['wait_events'][0][0] if perf['wait_events'] else 'N/A'
                summary_items.append(_t('report.summary_wait').format(event=top_wait))
            io = check_results.get('无效对象', {})
            if io.get('invalid_by_type'):
                total_inv = sum(int(r[2]) for r in io['invalid_by_type'] if len(r) >= 3 and str(r[2]).isdigit())
                summary_items.append(_t('report.summary_invalid').format(total=total_inv))
        if summary_items:
            for item in summary_items:
                p = doc.add_paragraph(item, style='List Bullet')
        else:
            p = doc.add_paragraph(_t('report.no_risk_found_oracle'))
            for run in p.runs:
                run.font.size = Pt(10.5)

    # ── 锁诊断（P0）──────────────────────────────────────────────────────────
    blocking_data = check_results.get('阻塞会话', {})
    deadlock_data = check_results.get('死锁检测', {})
    long_trx_data = check_results.get('长事务', {})
    blocking_rows = blocking_data.get('blocking_chain', [])
    deadlock_stats = deadlock_data.get('deadlock_stats', [])
    lock_waiters = deadlock_data.get('lock_waiters', [])
    long_trx_rows = long_trx_data.get('long_trx', [])

    has_lock_data = blocking_rows or deadlock_stats or lock_waiters or long_trx_rows
    if has_lock_data:
        _add_section(_t('report.oracle_lock_chapter'))

        # -- 阻塞会话分析 --
        _add_subsection(_t('report.oracle_sec_blocking'))
        if blocking_rows:
            blocking_headers = [
                _t('report.oracle_col_blocked_sid'), _t('report.oracle_col_serial'),
                _t('report.oracle_col_username'), _t('report.oracle_col_event'),
                _t('report.oracle_col_wait_sec'), _t('report.oracle_col_blocking_sid'),
                _t('report.oracle_col_blocking_username'),
                _t('report.oracle_col_lock_type'), _t('report.oracle_col_locked_obj'),
            ]
            blocking_display = []
            for r in blocking_rows:
                blocking_display.append([
                    str(r[0]) if len(r) > 0 else '',        # blocked_sid
                    str(r[1]) if len(r) > 1 else '',        # blocked_serial
                    str(r[2]) if len(r) > 2 else '',        # blocked_user
                    str(r[6]) if len(r) > 6 else '',        # blocked_event
                    str(r[7]) if len(r) > 7 else '',        # sec_in_wait
                    str(r[8]) if len(r) > 8 else '',        # blocking_sid
                    str(r[10]) if len(r) > 10 else '',      # blocking_user
                    str(r[14]) if len(r) > 14 else '',      # lock_type
                    str(r[17]) if len(r) > 17 else '',      # locked_object
                ])
            _docx_table(doc, blocking_headers, blocking_display)
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(_t('report.oracle_no_blocking'))
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(128, 128, 128)

        # -- 死锁检测 --
        _add_subsection(_t('report.oracle_sec_deadlock'))
        if deadlock_stats:
            p = doc.add_paragraph(_t('report.oracle_deadlock_stats_title'))
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.bold = True
            deadlock_stats_display = []
            for r in deadlock_stats:
                deadlock_stats_display.append([
                    str(r[0]) if len(r) > 0 else '',
                    str(r[1]) if len(r) > 1 else '',
                ])
            _docx_table(doc,
                [_t('report.oracle_col_stat_name'), _t('report.oracle_col_stat_value')],
                deadlock_stats_display)
            doc.add_paragraph()

        if lock_waiters:
            p = doc.add_paragraph(_t('report.oracle_lock_waiters_title'))
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.bold = True
            waiters_display = []
            for r in lock_waiters:
                waiters_display.append([
                    str(r[0]) if len(r) > 0 else '',        # sid
                    str(r[1]) if len(r) > 1 else '',        # serial#
                    str(r[2]) if len(r) > 2 else '',        # username
                    str(r[4]) if len(r) > 4 else '',        # event
                    str(r[5]) if len(r) > 5 else '',        # seconds_in_wait
                    str(r[7]) if len(r) > 7 else '',        # lock type
                    str(r[10]) if len(r) > 10 else '',      # locked object
                ])
            _docx_table(doc,
                [_t('report.oracle_col_sid'), _t('report.oracle_col_serial'),
                 _t('report.oracle_col_username'), _t('report.oracle_col_event'),
                 _t('report.oracle_col_wait_sec'), _t('report.oracle_col_lock_type'),
                 _t('report.oracle_col_locked_obj')],
                waiters_display)
            doc.add_paragraph()

        if not deadlock_stats and not lock_waiters:
            p = doc.add_paragraph(_t('report.oracle_no_deadlock'))
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(128, 128, 128)

        # -- 长事务检测 --
        _add_subsection(_t('report.oracle_sec_long_trx'))
        if long_trx_rows:
            long_trx_display = []
            for r in long_trx_rows:
                long_trx_display.append([
                    str(r[0]) if len(r) > 0 else '',        # sid
                    str(r[1]) if len(r) > 1 else '',        # serial#
                    str(r[2]) if len(r) > 2 else '',        # username
                    str(r[8]) if len(r) > 8 else '',        # trx_start
                    str(r[9]) if len(r) > 9 else '',        # trx_seconds
                    str(r[10]) if len(r) > 10 else '',      # undo_blocks
                ])
            _docx_table(doc,
                [_t('report.oracle_col_sid'), _t('report.oracle_col_serial'),
                 _t('report.oracle_col_username'), _t('report.oracle_col_trx_start'),
                 _t('report.oracle_col_trx_seconds'), _t('report.oracle_col_undo_blocks')],
                long_trx_display)
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(_t('report.oracle_no_long_trx'))
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(128, 128, 128)

    # ── 第24章 配置基线与索引健康 ────────────────────────────────────────
    # ── 24.1 配置基线检查 ─────────────────────────────────────────────────
    _add_section(_t('report.oracle_sec_config_baseline'))
    if config_baseline_result:
        cb = config_baseline_result
        summary = cb.get('summary', {})
        crit = summary.get('critical_count', 0)
        warn = summary.get('warning_count', 0)
        info = summary.get('info_count', 0)

        # 摘要行
        p = doc.add_paragraph()
        run = p.add_run(f"⚠ Critical: {crit}  |  ⚡ Warning: {warn}  |  ℹ Info: {info}")
        run.font.size = Pt(10.5)
        run.font.bold = True
        doc.add_paragraph()

        # 差距表格
        items = cb.get('items', [])
        if items:
            # 按严重程度排序
            severity_order = {'critical': 0, 'warning': 1, 'info': 2}
            items_sorted = sorted(items, key=lambda x: (severity_order.get(x.get('severity', 'info'), 2), x.get('param', '')))
            headers = [
                _t('report.tbl_col_key'),
                _t('report.oracle_col_current'),
                _t('report.oracle_col_recommended'),
                _t('report.oracle_col_gap'),
                _t('report.risk_level')
            ]
            rows = []
            for item in items_sorted:
                sev = item.get('severity', 'info')
                gap_str = f"{item.get('gap', '')} ({item.get('gap_pct', 0):.0f}%)"
                sev_str = {'critical': _t('report.risk_high'), 'warning': _t('report.risk_mid'), 'info': _t('report.risk_low')}.get(sev, sev)
                rows.append([
                    item.get('param', ''),
                    item.get('current', ''),
                    item.get('recommended', ''),
                    gap_str,
                    sev_str,
                ])
            _docx_table(doc, headers, rows)
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(_t('report.no_risk_found_oracle'))
            run = p.runs[0] if p.runs else p.add_run(_t('report.no_risk_found_oracle'))
            run.font.size = Pt(10.5)
            doc.add_paragraph()
    else:
        p = doc.add_paragraph(_t('report.no_config_baseline'))
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(128, 128, 128)
        doc.add_paragraph()

    # ── 24.2 索引健康分析 ─────────────────────────────────────────────────
    _add_section(_t('report.oracle_sec_index_health'))
    if index_health_result:
        ih = index_health_result
        sm = ih.get('summary', {})
        miss = sm.get('missing_count', 0)
        redun = sm.get('redundant_count', 0)
        unused = sm.get('unused_count', 0)

        # 摘要
        p = doc.add_paragraph()
        run = p.add_run(f"⚠ 缺失: {miss}  |  🔁 冗余: {redun}  |  💤 未使用: {unused}")
        run.font.size = Pt(10.5)
        run.font.bold = True
        doc.add_paragraph()

        # 未使用索引
        unused_list = ih.get('unused_indexes', [])
        if unused_list:
            _add_subsection(_t('report.oracle_idx_unused'))
            headers = [_t('report.oracle_col_schema'), _t('report.oracle_col_table'),
                       _t('report.oracle_col_idx_name'), _t('report.oracle_col_days_unused'),
                       _t('report.oracle_col_recommendation')]
            rows = [[i.get('table_schema', ''), i.get('table_name', ''), i.get('index_name', ''),
                     str(i.get('days_unused', '')), i.get('recommendation', '')] for i in unused_list[:50]]
            _docx_table(doc, headers, rows)
            doc.add_paragraph()

        # 冗余索引
        redun_list = ih.get('redundant_indexes', [])
        if redun_list:
            _add_subsection(_t('report.oracle_idx_redundant'))
            headers = [_t('report.oracle_col_schema'), _t('report.oracle_col_table'),
                       _t('report.oracle_col_idx_name'), _t('report.oracle_col_idx_name2'),
                       _t('report.oracle_col_column')]
            rows = [[i.get('table_schema', ''), i.get('table_name', ''), i.get('index_name', ''),
                     i.get('index_name2', ''), i.get('column_name', '')] for i in redun_list[:50]]
            _docx_table(doc, headers, rows)
            doc.add_paragraph()

        # 缺失索引
        miss_list = ih.get('missing_indexes', [])
        if miss_list:
            _add_subsection(_t('report.oracle_idx_missing'))
            headers = [_t('report.oracle_col_schema'), _t('report.oracle_col_table'),
                       _t('report.oracle_col_column'), _t('report.oracle_col_recommendation')]
            rows = [[i.get('table_schema', ''), i.get('table_name', ''), i.get('column_name', ''),
                     i.get('recommendation', '')] for i in miss_list[:50]]
            _docx_table(doc, headers, rows)
            doc.add_paragraph()

        if not unused_list and not redun_list and not miss_list:
            p = doc.add_paragraph(_t('report.no_risk_found_oracle'))
            for run in p.runs:
                run.font.size = Pt(10.5)
            doc.add_paragraph()
    else:
        p = doc.add_paragraph(_t('report.no_index_health'))
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(128, 128, 128)
        doc.add_paragraph()

    # ── AI 诊断章节（仅在启用且有诊断结果时生成，序号接续前面章节）───────
    if ai_advice:
        _add_section(_t('report.oracle_sec_ai'))
        _render_ai_advice(doc, ai_advice)



    # ── 页脚 ────────────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(_t('report.oracle_footer').format(VER=VER, time=time.strftime("%Y-%m-%d %H:%M:%S")))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc

# ═══════════════════════════════════════════════════════════════════════════
#                    主入口
# ═══════════════════════════════════════════════════════════════════════════

def print_banner():
    # Detect language
    try:
        from i18n import get_lang, t as _tt
        _lang = get_lang()
    except Exception:
        _lang = 'zh'

    def _t(key):
        try:
            return _tt(key, _lang)
        except Exception:
            return key

    banner_tool = _t('oracle_banner_tool')
    banner_sub  = _t('oracle_banner_subtitle')
    art = f"""
{CYAN}{BOLD}  ██████╗ ██████╗  ██████╗██╗  ██╗███████╗ ██████╗██╗  ██╗
  ██╔══██╗██╔══██╗██╔════╝██║  ██║██╔════╝██╔════╝██║ ██╔╝
  ██║  ██║██████╔╝██║     ███████║█████╗  ██║     █████╔╝
  ██║  ██║██╔══██╗██║     ██╔══██║██╔══╝  ██║     ██╔═██╗
  ██████╔╝██████╔╝╚██████╗██║  ██║███████╗╚██████╗██║  ██╗
  ╚═════╝ ╚═════╝  ╚═════╝╚═╝  ╚═╝╚══════╝ ╚═════╝╚═╝  ╚═╝{RESET}
{BOLD}      🗄️  {banner_tool}  {VER}{RESET}
{DIM}  ──────────────────────────────────────────────────────────{RESET}
{GREEN}  {banner_sub}{RESET}
{DIM}  ──────────────────────────────────────────────────────────{RESET}
"""
    print(art)

def single_inspection(args):
    """单机巡检主流程"""
    import paramiko

    # 返回值：供 Web UI 调用 run_full_analysis 使用
    context = {}

    # 标记是否由 getData 调用（Web UI 模式），跳过 report/history
    _web_ui_mode = getattr(args, '_web_ui_mode', False)
    # Get language and local _t for this function
    try:
        from i18n import get_lang
        _lang = get_lang()
    except Exception:
        _lang = 'zh'

    def _t(key):
        try:
            from i18n import t as _tt
            return _tt(key, _lang)
        except Exception:
            return key

    def _plural(cnt, singular, plural):
        """Return singular or plural form. EN: 1→singular else→plural; ZH: always singular."""
        return singular if _lang == 'en' and cnt == 1 else plural

    # Map internal Chinese check-item keys to i18n display names
    _item_i18n = {
        '实例信息':      'oracle_check_item_instance',
        '数据库信息':    'oracle_check_item_database',
        '版本/补丁':     'oracle_check_item_version',
        '表空间':        'oracle_check_item_tablespace',
        'Redo日志':      'oracle_check_item_redolog',
        '控制文件':      'oracle_check_item_controlfile',
        'SGA/PGA内存': 'oracle_check_item_sga_pga',
        '关键参数':      'oracle_check_item_params',
        'Undo信息':      'oracle_check_item_undo',
        '长SQL':         'oracle_check_item_long_sql',
        '性能指标':      'oracle_check_item_perf',
        'Top SQL':       'oracle_check_item_top_sql',
        '无效对象':     'oracle_check_item_invalid_obj',
        '用户安全':     'oracle_check_item_users',
        '备份信息':     'oracle_check_item_backup',
        '闪回/回收站': 'oracle_check_item_flashback',
        'Data Guard':   'oracle_check_item_dataguard',
        'RAC+ASM':      'oracle_check_item_rac_asm',
        'AWR快照':      'oracle_check_item_awr',
        '作业调度':     'oracle_check_item_jobs',
        'Alert日志':    'oracle_check_item_alert',
        '阻塞会话':     'oracle_check_item_blocking',
        '死锁检测':     'oracle_check_item_deadlock',
        '长事务':        'oracle_check_item_long_trx',
    }

    def _item_name(name):
        key = _item_i18n.get(name)
        return _t(key) if key else name

    print(f"\n{GREEN}▶ {_t('oracle_log_start')}{RESET}")
    t0 = time.time()

    # ── 1. Oracle 连接 ─────────────────────────────────────────────────────
    print(f"\n[{GREEN}1/6{RESET}] {_t('oracle_log_connect_db')}")
    try:
        if args.servicename:
            dsn = oracledb.makedsn(args.host, args.port, service_name=args.servicename)
        else:
            dsn = oracledb.makedsn(args.host, args.port, args.sid)
        # sys 用户默认以 SYSDBA 身份连接（oracle privilege model）
        if args.user.upper() == 'SYS' and not args.sysdba:
            args.sysdba = True
        mode = oracledb.SYSDBA if args.sysdba else None
        ssh_tunnel = None
        conn, ssh_tunnel = _get_oracle_conn_thunk_first(
            dsn, args.user, args.password, mode,
            ssh_host=args.ssh_host, ssh_port=args.ssh_port or 22,
            ssh_user=args.ssh_user, ssh_password=args.ssh_pass,
            ssh_key=args.ssh_key if hasattr(args, 'ssh_key') else None
        )
        print(f"  ✅ {_t('oracle_log_connect_ok')} (mode: {'SYSDBA' if args.sysdba else 'NORMAL'})")
    except Exception as e:
        print(f"  ❌ {_t('oracle_log_connect_fail')}: {e}")
        return

    # ── 2. 获取版本 ───────────────────────────────────────────────────────
    print(f"\n[{GREEN}2/6{RESET}] {_t('oracle_log_get_version')}")
    version_str, ver_major = get_db_version_and_major(conn)
    print(f"  {_t('oracle_log_version_fmt').format(ver=version_str, major=ver_major)}")

    # ── 3. OS 层采集 ──────────────────────────────────────────────────────
    print(f"\n[{GREEN}3/6{RESET}] {_t('oracle_log_collect_os')}")
    os_data = {}
    ssh_client = None
    if args.ssh_host:
        try:
            ssh_client = paramiko.SSHClient()
            ssh_client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            ssh_client.connect(
                args.ssh_host, port=args.ssh_port or 22,
                username=args.ssh_user, password=args.ssh_pass,
                timeout=15
            )
            collector = OSCollector(ssh_client)
            os_data = collector.collect()
            print(f"  ✅ {_t('oracle_log_ssh_ok')} (host: {args.ssh_host})")
        except Exception as e:
            print(f"  ⚠ {_t('oracle_log_ssh_fail')}: {e}，{_t('oracle_log_use_local')}")
            collector = OSCollector(None)
            os_data = collector.collect()
    else:
        collector = OSCollector(None)
        os_data = collector.collect()
        print(f"  ✅ {_t('oracle_log_local_ok')}")

    # ── 4. 数据库层巡检（版本自适应）──────────────────────────────────────
    _ver_suffix = 'g' if ver_major in ('10', '11') else 'c'
    print(f"\n[{GREEN}4/6{RESET}] {_t('oracle_log_db_inspect')} (Oracle {ver_major}{_ver_suffix})...")
    check_results = {}
    chapter_results = []

    template_id = getattr(args, 'template_id', None)

    # ── 若无指定 template_id，自动查询 DB 默认模板（is_default=1）─
    if template_id is None:
        try:
            from inspection_dal import get_default_template as _get_default_template
            _default = _get_default_template('oracle')
            if _default:
                template_id = _default.get('id')
        except Exception:
            pass  # 无默认模板则跳过配置章节

    # ── 4.1 硬编码检查函数 ──────────────────────────────────────────────
    # 当存在模板时，跳过模板章节已覆盖的项，避免重复查询同一数据
    _TEMPLATE_COVERED = {
        '实例信息', '数据库信息', '版本/补丁', '表空间', 'Redo日志',
        '控制文件', 'SGA/PGA内存', '关键参数', 'Undo信息', '性能指标',
        'Top SQL', '无效对象', '用户安全', '备份信息', 'Data Guard',
        'AWR快照', '作业调度',
    }
    checks = get_checks_for_version(ver_major)
    if template_id is not None:
        checks = [(n, f) for n, f in checks if n not in _TEMPLATE_COVERED]
    for name, fn in checks:
        try:
            result = fn(conn)
            if result and 'error' not in result:
                check_results[name] = result
                rows = list(result.values())[0] if result else []
                cnt = len(rows) if isinstance(rows, list) else '-'
                print(f"  ✅ {_item_name(name)}  ({cnt} {_plural(cnt, 'row', 'rows')})")
            elif result and 'error' in result:
                print(f"  ⚠ {_item_name(name)}  {_t('oracle_log_check_fail').format(error=result.get('error', 'unknown'))}")
            else:
                print(f"  ⚠ {_item_name(name)}  {_t('oracle_log_check_empty')}")
        except Exception as e:
            print(f"  ⚠ {_item_name(name)}  {_t('oracle_log_check_skip')}: {e}")

    # ── 4.2 配置驱动章节（仅在指定 template_id 时执行）───────────────────
    if template_id is not None:
        chapters = load_oracle_chapter_structure(template_id)
        if chapters:
            enabled_count = sum(1 for ch in chapters if ch.get('queries'))
            print(f"\n  ℹ 使用巡检配置模板，加载 {enabled_count} 个配置章节...")
            chapter_results = execute_oracle_chapter_queries(conn, chapters)
            for ch in chapter_results:
                title = ch['chapter_title_zh']
                qcount = len(ch['queries_results'])
                print(f"  📋 {title} ({qcount} 个查询)")

    # ── 4.5 AI 诊断（根据配置判断是否启用）───────────────────────────────────
    print(f"\n[{GREEN}4.5/6{RESET}] {_t('oracle_log_ai_diagnosis')}")
    ai_advice = ''
    risk_items = []  # 在 try 块外初始化，确保 AI 禁用时也能被 history 引用
    try:
        from analyzer import AIAdvisor
        cfg_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'dbc_config.json')
        ai_cfg = {}
        if os.path.exists(cfg_path):
            with open(cfg_path, 'r', encoding='utf-8') as f:
                ai_cfg = json.load(f).get('ai', {})
        advisor = AIAdvisor(
            backend=ai_cfg.get('backend'),
            api_key=ai_cfg.get('api_key'),
            api_url=ai_cfg.get('api_url'),
            model=ai_cfg.get('model')
        )
        if advisor.enabled:
            # 从 check_results 直接构建 db_info（避免引用报告生成阶段才定义的变量）
            _db_info = {}
            for row in check_results.get('数据库信息', {}).get('database', []):
                cols = ['DBID','NAME','DATABASE_ROLE','OPEN_MODE','LOG_MODE',
                        'CREATED','STARTUP_TIME','CDB','flashback_on','force_logging',
                        'block_size','sga_max_size','sga_target','pga_aggregate_target',
                        'spfile','global_name']
                for i, c in enumerate(cols):
                    if i < len(row):
                        _db_info[c] = row[i]
            for row in check_results.get('实例信息', {}).get('instance', []):
                if len(row) >= 4:
                    _db_info.setdefault('INSTANCE_NAME', row[1])
                    _db_info.setdefault('HOST_NAME', row[2])
                    _db_info.setdefault('VERSION', row[3])
            label = _db_info.get('NAME', args.servicename or args.sid or 'ORACLE')
            print(f"  🤖 {_t('oracle_log_ai_calling')} ({advisor.backend} / {advisor.model})...")
            # 收集风险项作为上下文
            risk_items = []
            ts = check_results.get('表空间', {})
            for row in ts.get('data_tablespaces', []):
                if len(row) >= 8 and row[7] != '-':
                    try:
                        used_pct = float(row[7])
                        if used_pct > 90:
                            risk_items.append({'col1': _t('report.risk_tablespace').format(name=row[0]), 'col2': _t('report.risk_high'),
                                'col3': _t('report.risk_ts_high').format(pct=f'{used_pct:.1f}')})
                        elif used_pct > 80:
                            risk_items.append({'col1': _t('report.risk_tablespace').format(name=row[0]), 'col2': _t('report.risk_mid'),
                                'col3': _t('report.risk_ts_mid').format(pct=f'{used_pct:.1f}')})
                    except (ValueError, TypeError):
                        pass
            # 收集等待事件 Top5
            perf = check_results.get('性能指标', {})
            wait_top5 = perf.get('wait_events', [])[:5]
            wait_summary = '\n'.join([_t('oracle_log_ai_wait_fmt').format(w0=w[0], w1=w[1], w2=w[2], w3=w[3])
                                      for w in wait_top5]) if wait_top5 else 'N/A'
            # 收集阻塞会话
            blocked_data = check_results.get('阻塞会话', {})
            blocking_rows = blocked_data.get('blocking_chain', [])
            blocked_sessions = len(blocking_rows)
            blocked_summary = _t('oracle_log_ai_blocked').format(n=blocked_sessions) if blocked_sessions else _t('oracle_log_ai_no_blocked')
            # 收集 Top SQL（按 Buffer Gets 前5）
            top_sql_raw = check_results.get('Top SQL', {})
            top_sql5 = top_sql_raw.get('top_sql_buffer_gets', [])[:5]
            top_sql_summary = '\n'.join([
                _t('oracle_log_ai_top_sql_fmt').format(s0=r[0], s1=str(r[1])[:60], s2=r[2], s4=r[4], s5=r[5])
                for r in top_sql5]) if top_sql5 else 'N/A'
            # 构建详细指标
            metrics = {
                'db_version': _db_info.get('VERSION', version_str),
                'hostname': _db_info.get('HOST_NAME', os_data.get('hostname', 'N/A')),
                'uptime': os_data.get('uptime', 'N/A'),
                'risk_count': len(risk_items),
                'tablespace_count': len(ts.get('data_tablespaces', [])),
                'wait_events_top5': wait_summary,
                'blocked_sessions': blocked_summary,
                'top_sql_top5': top_sql_summary,
            }
            ai_advice = advisor.diagnose('oracle', label, metrics, risk_items, timeout=600, lang=_lang)
            if ai_advice:
                print(f"  ✅ {_t('oracle_log_ai_ok')}")
            else:
                print(f"  ⚠ {_t('oracle_log_ai_empty')}")
        else:
            print(f"  ⏭ {_t('oracle_log_ai_disabled')} (backend: {advisor.backend})")
    except TimeoutError as e:
        ai_advice = f"⚠ {_t('oracle_log_ai_timeout')}"
        print(f"  {ai_advice}")
    except Exception as e:
        err_str = str(e)
        if 'connection' in err_str.lower() or 'refused' in err_str.lower():
            ai_advice = f"⚠ {_t('oracle_log_ai_conn_fail')}"
        else:
            ai_advice = f"⚠ {_t('oracle_log_ai_fail')}: {err_str[:120]}"
        print(f"  {ai_advice}")

    # ── 4.6 慢查询深度分析（P2）────────────────────────────────────────────
    slow_query_result = None
    try:
        from slow_query_analyzer import OracleSlowQueryAnalyzer
        analyzer = OracleSlowQueryAnalyzer()
        ai_advisor = None
        try:
            from analyzer import AIAdvisor
            ai_advisor = AIAdvisor(
                backend=ai_cfg.get('backend'),
                api_key=ai_cfg.get('api_key'),
                api_url=ai_cfg.get('api_url'),
                model=ai_cfg.get('model')
            )
        except Exception:
            pass
        print(f"\n[{GREEN}4.6/6{RESET}] {_t('oracle_log_slow_query')}")
        result = analyzer.analyze(conn, ai_advisor=ai_advisor, lang=_lang)
        slow_query_result = result.to_dict()
        if result.is_empty():
            print(f"  \u2139\ufe0f  {_t('oracle_log_slow_query_empty')}")
        else:
            print(f"  \u2705  {_t('oracle_log_slow_query_ok').format(count=len(result.top_sql_by_latency))}")
    except ImportError:
        print(f"  \u26a0  slow_query_analyzer 模块未找到，跳过慢查询深度分析")
    except Exception as e:
        print(f"  \u26a0  慢查询深度分析失败: {e}")

    # ── 4.7 配置基线检查（P3）────────────────────────────────────────────
    config_baseline_result = None
    try:
        from config_baseline import check_oracle_config_baseline
        print(f"\n[{GREEN}4.7/6{RESET}] {_t('oracle_cli_config_baseline_checking')}")
        config_baseline_result = check_oracle_config_baseline(conn)
        cb = config_baseline_result
        summary = cb.get('summary', {})
        crit = summary.get('critical_count', 0)
        warn = summary.get('warning_count', 0)
        info = summary.get('info_count', 0)
        print(f"  ✅  {_t('oracle_cli_config_baseline_ok') % (crit, warn, info)}")
    except ImportError:
        print(f"  ⚠  config_baseline 模块未找到，跳过配置基线检查")
    except Exception as e:
        print(f"  ⚠  配置基线检查失败: {e}")

    # ── 4.8 索引健康分析（P3）────────────────────────────────────────────
    index_health_result = None
    try:
        from index_health import analyze_oracle_indexes
        print(f"\n[{GREEN}4.8/6{RESET}] {_t('oracle_cli_index_health_checking')}")
        index_health_result = analyze_oracle_indexes(conn)
        ih = index_health_result
        sm = ih.get('summary', {})
        miss = sm.get('missing_count', 0)
        redun = sm.get('redundant_count', 0)
        unused = sm.get('unused_count', 0)
        print(f"  ✅  {_t('oracle_cli_index_health_ok') % (miss, redun, unused)}")
    except ImportError:
        print(f"  ⚠  index_health 模块未找到，跳过索引健康分析")
    except Exception as e:
        print(f"  ⚠  索引健康分析失败: {e}")

    conn.close()
    if ssh_tunnel:
        ssh_tunnel.close()  # 关闭数据库连接

    # ── 从 check_results 提取 db_info（报告生成 + 历史记录共用）───
    _db_info = {}
    _inst_rows = check_results.get('实例信息', {}).get('instance', [])
    _db_rows   = check_results.get('数据库信息', {}).get('database', [])
    if _db_rows:
        _cols = ['DBID','NAME','DATABASE_ROLE','OPEN_MODE','LOG_MODE',
                 'CREATED','STARTUP_TIME','CDB','flashback_on','force_logging',
                 'block_size','sga_max_size','sga_target','pga_aggregate_target',
                 'spfile','global_name']
        for row in _db_rows:
            for i, c in enumerate(_cols):
                if i < len(row):
                    _db_info[c] = row[i]
    if _inst_rows and len(_inst_rows[0]) >= 4:
        _db_info['INSTANCE_NAME'] = _inst_rows[0][1]
        _db_info['HOST_NAME']     = _inst_rows[0][2]
        _db_info['VERSION']       = _inst_rows[0][3]
        _db_info['STARTUP_TIME']  = _inst_rows[0][4]
        _db_info['STATUS']         = _inst_rows[0][5]
    _db_name = _db_info.get('NAME', args.servicename or args.sid or 'ORACLE')

    # ── 构建 context（供 smart_analyze_oracle 和 HistoryManager 使用）───
    def _ts_rows(data_ts):
        """将 check_results 表空间数据转为 context 格式"""
        rows = []
        for row in data_ts:
            if len(row) >= 8:
                rows.append({
                    'TABLESPACE_NAME': str(row[0]),
                    'STATUS': str(row[1]),
                    'CONTENTS': str(row[2]) if len(row) > 2 else '',
                    'TOTAL_MB': float(row[4]) if row[4] != '-' else 0,
                    'USED_MB': float(row[5]) if row[5] != '-' else 0,
                    'USED_PCT_WITH_MAXEXT': float(row[7]) if row[7] != '-' else 0,
                })
        return rows

    _perf = check_results.get('性能指标', {})
    _sess_rows = _perf.get('session_by_status', [])
    _total_sess = sum(int(r[1]) for r in _sess_rows if len(r) >= 2 and str(r[1]).isdigit())
    _ora_sessions_fmt = [{'TOTAL_SESSIONS': _total_sess}]

    _sga_rows = check_results.get('SGA/PGA内存', {}).get('sga_total', [])
    _sga_val = _sga_rows[0][0] if _sga_rows and _sga_rows[0] else 0.0
    _ora_sga_fmt = [{'SGA_TOTAL_MB': float(_sga_val)}]

    _params = check_results.get('关键参数', {})
    _sess_limit = 0
    for row in _params.get('params', []):
        if len(row) >= 2 and str(row[0]).lower() == 'sessions':
            try:
                _sess_limit = int(float(str(row[1])))
            except (ValueError, TypeError):
                pass
            break
    _ora_session_limit_fmt = [{'SESSIONS_LIMIT': _sess_limit}] if _sess_limit else []

    _blocking_data = check_results.get('阻塞会话', {})
    _blocking_rows = _blocking_data.get('blocking_chain', [])
    _ora_blocked_fmt = []
    for r in _blocking_rows:
        _ora_blocked_fmt.append({
            'BLOCKED_SID': str(r[0]) if len(r) > 0 else '',
            'BLOCKED_SERIAL': str(r[1]) if len(r) > 1 else '',
            'BLOCKED_USER': str(r[2]) if len(r) > 2 else '',
            'BLOCKED_EVENT': str(r[6]) if len(r) > 6 else '',
            'SEC_IN_WAIT': float(r[7]) if len(r) > 7 and r[7] is not None else 0,
            'BLOCKING_SID': str(r[8]) if len(r) > 8 else '',
            'BLOCKING_SERIAL': str(r[9]) if len(r) > 9 else '',
            'BLOCKING_USER': str(r[10]) if len(r) > 10 else '',
            'LOCK_TYPE': str(r[14]) if len(r) > 14 else '',
            'LOCKED_OBJECT': str(r[17]) if len(r) > 17 else '',
        })

    _deadlock_data = check_results.get('死锁检测', {})
    _deadlock_stats = _deadlock_data.get('deadlock_stats', [])
    _ora_deadlock_fmt = []
    for r in _deadlock_stats:
        if len(r) >= 2:
            _ora_deadlock_fmt.append({
                'STAT_NAME': str(r[0]),
                'STAT_VALUE': int(str(r[1])) if r[1] is not None else 0,
            })

    _long_trx_data = check_results.get('长事务', {})
    _long_trx_rows = _long_trx_data.get('long_trx', [])
    _ora_long_trx_fmt = []
    for r in _long_trx_rows:
        _ora_long_trx_fmt.append({
            'SID': str(r[0]) if len(r) > 0 else '',
            'SERIAL': str(r[1]) if len(r) > 1 else '',
            'USERNAME': str(r[2]) if len(r) > 2 else '',
            'TRX_START': str(r[8]) if len(r) > 8 else '',
            'TRX_SECONDS': float(r[9]) if len(r) > 9 and r[9] is not None else 0,
            'UNDO_BLOCKS': int(str(r[10])) if len(r) > 10 and r[10] is not None else 0,
        })

    context = {
        'ora_version': [{'BANNER': version_str}],
        'ora_tablespace': _ts_rows(check_results.get('表空间', {}).get('data_tablespaces', [])),
        'ora_sessions': _ora_sessions_fmt,
        'ora_sga_total': _ora_sga_fmt,
        'ora_session_limit': _ora_session_limit_fmt,
        'ora_blocked': _ora_blocked_fmt,
        'ora_deadlock': _ora_deadlock_fmt,
        'ora_long_trx': _ora_long_trx_fmt,
        'system_info': {
            'hostname': os_data.get('hostname', ''),
            'cpu': {'usage_percent': os_data.get('cpu_usage_pct', 0)},
            'memory': {'usage_percent': os_data.get('mem_usage_pct', os_data.get('mem_percent', 0))},
            'disk_list': [{'mountpoint': d.get('mount', '/'), 'usage_percent': d.get('percent', 0)}
                          for d in os_data.get('disk_list', [])],
            'disk_usage': os_data.get('disk_usage', ''),
        },
        'health_status': _t('report.health_good') if not risk_items else (_t('report.health_attention') if any(r.get('col2') == _t('report.risk_high') for r in risk_items) else _t('report.health_fair')),
        'auto_analyze': risk_items if risk_items else [],
        # Web UI 报告生成所需数据（非 CLI 模式使用）
        '_oracle_db_info': _db_info,
        '_oracle_os_data': os_data,
        '_oracle_check_results': check_results,
        '_oracle_version_str': version_str,
        '_oracle_ai_advice': ai_advice,
        '_oracle_lang': _lang,
        '_oracle_config_baseline': config_baseline_result,
        '_oracle_index_health': index_health_result,
        '_oracle_chapter_results': chapter_results,
    }

    # ── Web UI 模式：跳过报告生成和历史保存，由统一调度层处理 ──
    if _web_ui_mode:
        elapsed = time.time() - t0
        if ssh_client:
            ssh_client.close()
        return context

    # ── CLI 模式：生成报告 + 保存历史（保持独立运行能力）───
    # ── 5. 生成报告 ────────────────────────────────────────────────────────
    print(f"\n[{GREEN}5/6{RESET}] {_t('oracle_log_gen_report')}")
    docx = build_word_report(_db_info, os_data, check_results, version_str, ai_advice,
                              inspector=args.inspector or 'dbcheck', lang=_lang,
                              desensitize=bool(getattr(args, 'desensitize', False)),
                              config_baseline_result=config_baseline_result,
                              index_health_result=index_health_result,
                              host=args.host, health_status='',
                              chapter_results=chapter_results)

    # ── 6. 保存报告 ────────────────────────────────────────────────────────
    print(f"\n[{GREEN}6/6{RESET}] {_t('oracle_log_save_report')}")
    output_dir = args.output or os.path.join(os.getcwd(), 'reports')
    os.makedirs(output_dir, exist_ok=True)

    ver_tag  = ver_major or 'DB'
    ts = time.strftime('%Y%m%d%H%M%S')

    # Word
    if docx:
        fname_template = _t('webui.oracle_report_filename')
        docx_fname = fname_template.format(ip=args.host, name=_db_name, ts=ts) + '.docx'
        docx_path  = os.path.join(output_dir, docx_fname)
        try:
            docx.save(docx_path)
            print(f"   Word:  {docx_path}")
        except Exception as e:
            print(f"   {_t('oracle_log_word_report')}: {e}")

    # ── 保存历史记录 ──────────────────────────────────────────────────────
    try:
        from analyzer import HistoryManager
        script_dir = os.path.dirname(os.path.abspath(__file__))
        label = _db_info.get('NAME', args.servicename or args.sid or 'ORACLE')
        hm = HistoryManager(script_dir)
        hm.save_snapshot('oracle_full', args.host, args.port, label, context)
        print(f"  ✅ {_t('oracle_log_history_ok')}")
    except Exception as e:
        print(f"  ⚠ {_t('oracle_log_history_fail')}: {e}")

    elapsed = time.time() - t0

    if ssh_client:
        ssh_client.close()

    return context

def _input(prompt, default=''):
    """统一输入函数，带默认值显示"""
    if default:
        val = input(f"{prompt} [{default}]: ").strip()
        return val if val else default
    return input(f"{prompt}: ").strip()

def _password_input(prompt):
    """密码输入函数，隐藏用户输入"""
    return getpass.getpass(prompt)

def interactive_single_inspection():
    """交互式单机巡检（替代 argparse，适合无参数直接运行）"""
    from i18n import t
    print(f"\n{BOLD}{'='*52}{RESET}")
    print(f"{RED}{BOLD}   {t('oracle_banner_title')}{RESET}")
    print(f"{DIM}{'='*52}{RESET}\n")

    # ── Oracle 连接信息 ─────────────────────────────────────────
    host        = _input(f"{CYAN}{t('oracle_host_ip')}{RESET}",    'localhost')
    port        = _input(f"{CYAN}{t('oracle_port')}{RESET}",             '1521')
    connect_by  = _input(f"{CYAN}{t('oracle_connect_by')}{RESET}", 'S').upper()
    if connect_by == 'N':
        sid_or_svc = _input(f"{CYAN}{t('oracle_servicename')}{RESET}")
        sid, svc = None, sid_or_svc
    else:
        sid     = _input(f"{CYAN}{t('oracle_sid')}{RESET}",       'ORCL')
        svc     = None
    user        = _input(f"{CYAN}{t('oracle_username')}{RESET}",           'sys')
    password    = _password_input(f"{t('oracle_password')}: ")
    # sys 用户默认以 SYSDBA 登录，其他用户可自行选择
    if user.upper() == 'SYS':
        sysdba_default = 'Y'
    else:
        sysdba_default = 'N'
    sysdba_opt  = _input(f"{CYAN}{t('oracle_sysdba_prompt')}{RESET}", sysdba_default).upper()
    sysdba      = (sysdba_opt == 'Y')

    # ── SSH 信息（可选）────────────────────────────────────────
    use_ssh = _input(f"\n{GREEN}{t('oracle_ssh_use')}{RESET}", 'n').upper()
    ssh_host, ssh_port, ssh_user, ssh_pass = None, 22, None, None
    if use_ssh == 'Y':
        ssh_host = _input(f"{CYAN}{t('oracle_ssh_host')}{RESET}", host)
        ssh_port = _input(f"{CYAN}{t('oracle_ssh_port')}{RESET}",   '22')
        ssh_user = _input(f"{CYAN}{t('oracle_ssh_username')}{RESET}")
        ssh_pass = _password_input(f"{t('oracle_ssh_password')}: ")
        if not ssh_user or not ssh_pass:
            print(f"  {YELLOW}⚠ {t('oracle_ssh_skip_warning')}{RESET}")
            ssh_host, ssh_user, ssh_pass = None, None, None

    # ── 输出选项 ───────────────────────────────────────────────
    output_dir = _input(f"\n{GREEN}{t('oracle_output_dir')}{RESET}", 'reports')
    inspector  = _input(f"{GREEN}{t('oracle_inspector_name')}{RESET}", 'dbcheck')

    # ── 构造 args ───────────────────────────────────────────────
    class _Args:
        pass
    args = _Args()
    args.host        = host
    args.port        = int(port)
    args.sid         = sid
    args.servicename = svc
    args.user        = user
    args.password    = password
    args.sysdba      = sysdba
    args.ssh_host    = ssh_host
    args.ssh_port    = int(ssh_port) if ssh_port else 22
    args.ssh_user    = ssh_user
    args.ssh_pass    = ssh_pass
    args.output      = output_dir if output_dir else None
    args.inspector   = inspector or 'dbcheck'

    print_banner()
    single_inspection(args)

def main():
    import argparse
    parser = argparse.ArgumentParser(
        description=f'DBCheck Oracle 全面巡检工具 v{VER}（OS层+数据库层）',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 本地 Oracle（使用 SID）
  python main_oracle_full.py -h localhost -P 1521 -s ORCL -u system -p xxx

  # 使用 ServiceName 连接
  python main_oracle_full.py -h localhost -P 1521 -S ORCL -u system -p xxx

  # SSH 采集 OS 信息
  python main_oracle_full.py -h localhost -P 1521 -s ORCL -u system -p xxx \\
      --ssh-host localhost --ssh-user oracle --ssh-pass xxx

        """
    )
    parser.add_argument('--host',             required=False, help='Oracle 主机IP（交互模式无需指定）')
    parser.add_argument('-P', '--port',      type=int, default=1521, help='端口（默认1521）')
    parser.add_argument('-s', '--sid',        help='ORACLE_SID（SID和ServiceName二选一）')
    parser.add_argument('-S', '--servicename', help='ServiceName（SID和ServiceName二选一）')
    parser.add_argument('-u', '--user',       default='sys', help='用户名（默认sys）')
    parser.add_argument('-p', '--password',  default='',   help='密码')
    parser.add_argument('--sysdba',       action='store_true', help='以 SYSDBA 身份连接（sys用户默认开启）')
    parser.add_argument('--ssh-host',   help='SSH 主机（与 Oracle 主机相同则省略）')
    parser.add_argument('--ssh-port',   type=int, default=22, help='SSH 端口（默认22）')
    parser.add_argument('--ssh-user',   help='SSH 用户名')
    parser.add_argument('--ssh-pass',   help='SSH 密码')
    parser.add_argument('-o', '--output',   help='报告输出目录')
    parser.add_argument('--inspector',   help='巡检人姓名')

    args = parser.parse_args()

    # Local _t for main() error messages
    try:
        from i18n import get_lang
        _lang = get_lang()
    except Exception:
        _lang = 'zh'

    def _t(key):
        try:
            from i18n import t as _tt
            return _tt(key, _lang)
        except Exception:
            return key

    # 无参数时进入交互模式
    if len(sys.argv) == 1 or (
           not args.host and not args.sid and not args.servicename
    ):
        interactive_single_inspection()
        return

    if not args.sid and not args.servicename:
        print(f"❌ {_t('oracle_log_need_sid_svc')}")
        return

    if args.ssh_host and not (args.ssh_user and args.ssh_pass):
        print(f"❌ {_t('oracle_log_need_ssh_cred')}")
        return

    print_banner()
    single_inspection(args)

# ═══════════════════════════════════════════════════════════════════════════
# v19 兼容版覆盖函数（12c 基准不动，19c 出错项单独覆盖）
# ═══════════════════════════════════════════════════════════════════════════

def oracle_check_database_v19(conn):
    """数据库层信息（v19 专用版）：不使用 PLUGGABLE_DB（该列在目标环境不存在）"""
    results = {}
    cur = conn.cursor()
    try:
        # 只取 v$database 中稳定存在的列，不依赖 PLUGGABLE_DB
        cur.execute("""
            SELECT DBID, NAME, DATABASE_ROLE, CREATED, LOG_MODE, OPEN_MODE,
                   FLASHBACK_ON, FORCE_LOGGING, CREATED
            FROM v$database
        """)
        results['database'] = cur.fetchall()

        cur.execute("SELECT global_name FROM global_name")
        results['global_name'] = cur.fetchone()

        try:
            cur.execute("""
                SELECT parameter, value
                FROM nls_database_parameters
                WHERE parameter IN ('NLS_CHARACTERSET', 'NLS_NCHAR_CHARACTERSET')
            """)
            rows = cur.fetchall()
            results['charset'] = tuple(r[1] for r in rows) if rows else ('', '')
        except Exception:
            results['charset'] = ('', '')

        for param in ['db_block_size', 'sga_max_size', 'sga_target',
                      'pga_aggregate_target', 'memory_max_target', 'memory_target']:
            try:
                cur.execute(f"SELECT value FROM v$parameter WHERE name='{param}'")
                r = cur.fetchone()
                results[param] = r[0] if r else ''
            except Exception:
                results[param] = ''

        # ADR 相关路径
        try:
            cur.execute("SELECT value FROM v$parameter WHERE name='diagnostic_dest'")
            r = cur.fetchone()
            results['adr'] = r[0] if r else ''
        except Exception:
            results['adr'] = ''

        try:
            cur.execute("SELECT value FROM v$parameter WHERE name='db_create_file_dest'")
            r = cur.fetchone()
            results['omf'] = r[0] if r else ''
        except Exception:
            results['omf'] = ''

        try:
            cur.execute("SELECT log_mode FROM v$database")
            r = cur.fetchone()
            results['log_mode'] = r[0] if r else ''
        except Exception:
            results['log_mode'] = ''

        try:
            cur.execute("SELECT force_logging FROM v$database")
            r = cur.fetchone()
            results['force_logging'] = r[0] if r else ''
        except Exception:
            results['force_logging'] = ''

        try:
            cur.execute("SELECT flashback_on FROM v$database")
            r = cur.fetchone()
            results['flashback_on'] = r[0] if r else ''
        except Exception:
            results['flashback_on'] = ''

        try:
            cur.execute("SELECT TO_CHAR(CREATED, 'YYYY-MM-DD HH24:MI:SS') FROM v$database")
            r = cur.fetchone()
            results['created'] = r[0] if r else ''
        except Exception:
            results['created'] = ''

        try:
            cur.execute("SELECT TO_CHAR(STARTUP_TIME, 'YYYY-MM-DD HH24:MI:SS') FROM v$instance")
            r = cur.fetchone()
            results['startup_time'] = r[0] if r else ''
        except Exception:
            results['startup_time'] = ''

    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def _col_name(cur, view, pattern):
    """探测视图实际列名（模糊匹配），找不到返回 None"""
    try:
        cur.execute(f"""
            SELECT column_name FROM user_tab_columns
            WHERE table_name = UPPER('{view}')
              AND column_name LIKE '%{pattern}%'
            FETCH FIRST 1 ROWS ONLY
        """)
        r = cur.fetchone()
        return r[0] if r else None
    except Exception:
        return None

def oracle_check_tablespace_v19(conn):
    """表空间（v19 兼容版）：自动探测 dba_temp_free_space/dba_free_space 的实际列名，不再猜"""
    results = {}
    cur = conn.cursor()
    try:
        # ── 自动探测空闲列名 ─────────────────────────────────────────
        # dba_temp_free_space：可能有 FREE_SPACE / TABLESPACE_SIZE / ALLOCATED_SPACE 等
        tfs_col = (_col_name(cur, 'dba_temp_free_space', 'FREE')
                or _col_name(cur, 'dba_temp_free_space', 'SPACE')
                or _col_name(cur, 'dba_temp_free_space', 'SIZE'))

        # dba_free_space：可能有 FREE_SPACE / BYTES / BLOCKS
        fs_col = (_col_name(cur, 'dba_free_space', 'FREE')
              or _col_name(cur, 'dba_free_space', 'BYTES')
              or _col_name(cur, 'dba_free_space', 'BLOCKS'))

        # 永久表空间
        cur.execute("""
            SELECT bt.tablespace_name,
                   bt.status,
                   ROUND(NVL(df.curr_mb,0), 2) curr_mb,
                   ROUND(NVL(df.max_mb,0), 2) max_mb,
                   ROUND(NVL(seg.used_mb,0), 2) used_mb,
                   ROUND(NVL(df.curr_mb,0) - NVL(seg.used_mb,0), 2) free_mb,
                   ROUND(NVL(seg.used_mb,0) / NULLIF(GREATEST(NVL(df.curr_mb,0), NVL(df.max_mb,0)), 0) * 100, 2) pct_used
            FROM dba_tablespaces bt
            LEFT JOIN (SELECT tablespace_name,
                              SUM(bytes/1024/1024) curr_mb,
                              SUM(MAXBYTES/1024/1024) max_mb
                       FROM dba_data_files GROUP BY tablespace_name) df
                   ON bt.tablespace_name = df.tablespace_name
            LEFT JOIN (SELECT tablespace_name,
                              SUM(bytes/1024/1024) used_mb
                       FROM dba_segments GROUP BY tablespace_name) seg
                   ON bt.tablespace_name = seg.tablespace_name
            WHERE bt.contents = 'PERMANENT'
            ORDER BY pct_used DESC NULLS LAST
        """)
        results['data_tablespaces'] = cur.fetchall()

        # 临时表空间：只查 dba_temp_files，不依赖 dba_temp_free_space（列名不稳定）
        cur.execute("""
            SELECT bt.tablespace_name,
                   bt.status,
                   ROUND(NVL(tf.curr_mb,0), 2) curr_mb,
                   ROUND(NVL(tf.max_mb,0), 2) max_mb,
                   '-' used_mb,
                   '-' free_mb,
                   '-' pct_used
            FROM dba_tablespaces bt
            LEFT JOIN (SELECT tablespace_name,
                              SUM(bytes/1024/1024) curr_mb,
                              SUM(MAXBYTES/1024/1024) max_mb
                       FROM dba_temp_files GROUP BY tablespace_name) tf
                   ON bt.tablespace_name = tf.tablespace_name
            WHERE bt.contents = 'TEMPORARY'
            ORDER BY bt.tablespace_name
        """)
        results['temp_tablespaces'] = cur.fetchall()

        # 自动扩展文件（dba_data_files 有 BYTES/MAXBYTES 是确定的）
        cur.execute("""
            SELECT tablespace_name, file_name,
                   ROUND(bytes/1024/1024,2) curr_mb,
                   ROUND(MAXBYTES/1024/1024,2) max_mb,
                   AUTOEXTENSIBLE
            FROM dba_data_files
            WHERE AUTOEXTENSIBLE = 'YES'
            ORDER BY tablespace_name
        """)
        results['autoextend_files'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_redolog_v19(conn):
    """Redo 日志（v19 兼容版）：不使用 v$loghist（列不稳定），直接查 v$log"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT GROUP#, THREAD#, SEQUENCE#, ROUND(BYTES/1024/1024,2) size_mb,
                   STATUS, MEMBERS, ARCHIVED
            FROM v$log
            ORDER BY THREAD#, GROUP#
        """)
        results['logs'] = cur.fetchall()

        cur.execute("""
            SELECT GROUP#, MEMBER, TYPE, STATUS
            FROM v$logfile
            ORDER BY GROUP#
        """)
        results['logfiles'] = cur.fetchall()

        # 直接查 v$log，统计 CURRENT 组大小（不依赖 v$loghist）
        cur.execute("""
            SELECT THREAD#,
                   COUNT(*) switch_cnt,
                   ROUND(SUM(BYTES)/1024/1024/1024, 2) total_mb
            FROM v$log
            WHERE STATUS = 'CURRENT'
            GROUP BY THREAD#
        """)
        results['redo_switch'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_top_sql_v19(conn, limit=20):
    """Top SQL（v19 兼容版）：去除中文字段别名，避免字符集解析问题"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute(f"""
            SELECT * FROM (
                SELECT sql_id,
                       SUBSTR(sql_text,1,80) AS sql_text,
                       ROUND(buffer_gets/1024/1024,2) AS buf_mb,
                       ROUND(disk_reads/1024/1024,2) AS disk_mb,
                       executions,
                       ROUND(elapsed_time/1000000,2) AS elapsed_sec,
                       ROUND(buffer_gets/DECODE(executions,0,1,executions)) AS gets_per_exec,
                       module
                FROM v$sql
                WHERE executions > 0
                ORDER BY buffer_gets DESC
            ) WHERE ROWNUM <= {limit}
        """)
        results['top_sql_buffer_gets'] = cur.fetchall()

        cur.execute(f"""
            SELECT * FROM (
                SELECT sql_id,
                       SUBSTR(sql_text,1,80) AS sql_text,
                       ROUND(disk_reads/1024/1024,2) AS disk_mb,
                       executions,
                       ROUND(elapsed_time/1000000,2) AS elapsed_sec,
                       module
                FROM v$sql
                WHERE executions > 0
                ORDER BY disk_reads DESC
            ) WHERE ROWNUM <= {limit}
        """)
        results['top_sql_disk_reads'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_backup_v19(conn):
    """RMAN 备份信息（v19 兼容版）：v$rman_backup_job_details 无 bytes 列，用 TIME_TAKEN_DISPLAY"""
    results = {}
    cur = conn.cursor()
    try:
        # v$rman_backup_job_details 标准列（19c 通用）：SESSION_KEY/INPUT_TYPE/STATUS/START_TIME/END_TIME/TIME_TAKEN_DISPLAY
        cur.execute("""
            SELECT SESSION_KEY, INPUT_TYPE, STATUS,
                   TO_CHAR(START_TIME,'YYYY-MM-DD HH24:MI') start_t,
                   TO_CHAR(END_TIME,'YYYY-MM-DD HH24:MI') end_t,
                   TIME_TAKEN_DISPLAY AS elapsed_disp
            FROM v$rman_backup_job_details
            WHERE end_time > SYSDATE - 30
            ORDER BY end_time DESC
        """)
        results['rman_jobs'] = cur.fetchall()

        # 备份集大小从 v$backup_piece（而非 v$backup_set）取
        try:
            cur.execute("""
                SELECT p.handle, s.INPUT_TYPE,
                       ROUND(SUM(p.bytes)/1024/1024/1024, 2) size_gb,
                       MAX(p.compressed) compressed
                FROM v$backup_set s, v$backup_piece p
                WHERE p.set_stamp = s.set_stamp
                  AND p.set_count = s.set_count
                  AND p.completion_time > SYSDATE - 30
                GROUP BY p.handle, s.INPUT_TYPE
                ORDER BY MAX(p.completion_time) DESC
            """)
            results['backup_pieces'] = cur.fetchall()
        except Exception:
            results['backup_pieces'] = []
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_flashback_v19(conn):
    """闪回配置（v19 兼容版）：不用 v$flashback_database_stat，直接查 v$database 稳定列"""
    results = {}
    cur = conn.cursor()
    try:
        # v$database 标准列：FLASHBACK_ON / CREATED / LOG_MODE
        cur.execute("""
            SELECT FLASHBACK_ON,
                   TO_CHAR(CREATED,'YYYY-MM-DD HH24:MI') created,
                   LOG_MODE
            FROM v$database
        """)
        results['flashback'] = cur.fetchall()

        try:
            cur.execute("""
                SELECT owner, original_name, type,
                       ROUND(space * (SELECT TO_NUMBER(value) FROM v$parameter WHERE name='db_block_size')/1024/1024,2) mb,
                       can_undrop, can_purge
                FROM dba_recyclebin
                ORDER BY mb DESC
            """)
            results['recyclebin'] = cur.fetchall()
        except Exception:
            try:
                cur.execute("""
                    SELECT owner, original_name, type,
                           ROUND(space * (SELECT TO_NUMBER(value) FROM v$parameter WHERE name='db_block_size')/1024/1024,2) mb,
                           can_undrop, can_purge
                    FROM cdb_recyclebin
                    ORDER BY mb DESC
                """)
                results['recyclebin'] = cur.fetchall()
            except Exception:
                results['recyclebin'] = []
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_dataguard_v19(conn):
    """Data Guard（v19 兼容版）：不使用 STANDBY_DB_UNIQUE_NAME（可能不存在），改用安全列"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT GROUP#, TYPE, MEMBER, IS_RECOVERY_DEST_FILE
            FROM v$logfile
            WHERE TYPE = 'STANDBY'
        """)
        results['standby_logs'] = cur.fetchall()
    except Exception:
        results['standby_logs'] = []

    try:
        cur.execute("""
            SELECT dest_id, status, destination, archiver, transmit_mode
            FROM v$archive_dest
            WHERE destination IS NOT NULL
        """)
        results['archive_dest'] = cur.fetchall()
    except Exception:
        results['archive_dest'] = []

    try:
        # 不用 STANDBY_DB_UNIQUE_NAME（列名在部分环境不存在），只取存在的列
        cur.execute("""
            SELECT dest_id, database_mode, recovery_mode, protection_mode, status
            FROM v$archive_dest_status
            WHERE status != 'INACTIVE'
        """)
        results['dg_status'] = cur.fetchall()
    except Exception as e:
        results['dg_status'] = []
        results['dg_error'] = str(e)

    return results

def oracle_check_awr_v19(conn):
    """AWR 快照（v19 兼容版）：INTERVAL DAY TO SECOND 不能直接 * 24，用 EXTRACT 转换"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT instance_number, snap_id,
                   TO_CHAR(begin_interval_time,'YYYY-MM-DD HH24:MI') bt,
                   TO_CHAR(end_interval_time,'YYYY-MM-DD HH24:MI') et,
                   ROUND(EXTRACT(DAY FROM (end_interval_time - begin_interval_time)) * 24 +
                         EXTRACT(HOUR FROM (end_interval_time - begin_interval_time)) +
                         EXTRACT(MINUTE FROM (end_interval_time - begin_interval_time)) / 60, 2) elapsed_hr,
                   ERROR_COUNT
            FROM dba_hist_snapshot
            WHERE end_interval_time > SYSDATE - 7
            ORDER BY instance_number, snap_id DESC
        """)
        results['awr_snaps'] = cur.fetchall()

        cur.execute("SELECT * FROM dba_hist_wr_control")
        results['awr_settings'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results

def oracle_check_alert_v19(conn, days=7):
    """Alert 日志（v19 兼容版）：三重容错 v$diag_alert_text → v$diag_alert_xml → 直接读 ADR"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT TO_CHAR(alert_time,'YYYY-MM-DD HH24:MI:SS') t,
                   SUBSTR(message_text,1,200) message
            FROM v$diag_alert_text
            WHERE alert_time > SYSDATE - :days
              AND (message_text LIKE '%ORA-%' OR message_text LIKE '%ERROR%')
            ORDER BY alert_time DESC
        """, days=days)
        results['alert_errors'] = cur.fetchall()
    except Exception:
        try:
            cur.execute("""
                SELECT TO_CHAR(trap_time,'YYYY-MM-DD HH24:MI:SS') t,
                       SUBSTR(message_text,1,200) message
                FROM v$diag_alert_xml
                WHERE trap_time > SYSDATE - :days
                  AND (message_text LIKE '%ORA-%' OR message_text LIKE '%ERROR%')
                ORDER BY trap_time DESC
            """, days=days)
            results['alert_errors'] = cur.fetchall()
        except Exception:
            # ADR HOME 路径查不到时直接标记为无权限/不可访问
            results['alert_errors'] = []
            results['alert_hint'] = 'ADR视图不可访问（需 SYSDBA 权限或 CDB 环境）'
    finally:
        cur.close()
    return results

def getData(host, port, user, password, **kwargs):
    """Web UI 统一接口 - 返回 CompatWrapper"""
    ssh_info = kwargs.get('ssh_info', {})
    inspector_name = kwargs.get('inspector_name', 'Jack')
    service_name = kwargs.get('service_name')
    sid = kwargs.get('sid')
    sysdba = kwargs.get('sysdba', False)
    desensitize = kwargs.get('desensitize', False)
    template_id = kwargs.get('template_id')

    class _Args:
        pass
    args = _Args()
    args.host = host
    args.port = int(port)
    args.user = user
    args.password = password
    args.servicename = service_name
    args.sid = sid
    args.sysdba = bool(sysdba or user.upper() == 'SYS')
    if not args.sid and not args.servicename:
        args.sid = 'orcl'
    args.ssh_host = ssh_info.get('host') if ssh_info else None
    args.ssh_port = int(ssh_info.get('port', 22)) if ssh_info else 22
    args.ssh_user = ssh_info.get('user') if ssh_info else None
    args.ssh_pass = ssh_info.get('password') if ssh_info else None
    args.ssh_key = ssh_info.get('key_file') if ssh_info else None
    args.output = None
    args.zip = False
    args.inspector = inspector_name
    args.desensitize = desensitize
    args.template_id = template_id
    # Web UI 模式：跳过报告生成和历史保存
    args._web_ui_mode = True

    context = single_inspection(args)

    class CompatWrapper:
        def __init__(self, ctx, args_obj):
            self._context = ctx
            self._args = args_obj
            self.conn_db = None  # Oracle 连接在 single_inspection 内完成并关闭

        def checkdb(self, sqlfile=''):
            return self._context

        def generate_report(self, output_file, inspector_name="Jack"):
            db_info = self._context.get('_oracle_db_info', {})
            os_data = self._context.get('_oracle_os_data', {})
            check_results = self._context.get('_oracle_check_results', {})
            version_str = self._context.get('_oracle_version_str', '')
            ai_advice = self._context.get('_oracle_ai_advice', '')
            lang = self._context.get('_oracle_lang', 'zh')
            cb_result = self._context.get('_oracle_config_baseline')
            ih_result = self._context.get('_oracle_index_health')
            chapter_results = self._context.get('_oracle_chapter_results')
            has_template = chapter_results is not None and len(chapter_results) > 0
            docx = build_word_report(db_info, os_data, check_results, version_str, ai_advice,
                                      inspector=inspector_name, lang=lang,
                                      desensitize=bool(getattr(self._args, 'desensitize', False)),
                                      config_baseline_result=cb_result,
                                      index_health_result=ih_result,
                                      host=getattr(self._args, 'host', ''), health_status=self._context.get('health_status', ''),
                                      chapter_results=chapter_results if has_template else None)
            if docx:
                docx.save(output_file)
                return output_file
            return None

    return CompatWrapper(context, args)

if __name__ == '__main__':
    main()
