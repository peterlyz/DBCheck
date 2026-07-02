# -*- coding: utf-8 -*-
#
# Copyright (c) 2025-2026 fiyo (Jack Ge) <sdfiyon@gmail.com>
#
# This file is part of DBCheck, an open-source database health inspection tool.
# DBCheck is released under the MIT License with Attribution Requirements.
# See LICENSE for full license text.
#

"""
通用数据库巡检引擎
所有数据库类型的巡检模块都应继承此类！

架构：
1. BaseInspectionEngine - 包含所有通用逻辑（数据采集、风险分析、报告生成）
2. 子类只需实现：
   - connect(self): 连接数据库，返回 (ok: bool, version: str)
   - get_template_id(self): 返回 inspection_template 表的 template_id

通用引擎自动完成：
- 调用 connect() 连接数据库
- 调用 collect_data() 采集数据
- 调用 generate_report() 生成报告
"""
import os
import sys
import sqlite3
import traceback
import time
import re
from datetime import datetime
# docx 相关导入改为延迟导入（在 generate_report 相关方法中导入）
# from docx import Document
# from docx.shared import Pt, RGBColor, Inches, Cm
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.table import WD_TABLE_ALIGNMENT
# from docx.oxml.ns import qn
# from docx.oxml import parse_xml

import configparser
import importlib
import io
import json
import platform
import getpass

try:
    import psutil
except ImportError:
    psutil = None

try:
    import paramiko
except ImportError:
    paramiko = None

# ── 健康检查评分阈值 ─────────────────────────────
HEALTH_THRESHOLD = {'excellent': 90, 'good': 75, 'fair': 60, 'poor': 0}

# ── 磁盘采集时忽略的外接 ISO / Media 挂载点前缀 ─────────────────────────────────
IGNORE_MOUNTS = {'/mnt/iso', '/media', '/run/media', '/iso', '/cdrom'}


# ============================================================
# 远程系统信息收集器
# ============================================================
class RemoteSystemInfoCollector:
    """远程系统信息收集器 - 通过SSH连接获取远程主机信息"""
    
    def __init__(self, host, port=22, username='root', password=None, key_file=None):
        self.host = host
        self.port = port
        self.username = username
        self.password = password
        self.key_file = key_file
        self.ssh_client = None
    
    def connect(self):
        if paramiko is None:
            print("[WARN] paramiko 未安装，无法使用 SSH 功能")
            return False
        try:
            self.ssh_client = paramiko.SSHClient()
            self.ssh_client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            self.ssh_client.connect(
                hostname=self.host, port=self.port,
                username=self.username, password=self.password,
                timeout=10, look_for_keys=False, allow_agent=False,
                disabled_algorithms={'pubkeys': ['ssh-rsa']}
            )
            return True
        except Exception as e:
            print(f"SSH 连接失败 ({self.host}:{self.port}): {e}")
            return False
    
    def disconnect(self):
        if self.ssh_client:
            try: self.ssh_client.close()
            except Exception: pass
    
    def _run(self, cmd):
        stdin, stdout, stderr = self.ssh_client.exec_command(cmd, timeout=20)
        return stdout.read().decode('utf-8', errors='ignore'), stderr.read().decode('utf-8', errors='ignore')
    
    def get_cpu_info(self):
        out, _ = self._run("cat /proc/cpuinfo | grep 'model name' | head -1 && top -bn1 | grep 'Cpu' | awk '{print $2}'")
        info = {'Model name': '', 'usage_percent': 0.0}
        lines = out.strip().split('\n')
        if lines:
            info['Model name'] = lines[0].replace('model name', '').strip()
        if len(lines) > 1:
            try: info['usage_percent'] = float(lines[1].replace('%us',''))
            except: pass
        return info
    
    def get_memory_info(self):
        out, _ = self._run("free -m | grep Mem")
        info = {'total_mb': 0, 'used_mb': 0, 'usage_percent': 0.0}
        try:
            parts = out.strip().split()
            if len(parts) >= 3:
                info['total_mb'] = float(parts[1])
                info['used_mb'] = float(parts[2])
                if info['total_mb'] > 0:
                    info['usage_percent'] = round(info['used_mb'] / info['total_mb'] * 100, 1)
        except Exception: pass
        return info
    
    def get_disk_info(self):
        out, _ = self._run("df -h | grep -vE 'tmpfs|udev|overlay|mnt/iso|/iso|/media/|/run/media/|/cdrom|Filesystem'")
        disks = {}
        for line in out.split('\n')[1:]:
            parts = line.split()
            if len(parts) >= 6:
                device, size, used, avail, use_pct, mountpoint = parts[0], parts[1], parts[2], parts[3], parts[4], parts[5]
                try:
                    total_gb = float(size.replace('G','').replace('M',''))
                    used_gb = float(used.replace('G','').replace('M',''))
                    free_gb = float(avail.replace('G','').replace('M',''))
                    usage = float(use_pct.replace('%',''))
                    if 'M' in avail:
                        free_gb /= 1024; used_gb /= 1024; total_gb /= 1024
                    disks[mountpoint] = {
                        'device': device, 'mountpoint': mountpoint,
                        'total_gb': round(total_gb, 2), 'used_gb': round(used_gb, 2),
                        'free_gb': round(free_gb, 2), 'usage_percent': usage,
                        'fstype': 'unknown'
                    }
                except (ValueError, IndexError): continue
        return disks
    
    def get_system_info(self):
        import re
        boot_out, _ = self._run("uptime -s")
        platform_out, _ = self._run("uname -sr")
        os_release_out, _ = self._run("cat /etc/os-release 2>/dev/null || uname -a")
        # 从 /etc/os-release 解析 PRETTY_NAME（如 "CentOS Linux 7.9.2009"）
        platform_text = '未知'
        if os_release_out:
            m = re.search(r'PRETTY_NAME="([^"]+)"', os_release_out)
            if m:
                platform_text = m.group(1)
            elif os_release_out.strip():
                # fallback：使用 uname -a 前3个字段
                parts = os_release_out.strip().split(maxsplit=3)
                platform_text = ' '.join(parts[:3]) if len(parts) >= 3 else os_release_out.strip()
        return {
            'platform': platform_out.strip() if platform_out.strip() else 'Linux',
            'platform_text': platform_text,
            'boot_time': boot_out.strip() if boot_out.strip() else '未知',
            'cpu': self.get_cpu_info(),
            'memory': self.get_memory_info(),
            'disk': self.get_disk_info()
        }


# ============================================================
# 本地系统信息收集器
# ============================================================
class LocalSystemInfoCollector:
    """本地系统信息收集器"""
    
    def __init__(self):
        """初始化本地系统信息收集器。"""
        
    def get_cpu_info(self):
        info = {'Model name': '', 'usage_percent': 0.0}
        try:
            import platform as _pf
            info['Model name'] = f"{_pf.processor()} ({_pf.machine()})"
            if psutil:
                info['usage_percent'] = round(psutil.cpu_percent(interval=1), 1)
        except Exception: pass
        return info
    
    def get_memory_info(self):
        info = {'total_mb': 0, 'used_mb': 0, 'usage_percent': 0.0}
        try:
            if psutil:
                mem = psutil.virtual_memory()
                info['total_mb'] = round(mem.total / 1024 / 1024, 2)
                info['used_mb'] = round(mem.used / 1024 / 1024, 2)
                info['usage_percent'] = round(mem.percent, 1)
        except Exception: pass
        return info
    
    def get_disk_info(self):
        disks = {}
        try:
            if psutil:
                for part in psutil.disk_partitions():
                    try:
                        usage = psutil.disk_usage(part.mountpoint)
                        if part.mountpoint in IGNORE_MOUNTS:
                            continue
                        disks[part.mountpoint] = {
                            'device': part.device,
                            'mountpoint': part.mountpoint,
                            'fstype': part.fstype,
                            'total_gb': round(usage.total / 1024**3, 2),
                            'used_gb': round(usage.used / 1024**3, 2),
                            'free_gb': round(usage.free / 1024**3, 2),
                            'usage_percent': round(usage.percent, 1),
                        }
                    except (PermissionError, OSError): continue
        except Exception: pass
        return disks
    
    def get_system_info(self):
        import platform as _pf
        boot_time = ''
        try: 
            if psutil:
                boot_time = datetime.fromtimestamp(psutil.boot_time()).strftime('%Y-%m-%d %H:%M:%S')
        except: pass
        return {
            'platform': f"{_pf.system()} {_pf.release()} {_pf.machine()}",
            'platform_text': f"{_pf.system()} {_pf.release()} ({_pf.machine()})",
            'boot_time': boot_time,
            'cpu': self.get_cpu_info(),
            'memory': self.get_memory_info(),
            'disk': self.get_disk_info()
        }


def get_host_disk_usage():
    """获取本地磁盘使用情况（Windows 兼容）"""
    disks = []
    try:
        if psutil:
            for part in psutil.disk_partitions():
                if part.mountpoint in IGNORE_MOUNTS: continue
                try:
                    usage = psutil.disk_usage(part.mountpoint)
                    disks.append({
                        'device': part.device, 'mountpoint': part.mountpoint,
                        'fstype': part.fstype,
                        'total_gb': round(usage.total / 1024**3, 2),
                        'used_gb': round(usage.used / 1024**3, 2),
                        'free_gb': round(usage.free / 1024**3, 2),
                        'usage_percent': round(usage.percent, 1),
                    })
                except (PermissionError, OSError): continue
    except Exception: pass
    return disks


# ============================================================
# 通用巡检引擎基类
# ============================================================
class BaseInspectionEngine:
    """
    通用数据库巡检引擎 - 所有数据库类型的巡检模块都应继承此类！
    
    子类只需实现：
    - connect(self): 连接数据库，返回 (ok: bool, version: str)
    - get_template_id(self): 返回 inspection_template 表的 template_id
    
    通用引擎自动完成：
    - 调用 connect() 连接数据库
    - 调用 collect_data() 采集数据
    - 调用 generate_report() 生成报告
    """
    
    def __init__(self, host, port, user, password, database=None, ssh_info=None, template_id=None):
        self.host = host
        self.port = int(port)
        self.user = user
        self.password = password
        self.database = database
        self.ssh_info = ssh_info or {}
        self._template_id = template_id  # 用户显式指定的模板 ID
        self.conn = None
        self.cursor = None
        self.context = {}
        self._lang = 'zh'
        self.db_type = None  # 子类必须设置！
        self.output_file = None
        self.template_file = None
        self._safe_errors = []  # 记录安全跳过的章节及原因
        self._query_errors = {}  # {q_key: 友好错误描述}，渲染时使用
        
        try:
            from i18n import get_lang, t as _t
            self._lang = get_lang()
            self._t = _t
        except Exception:
            self._lang = 'zh'
            self._t = lambda x, **kw: x  # 备用：直接返回键名
    
    @staticmethod
    def _classify_sql_error(item_name, err_str):
        """将 SQL 执行错误分类为友好描述"""
        err_lower = err_str.lower()
        if 'invalid object name' in err_lower or '42s02' in err_lower or 'does not exist' in err_lower:
            return f'⚠️ 章节「{item_name}」：所需表/视图不存在（功能未配置），已跳过'
        if 'permission' in err_lower or 'denied' in err_lower or '42501' in err_lower:
            return f'⚠️ 章节「{item_name}」：权限不足，无法查询，已跳过'
        if 'connection' in err_lower or 'communic' in err_lower:
            return f'⚠️ 章节「{item_name}」：数据库连接已断开，已跳过'
        return f'⚠️ 章节「{item_name}」：查询失败（{err_str[:80]}），已跳过'

    @staticmethod
    def _clean_xml_str(v, max_len=200):
        """清理字符串中的 XML 不兼容控制字符（如 \x00 等）"""
        if v is None:
            return ''
        s = str(v)[:max_len]
        # 去掉 XML 不兼容的控制字符（ASCII 0-31，保留 \t \n \r）
        allowed = {'\t', '\n', '\r'}
        s = ''.join(c for c in s if ord(c) >= 32 or c in allowed)
        return s
    
    # ── 子类必须实现的方法 ────────────────────────
    def connect(self):
        """连接数据库 - 子类必须实现！"""
        raise NotImplementedError("子类必须实现 connect() 方法")
    
    def get_template_id(self):
        """获取模板 ID - 优先使用用户显式指定的 _template_id，否则根据 db_type 自动查询"""
        # 1. 优先返回用户显式指定的 template_id
        if self._template_id is not None:
            return self._template_id
        # 2. 否则根据 db_type 从数据库自动查询默认模板
        try:
            _db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'inspection.db')
            if not os.path.exists(_db_path):
                return None
            _conn = sqlite3.connect(_db_path)
            _cur = _conn.cursor()
            _cur.execute(
                "SELECT id FROM inspection_template WHERE db_type=? ORDER BY is_default DESC, id LIMIT 1",
                (self.db_type,))
            row = _cur.fetchone()
            _conn.close()
            return row[0] if row else None
        except Exception:
            return None
    
    
    # ── 通用方法（所有数据库共用）────────────────────────
    def _customize_queries(self, sql_dict):
        """子类 SQL 定制钩子 - 加载模板后可覆盖特定查询"""
        pass

    def collect_data(self, sql_templates=''):
        """采集数据 - 通用逻辑！"""
        print("\n" + self._t(f'{self.db_type}_start_ing', default='dm8_start_ing'))
        total_steps = 25
        current_step = 0
        
        # 1. 连接数据库
        ok, version = self.connect()
        if not ok:
            return False, version
        # 统一保存版本号到 context（供报告 + AI 诊断使用）
        # 注意：后续 init_keys 和 SQL 查询可能会覆盖此值，
        # 但 _append_chapters 读取时会兼容 'VERSION'/'version' 两种键名
        self.context['version'] = [{'VERSION': version}]
        
        # 2. 加载 SQL 模板（从 inspection.db 或内置）
        sql_dict = {}  # {query_key: sql_text}
        cfg = None     # 兼容 configparser（仅当 sql_templates 非空时使用）
        try:
            if not sql_templates or sql_templates == '':
                # 从 inspection.db 加载查询，直接用 dict 存储，避免 configparser 解析问题
                _db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'inspection.db')
                if not os.path.exists(_db_path):
                    print(self._t(f'{self.db_type}_sql_template_fail', default='dm8_sql_template_fail').format(e="inspection.db 不存在"))
                    return self.context

                import sqlite3
                conn_db = sqlite3.connect(_db_path)
                c = conn_db.cursor()
                if self._template_id is not None:
                    c.execute('''
                        SELECT q.query_key, q.query_sql
                        FROM inspection_query q
                        JOIN inspection_chapter ch ON q.chapter_id = ch.id
                        WHERE ch.template_id = ? AND q.enabled = 1
                        ORDER BY ch.sort_order, q.sort_order
                    ''', (self._template_id,))
                else:
                    c.execute('''
                        SELECT q.query_key, q.query_sql
                        FROM inspection_query q
                        JOIN inspection_chapter ch ON q.chapter_id = ch.id
                        JOIN inspection_template t ON ch.template_id = t.id
                        WHERE t.db_type = ? AND q.enabled = 1
                        ORDER BY ch.sort_order, q.sort_order
                    ''', (self.db_type,))
                queries = c.fetchall()

                # 直接用 dict 存储，不再经过 configparser
                for name, sql in queries:
                    if sql:
                        sql_clean = sql.replace('\n', ' ').replace('\r', ' ').strip()
                        # 自动修复旧版 pg_blocking_chain SQL（缺少 blocking_activity JOIN）
                        if name == 'pg_blocking_chain' and 'blocking_activity_block' not in sql_clean:
                            sql_clean = (
                                "SELECT blocked_locks.pid AS blocked_pid, blocked_activity.usename AS blocked_user, "
                                "left(blocked_activity.query, 200) AS blocked_query, "
                                "blocking_locks.pid AS blocking_pid, blocking_activity_block.usename AS blocking_user, "
                                "left(blocking_activity_block.query, 200) AS blocking_query "
                                "FROM pg_catalog.pg_locks blocked_locks "
                                "JOIN pg_catalog.pg_stat_activity blocked_activity ON blocked_activity.pid = blocked_locks.pid "
                                "JOIN pg_catalog.pg_locks blocking_locks ON blocking_locks.locktype = blocked_locks.locktype "
                                "AND blocking_locks.pid != blocked_locks.pid AND blocking_locks.granted "
                                "JOIN pg_catalog.pg_stat_activity blocking_activity_block ON blocking_activity_block.pid = blocking_locks.pid "
                                "WHERE NOT blocked_locks.granted ORDER BY blocked_activity.query_start;"
                            )
                            # 同步修复数据库
                            try:
                                conn_db.execute(
                                    "UPDATE inspection_query SET query_sql=? WHERE query_key='pg_blocking_chain'",
                                    (sql_clean,)
                                )
                                conn_db.commit()
                            except Exception:
                                pass
                        sql_dict[name] = sql_clean
                conn_db.close()
            else:
                # 兼容模式：外部传入 INI 字符串时用 configparser
                cfg = configparser.RawConfigParser()
                cfg.read_string(sql_templates)
        except Exception as e:
            print(self._t(f'{self.db_type}_sql_template_fail', default='dm8_sql_template_fail').format(e=e))
            import traceback; traceback.print_exc()
            return self.context

        # 2.5 子类 SQL 定制钩子（如 TiDB 覆盖 MySQL 不兼容的查询）
        self._customize_queries(sql_dict)

        # 3. 初始化 context keys
        init_keys = [
            "version", "instance", "database", "uptime",
            "sessions", "session_limit", "blocked", "trx",
            "tablespace", "temp_ts",
            "sga", "memory_info", "pga",
            "redo_logs", "redo_curr",
            "arch_config", "archive_lag",
            "params",
            "invalid_objs", "invalid_cnt",
            "users", "sys_privs", "default_pws",
            "long_sql", "top_sql_cpu",
            "rss_status", "rss_apply",
            "dcinfo", "inst_info",
            "undo_info", "transactions",
            "lock_blocking", "lock_deadlock", "lock_long_trx",
            "recyclebin",
            "datafiles",
            "profile_pwd",
            "top_waits", "wait_class",
            "stale_stats",
            "partition_info"
        ]
        for key in init_keys:
            self.context.update({key: []})

        # 4. 执行所有 SQL（容错执行器，单个失败不中断）
        _prog_prefix = self._t(f'{self.db_type}_progress_prefix', default='dm8_progress_prefix')
        try:
            cursor = self.conn.cursor()
            # IvorySQL/PG: 延迟检测 pg_stat_statements 扩展是否可用
            self._pg_stat_statements_available = None
            # 统一获取查询列表：优先 dict，兼容 configparser
            if sql_dict:
                variables_items = list(sql_dict.items())
            elif cfg and cfg.has_section("variables"):
                variables_items = list(cfg.items("variables"))
            else:
                variables_items = []
            for i, (name, stmt) in enumerate(variables_items):
                # IvorySQL/PG: 跳过需 pg_stat_statements 扩展的查询（未安装时静默跳过）
                if self.db_type in ('ivorysql', 'pg') and name in ('pg_top_elapsed', 'pg_top_calls'):
                    if not getattr(self, '_pg_stat_statements_available', None):
                        self.context[name] = []
                        continue
                    # 首次遇到时检测扩展
                    try:
                        cursor.execute("SELECT 1 FROM pg_extension WHERE extname='pg_stat_statements'")
                        self._pg_stat_statements_available = cursor.fetchone() is not None
                        if not self._pg_stat_statements_available:
                            self.context[name] = []
                            continue
                    except Exception:
                        self._pg_stat_statements_available = False
                        self.context[name] = []
                        continue
                current_step = int((i + 1) / len(variables_items) * (total_steps - 6)) + 1
                self.print_progress_bar(current_step, total_steps, prefix=_prog_prefix, suffix=f'{name} ({i+1}/{len(variables_items)})')
                # 清洗 SQL: 去除换行、CLI终止符(\G/\g)、末尾分号
                import re
                clean_sql = stmt.replace('\n', ' ').replace('\r', ' ')
                # 去除 MySQL CLI 专用的 \G/\g 终止符（处理单/双反斜杠情况）
                clean_sql = re.sub(r'\\+G\s*$', '', clean_sql, flags=re.IGNORECASE)
                clean_sql = clean_sql.strip().rstrip(';').strip()
                result = self._execute_query_safe(cursor, clean_sql, item_name=name)
                self.context[name] = result.get('data', [])
                time.sleep(0.03)
            cursor.close()
        except Exception as e:
            print(self._t(f'{self.db_type}_query_loop_fail', default='dm8_query_loop_fail').format(e=e))

        # MySQL: 从 cache_hit_ratio / cache_hit_requests 计算缓冲池命中率
        # SHOW GLOBAL STATUS LIKE 返回 Variable_name / Value 两列（所有 MySQL 版本兼容）
        if self.db_type == 'mysql':
            def _get_show_status_value(data):
                """从 SHOW GLOBAL STATUS 结果提取 Value"""
                if not data or not isinstance(data, list):
                    return 0
                row = data[0]
                if not isinstance(row, dict):
                    return 0
                # 兼容不同驱动的大小写差异: Value / value / VALUE
                val = row.get('Value') or row.get('value') or row.get('VALUE')
                if val is None:
                    return 0
                try:
                    return int(str(val).replace(',', ''))
                except (ValueError, TypeError):
                    return 0

            reads = _get_show_status_value(self.context.get('cache_hit_ratio', []))
            requests = _get_show_status_value(self.context.get('cache_hit_requests', []))
            if requests > 0:
                self.context['cache_hit_ratio_pct'] = round((1.0 - reads / requests) * 100, 2)
            else:
                self.context['cache_hit_ratio_pct'] = 0.0

        # 容错执行结果存入 context
        self.context['_safe_errors'] = self._safe_errors if hasattr(self, '_safe_errors') else []

        # 5. 加载章节结构（用于报告动态生成章节）
        self._load_chapters_from_db()

        # 6. 收集系统信息
        current_step = total_steps - 4
        self.print_progress_bar(current_step, total_steps, prefix=_prog_prefix, suffix=self._t(f'{self.db_type}_progress_sysinfo', default='dm8_progress_sysinfo'))
        try:
            if self.ssh_info and self.ssh_info.get('ssh_host'):
                print(self._t(f'{self.db_type}_ssh_collecting', default='dm8_ssh_collecting').format(host=self.ssh_info['ssh_host']))
                collector = RemoteSystemInfoCollector(
                    host=self.ssh_info['ssh_host'], port=self.ssh_info.get('ssh_port', 22),
                    username=self.ssh_info.get('ssh_user', 'root'),
                    password=self.ssh_info.get('ssh_password'), key_file=self.ssh_info.get('ssh_key_file')
                )
                if not collector.connect():
                    print(self._t(f'{self.db_type}_ssh_conn_fail_skip', default='dm8_ssh_conn_fail_skip'))
                    collector = LocalSystemInfoCollector()
            else:
                collector = LocalSystemInfoCollector()
            system_info = collector.get_system_info()
            disk_list = system_info.get('disk_list') or system_info.get('disk') or get_host_disk_usage()
            if isinstance(disk_list, dict):
                disk_list = list(disk_list.values())
            system_info['disk_list'] = disk_list
            self.context.update({"system_info": system_info})
        except Exception as e:
            print(self._t(f'{self.db_type}_sysinfo_fail', default='dm8_sysinfo_fail').format(e=e))
            self.context.update({"system_info": {
                'platform': '未知', 'boot_time': '未知',
                'cpu': {}, 'memory': {},
                'disk_list': [{'device':'C:','mountpoint':'C:\\','fstype':'NTFS',
                               'total_gb':0,'used_gb':0,'free_gb':0,'usage_percent':0}]
            }})
        
        # 7. 风险分析
        current_step = total_steps - 3
        self.print_progress_bar(current_step, total_steps, prefix=_prog_prefix, suffix=self._t(f'{self.db_type}_progress_health', default='dm8_progress_health'))
        self.context.update({"auto_analyze": []})
        self._basic_risk_check()
        
        # 始终执行健康评分
        health = self._analyze_health_status()
        self.context['health_analysis'] = health
        
        # 将告警合并到 auto_analyze
        existing_items = {a.get('col3', '') for a in self.context.get('auto_analyze', [])}
        for alert in health['alerts']:
            level = '高' if alert.startswith('[紧急]') else '中'
            title = alert.split(']', 1)[1].strip() if ']' in alert else alert[:50]
            if alert not in existing_items:
                self.context['auto_analyze'].append({
                    'col1': title,
                    'col2': f'{level}风险',
                    'col3': alert.replace('[紧急] ', '').replace('[关注] ', ''),
                    'col4': level,
                    'col5': 'DBA/系统管理员',
                    'fix_sql': ''
                })
        
        # 更新整体健康状态
        if health['status'] in ('严重',):
            self.context.update({"health_status": "需紧急处理"})
        elif health['status'] == '警告':
            self.context.update({"health_status": "需关注"})
        elif health['status'] == '一般':
            self.context.update({"health_status": "良好"})
        else:
            problem_count = len(self.context.get("auto_analyze", []))
            if problem_count == 0:
                self.context.update({"health_status": "优秀"})
            elif problem_count <= 3:
                self.context.update({"health_status": "良好"})
        
        # 8. AI 诊断
        current_step = total_steps - 2
        self.print_progress_bar(current_step, total_steps, prefix=_prog_prefix, suffix=self._t(f'{self.db_type}_progress_ai', default='dm8_progress_ai'))
        self.context['ai_advice'] = ''
        try:
            from analyzer import AIAdvisor
            import json as _json
            cfg_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'dbc_config.json')
            ai_cfg = {}
            if os.path.exists(cfg_path):
                with open(cfg_path, 'r', encoding='utf-8') as f:
                    ai_cfg = _json.load(f).get('ai', {})
            advisor = AIAdvisor(
                backend=ai_cfg.get('backend'),
                api_key=ai_cfg.get('api_key'),
                api_url=ai_cfg.get('api_url'),
                model=ai_cfg.get('model')
            )
            if advisor.enabled:
                label = self.context.get('co_name', [{}])[0].get('DB_NAME', 'Unknown')
                print(self._t(f'{self.db_type}_ai_calling', default="dm8_ai_calling").format(backend=advisor.backend, model=advisor.model))
                ai_advice = advisor.diagnose(self.db_type, label, self.context, self.context.get('auto_analyze', []), lang=self._lang)
                self.context['ai_advice'] = ai_advice
        except Exception:
            self.context['ai_advice'] = ''
        
        # 9. 慢查询深度分析（P2）— 按 db_type 选择分析器
        self.context['slow_query_result'] = None
        try:
            from slow_query_analyzer import get_slow_query_analyzer
            analyzer = get_slow_query_analyzer(self.db_type)
            if self.conn:
                ai_advisor = None
                try:
                    from analyzer import AIAdvisor
                    import json as _json
                    cfg_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'dbc_config.json')
                    ai_cfg = {}
                    if os.path.exists(cfg_path):
                        with open(cfg_path, 'r', encoding='utf-8') as f:
                            ai_cfg = _json.load(f).get('ai', {})
                    ai_advisor = AIAdvisor(
                        backend=ai_cfg.get('backend'),
                        api_key=ai_cfg.get('api_key'),
                        api_url=ai_cfg.get('api_url'),
                        model=ai_cfg.get('model')
                    )
                except Exception:
                    pass
                print("\n\U0001F50D " + self._t(f'{self.db_type}_slow_query_analyzing', default='正在进行慢查询深度分析...'))
                result = analyzer.analyze(self.conn, ai_advisor=ai_advisor, lang=self._lang)
                self.context['slow_query_result'] = result.to_dict()
                if result.is_empty():
                    print("  \u2139\uFE0F " + self._t(f'{self.db_type}_slow_query_unavailable', default='无慢查询数据，跳过深度分析'))
                else:
                    print("  \u2705 " + self._t(f'{self.db_type}_slow_query_ok', default='慢查询深度分析完成，采集到 {{count}} 条 Top SQL').format(
                        count=len(result.top_sql_by_latency)))
        except ValueError:
            pass  # 不支持的数据库类型，跳过
        except ImportError:
            pass
        except Exception as e:
            print("\u26A0\uFE0F 慢查询深度分析失败: %s" % e)
        
        # 10. 配置基线检查（P3）— 仅 DM8 支持
        self.context['config_baseline_result'] = None
        if self.db_type == 'dm':
            try:
                from config_baseline import check_dm_config_baseline
                if self.conn:
                    print("\n\U0001F539 " + self._t('dm8_cli_config_baseline_checking'))
                    cb_result = check_dm_config_baseline(self.conn)
                    self.context['config_baseline_result'] = cb_result
                    summary = cb_result.get('summary', {})
                    crit = summary.get('critical_count', 0)
                    warn = summary.get('warning_count', 0)
                    info = summary.get('info_count', 0)
                    print("  \u2705 " + self._t('dm8_cli_config_baseline_ok') % (crit, warn, info))
            except ImportError:
                pass
            except Exception as e:
                print("  \u26A0  配置基线检查失败: %s" % e)
        
        # 11. 索引健康分析（P3）— 仅 DM8 支持
        self.context['index_health_result'] = None
        if self.db_type == 'dm':
            try:
                from index_health import analyze_dm_indexes
                if self.conn:
                    print("\n\U0001F50D " + self._t('dm8_cli_index_health_checking'))
                    ih_result = analyze_dm_indexes(self.conn)
                    self.context['index_health_result'] = ih_result
                    sm = ih_result.get('summary', {})
                    miss = sm.get('missing_count', 0)
                    redun = sm.get('redundant_count', 0)
                    unused = sm.get('unused_count', 0)
                    print("  \u2705 " + self._t('dm8_cli_index_health_ok') % (miss, redun, unused))
            except ImportError:
                pass
            except Exception as e:
                print("  \u26A0  索引健康分析失败: %s" % e)
        
        self.print_progress_bar(total_steps, total_steps, prefix=_prog_prefix, suffix=self._t(f'{self.db_type}_progress_done', default='dm8_progress_done'))
        return True, version
    
    def generate_report(self, output_file, inspector_name="Jack"):
        """生成报告 - 通用逻辑！"""
        try:
            # 1. 加载 Word 模板（根据 db_type）
            self.template_file = self._load_word_template(inspector_name)
            
            # 2. 渲染上下文
            success = self._render_context(output_file, inspector_name)
            
            # 3. 追加章节（success 为 True 或 'fallback' 时都追加）
            if success:
                self._append_chapters(output_file)
            
            return output_file
        except Exception as e:
            print(f"[ERROR] 生成报告失败: {e}")
            import traceback; traceback.print_exc(file=sys.stdout)
            return None
    
    
    # ── 辅助方法（所有数据库共用）───────────────────────
    def _load_chapters_from_db(self):
        """从 inspection.db 加载章节结构"""
        try:
            _db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'inspection.db')
            if not os.path.exists(_db_path):
                return []
            
            _conn = sqlite3.connect(_db_path)
            _cur = _conn.cursor()
            
            # 获取模板 ID
            template_id = self.get_template_id()
            if not template_id:
                _conn.close()
                return []
            
            # 加载章节
            _cur.execute(
                "SELECT id, chapter_number, chapter_title_zh, chapter_title_en "
                "FROM inspection_chapter WHERE template_id=? ORDER BY chapter_number",
                (template_id,))
            
            _chapters = []
            for _ch_row in _cur.fetchall():
                _chapter_id = _ch_row[0]
                _chapter = {
                    'chapter_number': _ch_row[1],
                    'chapter_title_zh': _ch_row[2],
                    'chapter_title_en': _ch_row[3],
                    'queries': []
                }
                
                # 加载查询
                _cur.execute(
                    "SELECT query_key, query_description_zh, query_description_en "
                    "FROM inspection_query WHERE chapter_id=? ORDER BY sort_order, id",
                    (_chapter_id,))
                
                _chapter['queries'] = [
                    {'key': r[0], 'desc_zh': r[1], 'desc_en': r[2]}
                    for r in _cur.fetchall()
                ]
                
                _chapters.append(_chapter)
            
            _conn.close()
            self.context['_chapters'] = _chapters
            print("[OK] 已从 inspection.db 加载 %d 个章节" % len(_chapters))
            return _chapters
            
        except Exception as e:
            print("[WARN] 加载章节结构失败: %s" % e)
            self.context['_chapters'] = []
            return []
    
    def _execute_query_safe(self, cursor, sql, item_name=""):
        """安全执行 SQL — 报错时 rollback 当前事务，防止 PostgreSQL 级联失败
        
        返回: {"columns": [...], "data": [...], "_error": "友好错误描述"} 
        _error 键仅在有错误时存在，渲染时用于区分"查询失败"和"无数据"
        """
        try:
            cursor.execute(sql)
            columns = [col[0] for col in cursor.description]
            data = []
            for row in cursor.fetchall():
                data.append(dict(zip(columns, row)))
            return {"columns": columns, "data": data}
        except Exception as e:
            # 分类错误，生成友好描述
            err_str = str(e)
            friendly = self._classify_sql_error(item_name, err_str)
            print(f"[WARN] {friendly}")
            # 记录到实例字典和列表，供渲染和"巡检跳过记录"章节使用
            if not hasattr(self, '_safe_errors'):
                self._safe_errors = []
            self._safe_errors.append(friendly)
            if item_name:
                self._query_errors[item_name] = friendly
            # PostgreSQL: 一条 SQL 报错后事务被 abort，必须 rollback 才能继续执行后续 SQL
            if self.conn:
                try:
                    self.conn.rollback()
                except Exception:
                    pass
            return {"columns": [], "data": [], "_error": friendly}
    
    def _analyze_health_status(self):
        """健康评分"""
        alerts_critical = []
        alerts_warning = []
        
        # 1. 表空间使用率
        ts_list = self.context.get('tablespace', [])
        if ts_list and isinstance(ts_list, list):
            for ts in ts_list:
                if not isinstance(ts, dict): continue
                used_pct = self._safe_float_val(ts.get('USED_PCT') or ts.get('used_pct', 0))
                name = ts.get('TABLESPACE_NAME', ts.get('tablespace_name', '?'))
                total_bytes = self._safe_float_val(ts.get('TOTAL_BYTES') or ts.get('total_bytes', 0))
                free_bytes = self._safe_float_val(ts.get('FREE_BYTES') or ts.get('free_bytes', 0))
                total_mb = total_bytes / 1024 / 1024
                free_mb = free_bytes / 1024 / 1024
                
                if used_pct >= 95:
                    alerts_critical.append(
                        f"[紧急] 表空间 {name} 使用率 {used_pct:.1f}% "
                        f"(总计 {total_mb:.0f}MB，剩余 {free_mb:.0f}MB)，需立即扩容"
                    )
                elif used_pct >= 85:
                    alerts_warning.append(
                        f"[关注] 表空间 {name} 使用率 {used_pct:.1f}% "
                        f"(剩余 {free_mb:.0f}MB)，建议提前规划扩容"
                    )
        
        # 2. 会话数接近上限
        sess = self.context.get('sessions', [])
        limit = self.context.get('session_limit', [])
        if sess and limit and isinstance(sess, list) and isinstance(limit, list):
            try:
                total = self._safe_int_val(sess[0].get('TOTAL_SESSIONS', 0))
                max_proc = 0
                for lim in limit:
                    if isinstance(lim, dict) and lim.get('NAME', '') == 'processes':
                        max_proc = self._safe_int_val(lim.get('VALUE', 0))
                        break
                if max_proc > 0 and total > 0:
                    pct = total * 100 / max_proc
                    if pct >= 95:
                        alerts_critical.append(f"[紧急] 会话数 {total}/{max_proc} ({pct:.0f}%)，即将耗尽资源")
                    elif pct >= 85:
                        alerts_warning.append(f"[关注] 会话数 {total}/{max_proc} ({pct:.0f}%)，接近上限")
            except (IndexError, KeyError, TypeError):
                pass
        
        # 3. 锁等待与阻塞链
        blocked = self.context.get('blocked', [])
        blocking_detail = self.context.get('lock_blocking', [])
        deadlock = self.context.get('lock_deadlock', [])
        long_trx = self.context.get('lock_long_trx', [])
        
        # 3.1 锁等待计数
        if blocked and isinstance(blocked, list):
            lock_cnt = len(blocked)
            if lock_cnt >= 5:
                alerts_critical.append(f"[紧急] 发现 {lock_cnt} 个锁等待，可能导致业务阻塞")
            elif lock_cnt >= 2:
                alerts_warning.append(f"[关注] 发现 {lock_cnt} 个锁等待，建议排查")
        
        # 3.2 锁阻塞链详情
        if blocking_detail and isinstance(blocking_detail, list):
            blocking_cnt = len(blocking_detail)
            if blocking_cnt > 0:
                # 提取阻塞者与被阻塞者信息
                blocker_users = set()
                waiter_users = set()
                max_wait_ms = 0
                for bd in blocking_detail:
                    if not isinstance(bd, dict): continue
                    blocker_users.add(bd.get('blocker_user', ''))
                    waiter_users.add(bd.get('waiter_user', ''))
                    wait_ms = self._safe_int_val(bd.get('wait_ms', 0))
                    if wait_ms > max_wait_ms:
                        max_wait_ms = wait_ms
                alerts_warning.append(
                    f"[关注] 发现 {blocking_cnt} 个锁阻塞链，"
                    f"阻塞者: {', '.join(u for u in blocker_users if u)}，"
                    f"最长等待: {max_wait_ms}ms"
                )
        
        # 3.3 死锁检测
        if deadlock and isinstance(deadlock, list):
            dl_count = self._safe_int_val(deadlock[0].get('deadlock_count', 0)) if deadlock else 0
            if dl_count > 0:
                alerts_critical.append(f"[紧急] 检测到 {dl_count} 个死锁，需立即处理")
        
        # 3.4 长事务（>60秒）
        if long_trx and isinstance(long_trx, list):
            long_trx_cnt = len(long_trx)
            if long_trx_cnt > 0:
                max_dur = 0
                long_trx_users = set()
                for lt in long_trx:
                    if not isinstance(lt, dict): continue
                    dur = self._safe_int_val(lt.get('duration_sec', 0))
                    if dur > max_dur:
                        max_dur = dur
                    long_trx_users.add(lt.get('USER_NAME', ''))
                if max_dur > 300:  # > 5 分钟
                    alerts_critical.append(
                        f"[紧急] 发现 {long_trx_cnt} 个长事务，最长 {max_dur} 秒，"
                        f"用户: {', '.join(u for u in long_trx_users if u)}"
                    )
                else:
                    alerts_warning.append(
                        f"[关注] 发现 {long_trx_cnt} 个长事务（>60秒），最长 {max_dur} 秒"
                    )
        
        # 4. 无效对象
        invalid = self.context.get('invalid_cnt', [])
        if invalid and isinstance(invalid, list):
            total_invalid = sum(self._safe_int_val(iv.get('INVALID_COUNT', 0)) for iv in invalid if isinstance(iv, dict))
            if total_invalid >= 20:
                alerts_warning.append(f"[关注] 存在 {total_invalid} 个无效对象，建议编译或清理")
            elif total_invalid > 0:
                pass
        
        # 5. 系统内存
        sys_info = self.context.get('system_info', {})
        mem = sys_info.get('memory', {}) if isinstance(sys_info, dict) else {}
        if isinstance(mem, dict):
            mem_pct = self._safe_float_val(mem.get('usage_percent', 0))
            if mem_pct >= 95:
                alerts_critical.append(f"[紧急] 系统内存使用率 {mem_pct:.1f}%，存在 OOM 风险")
            elif mem_pct >= 90:
                alerts_warning.append(f"[关注] 系统内存使用率 {mem_pct:.1f}%")
        
        # 6. 磁盘空间
        disks = sys_info.get('disk_list', []) if isinstance(sys_info, dict) else []
        if disks and isinstance(disks, list):
            for d in disks:
                if not isinstance(d, dict): continue
                mp = d.get('mountpoint', '/')
                if mp in IGNORE_MOUNTS: continue
                usage = self._safe_float_val(d.get('usage_percent', 0))
                if usage >= 98:
                    alerts_critical.append(f"[紧急] 磁盘 {mp} 使用率 {usage:.1f}%，即将写满")
                elif usage >= 90:
                    alerts_warning.append(f"[关注] 磁盘 {mp} 使用率 {usage:.1f}%")
        
        # ── 综合评分 ─────────────────────────────────
        critical_n = len(alerts_critical)
        warning_n = len(alerts_warning)
        score = 100.0 - critical_n * 15 - warning_n * 5
        score = max(0, min(100, score))
        
        if critical_n > 0:
            status = self._t(f'report.{self.db_type}_status_critical', default='严重')
        elif warning_n >= 4:
            status = self._t(f'report.{self.db_type}_status_warning', default='警告')
        elif warning_n > 0:
            status = self._t(f'report.{self.db_type}_status_general', default='一般')
        else:
            status = self._t(f'report.{self.db_type}_health_status_ok', default='健康')
        
        return {
            "status": status,
            "score": round(score, 1),
            "critical_count": critical_n,
            "warning_count": warning_n,
            "alerts": alerts_critical + alerts_warning
        }
    
    def _basic_risk_check(self):
        """基础风险检查"""
        ts_list = self.context.get('tablespace', [])
        for ts in ts_list:
            if not isinstance(ts, dict): continue
            used_pct = self._safe_float_val(ts.get('USED_PCT', 0))
            name = ts.get('TABLESPACE_NAME', '?')
            if used_pct > 90:
                self.context['auto_analyze'].append({
                    'col1': f'表空间 {name}', 'col2': '高风险',
                    'col3': f'表空间使用率 {used_pct:.1f}%，超过 90% 告警线',
                    'col4': '高', 'col5': 'DBA',
                    'fix_sql': f"-- 查询表空间使用情况:\nSELECT * FROM DBA_TABLESPACES WHERE NAME='{name}';"
                })
            elif used_pct > 80:
                self.context['auto_analyze'].append({
                    'col1': f'表空间 {name}', 'col2': '中风险',
                    'col3': f'表空间使用率 {used_pct:.1f}%，建议关注',
                    'col4': '中', 'col5': 'DBA', 'fix_sql': ''
                })
        
        sess = self.context.get('sessions', [])
        limit = self.context.get('session_limit', [])
        if sess and limit:
            total = self._safe_int_val(sess[0], 'TOTAL_SESSIONS')
            max_sess = 0
            for l in limit:
                if isinstance(l, dict) and l.get('NAME') == 'processes':
                    max_sess = self._safe_int_val(l, 'VALUE')
                    break
            if max_sess > 0 and (total / max_sess) * 100 > 85:
                self.context['auto_analyze'].append({
                    'col1': '会话数接近上限', 'col2': '高风险',
                    'col3': f'当前会话 {total} / 上限 {max_sess}',
                    'col4': '高', 'col5': 'DBA',
                    'fix_sql': '-- 查看会话:\nSELECT * FROM V$SESSION WHERE TYPE=\'USER\';'
                })
        
        mem = self.context.get('system_info', {}).get('memory', {})
        if isinstance(mem, dict) and mem.get('usage_percent', 0) > 90:
            self.context['auto_analyze'].append({
                'col1': '系统内存紧张', 'col2': '高风险',
                'col3': f'内存使用率 {mem["usage_percent"]:.1f}%',
                'col4': '高', 'col5': '系统管理员', 'fix_sql': ''
            })
        
        for disk in self.context.get('system_info', {}).get('disk_list', []):
            if not isinstance(disk, dict): continue
            mp = disk.get('mountpoint', '/')
            if mp in IGNORE_MOUNTS: continue
            usage = self._safe_float_val(disk.get('usage_percent', 0))
            if usage > 90:
                self.context['auto_analyze'].append({
                    'col1': f'磁盘空间不足 ({mp})', 'col2': '高风险',
                    'col3': f'磁盘 {mp} 使用率 {usage:.1f}%',
                    'col4': '高', 'col5': '系统管理员', 'fix_sql': ''
                })
    
    def _t(self, key):
        """i18n 翻译"""
        try:
            from i18n import t
            return t(key, self._lang)
        except Exception:
            return key
    
    def _safe_float_val(self, val, default=0.0):
        try:
            if val is None: return default
            return float(str(val).replace(',', '').replace('%', '').strip())
        except (ValueError, TypeError):
            return default
    
    def _safe_int_val(self, val, default=0):
        try:
            if val is None: return default
            return int(str(val).replace(',', '').strip())
        except (ValueError, TypeError):
            return default
    
    def print_progress_bar(self, iteration, total, prefix='', suffix='', decimals=1, length=50, fill='█'):
        """打印进度条"""
        percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
        filled_length = int(length * iteration // total)
        bar = fill * filled_length + '-' * (length - filled_length)
        print(f'\r{prefix} |{bar}| {percent}% {suffix}', end='\r')
        if iteration == total:
            print()
    
    def _load_word_template(self, inspector_name="Jack"):
        """加载 Word 模板（创建中文封面模板，支持 docxtpl 渲染）"""
        # 延迟导入 docx 相关模块
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import parse_xml
        
        template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")
        tpl_file = os.path.join(template_path, f"{self.db_type}_wordtemplates_v1.0.docx")
        
        if os.path.exists(tpl_file):
            return tpl_file
        
        # 创建模板（中文封面）
        doc = Document()
        
        db_type_display = {
            'dm8': 'DM8', 'mysql': 'MySQL', 'postgresql': 'PostgreSQL',
            'oracle': 'Oracle', 'sqlserver': 'SQL Server',
            'tidb': 'TiDB', 'ivorysql': 'IvorySQL'
        }.get(self.db_type, self.db_type.upper())
        
        # Logo
        logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'dbcheck_logo.png')
        if os.path.exists(logo_path):
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_para.add_run()
            logo_run.add_picture(logo_path, width=Cm(3.5))
        
        # 主标题（中文）
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(f'{db_type_display} 数据库健康巡检报告')
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(15, 75, 135)
        title_run.font.name = '微软雅黑'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 副标题（英文）
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run('Database Health Inspection Report')
        sub_run.font.size = Pt(14)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(100, 100, 100)
        sub_run.font.name = 'Times New Roman'
        
        doc.add_paragraph()
        
        # 封面信息表格
        info_table = doc.add_table(rows=5, cols=2, style='Table Grid')
        info_labels = ['服务器地址', '实例启动时间', '巡检结果', '巡检人员', '报告生成时间']
        for i, label in enumerate(info_labels):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = ''  # 空值，由渲染时填充
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = '微软雅黑'
                cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                if cell == row.cells[0]:
                    from docx.oxml.ns import nsdecls
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="336699"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    cell.paragraphs[0].runs[0].font.bold = True
        
        doc.add_page_break()
        
        os.makedirs(template_path, exist_ok=True)
        doc.save(tpl_file)
        return tpl_file
    
    def _render_context(self, output_file, inspector_name="Jack"):
        """从数据库加载模板配置，渲染 _chapters 到 context"""
        try:
            from inspection_dal import (
                get_default_template,
                get_chapters_by_template,
                get_queries_by_chapter,
            )

            # 1. 获取当前 db_type 的默认模板
            template = get_default_template(self.db_type)
            if not template:
                print(f"[WARN] 未找到 {self.db_type} 的模板配置，跳过章节渲染")
                self.context['_chapters'] = []
                return self._fallback_render(output_file, inspector_name)

            template_id = template['id']
            tpl_name = template.get('template_name_zh', '') or template.get('template_name', '')
            print(f"[INFO] 使用模板: {tpl_name} (id={template_id})")

            # 2. 加载章节和查询
            chapters = get_chapters_by_template(template_id)
            _chapters = []

            total_queries = 0
            for ch in chapters:
                chapter_id = ch['id']
                queries = get_queries_by_chapter(chapter_id)
                chapter_queries = []
                for q in queries:
                    q_key = q['query_key']
                    sql = q.get('query_sql', '')
                    desc_zh = q.get('query_description_zh', '') or ''
                    desc_en = q.get('query_description_en', '') or ''
                    name_zh = q.get('query_name_zh', '') or ''
                    name_en = q.get('query_name_en', '') or ''

                    # 数据已在 collect_data 阶段执行并存入 self.context，此处直接使用
                    if q_key not in self.context:
                        self.context[q_key] = []

                    chapter_queries.append({
                        'key': q_key,
                        'query_key': q_key,
                        'query_name_zh': name_zh,
                        'query_name_en': name_en,
                        'query_description_zh': desc_zh,
                        'query_description_en': desc_en,
                    })
                    total_queries += 1

                _chapters.append({
                    'chapter_number': ch.get('chapter_number', ch.get('sort_order', 0)),
                    'chapter_title_zh': ch.get('chapter_title_zh', ''),
                    'chapter_title_en': ch.get('chapter_title_en', ''),
                    'queries': chapter_queries,
                })

            self.context['_chapters'] = _chapters
            print(f"[INFO] 已加载 {len(_chapters)} 个章节，共 {total_queries} 个查询")

            # 3. 填充必要的 context 字段（兼容旧逻辑）
            required_keys = ['auto_analyze', 'version', 'co_name', 'system_info']
            for key in required_keys:
                if key not in self.context or not self.context[key]:
                    if key == 'auto_analyze':
                        self.context[key] = []
                    elif key == 'version':
                        self.context[key] = [{'VERSION': 'Unknown'}]
                    elif key == 'system_info':
                        self.context[key] = {}
                    elif key == 'co_name':
                        self.context[key] = [{'DB_NAME': f'{self.user}@{self.host}:{self.port}'}]
                    else:
                        self.context[key] = [{'placeholder': self._t("report.data_missing")}]

            if 'disk_list' not in self.context.get('system_info', {}) or not self.context['system_info'].get('disk_list'):
                self.context.setdefault('system_info', {})['disk_list'] = [{
                    'device': 'C:', 'mountpoint': 'C:\\', 'fstype': 'NTFS',
                    'total_gb': 50.0, 'used_gb': 25.0, 'free_gb': 25.0, 'usage_percent': 50.0
                }]

            self.context.update({"report_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S')})
            self.context.update({"inspector_name": inspector_name})

            problem_count = len(self.context.get("auto_analyze", []))
            self.context.update({"problem_count": problem_count})

            if problem_count == 0: health_status = self._t("report.health_excellent")
            elif problem_count <= 3: health_status = self._t("report.health_good")
            elif problem_count <= 6: health_status = self._t("report.health_fair")
            else: health_status = self._t("report.health_attention")
            self.context.update({"health_status": health_status})

            # 4. 执行基线配置检查
            self._check_baselines()
            
            # 5. 调用 fallback 渲染
            return self._fallback_render(output_file, inspector_name)
        
        except Exception as e:
            print(f"[ERROR] 渲染上下文失败: {e}")
            import traceback; traceback.print_exc(file=sys.stdout)
            return False
    
    def _check_baselines(self):
        """执行基线配置检查，结果存入 self.context['baseline_results']"""
        try:
            from inspection_dal import get_baselines_by_db_type
            
            baselines = get_baselines_by_db_type(self.db_type, enabled_only=True)
            if not baselines:
                print(f"[INFO] 未找到 {self.db_type} 的基线配置")
                self.context['baseline_results'] = []
                return
            
            results = []
            for bl in baselines:
                param_name = bl['param_name']
                query_sql = bl['query_sql']
                operator = bl['operator']
                expected_value = bl['expected_value']
                expected_value_min = bl['expected_value_min']
                expected_value_max = bl['expected_value_max']
                risk_level = bl['risk_level']
                description_zh = bl.get('description_zh', '')
                description_en = bl.get('description_en', '')
                
                # 执行查询获取当前值
                try:
                    _cur = self.conn.cursor() if self.conn else None
                    if not _cur:
                        raise Exception("数据库连接已关闭")
                    _cur.execute(query_sql)
                    rows = _cur.fetchall()
                    _cur.close()
                    
                    if not rows:
                        results.append({
                            'param_name': param_name,
                            'current_value': 'N/A',
                            'expected_value': expected_value,
                            'operator': operator,
                            'status': 'UNKNOWN',
                            'risk_level': risk_level,
                            'description_zh': description_zh,
                            'description_en': description_en,
                            'message': '查询无返回结果'
                        })
                        continue
                    
                    # 假设查询结果只有一行一列
                    current_value = str(rows[0][0]) if rows[0] else 'N/A'
                    
                    # 执行基线对比
                    status = self._compare_baseline(
                        current_value, operator,
                        expected_value, expected_value_min, expected_value_max
                    )
                    
                    results.append({
                        'param_name': param_name,
                        'current_value': current_value,
                        'expected_value': expected_value,
                        'expected_value_min': expected_value_min,
                        'expected_value_max': expected_value_max,
                        'operator': operator,
                        'status': status,
                        'risk_level': risk_level,
                        'description_zh': description_zh,
                        'description_en': description_en,
                    })
                    
                except Exception as e:
                    error_str = str(e)
                    # PG/IvorySQL: 回滚防止事务级联 abort
                    if self.db_type in ('pg', 'ivorysql') and self.conn:
                        try:
                            self.conn.rollback()
                        except Exception:
                            pass
                    # MySQL: 不存在 query_cache_size 等系统变量时跳过
                    if 'Unknown system variable' in error_str:
                        results.append({
                            'param_name': param_name,
                            'current_value': 'N/A (变量不存在)',
                            'expected_value': expected_value,
                            'operator': operator,
                            'status': 'SKIP',
                            'risk_level': risk_level,
                            'description_zh': description_zh,
                            'description_en': description_en,
                            'message': '当前版本不存在此系统变量'
                        })
                    # PG/IvorySQL: 不存在的配置参数跳过（如非 ORAMODE 下的 ivorysql.oracle_compatibility）
                    elif 'unrecognized configuration parameter' in error_str:
                        results.append({
                            'param_name': param_name,
                            'current_value': 'N/A (参数不存在)',
                            'expected_value': expected_value,
                            'operator': operator,
                            'status': 'SKIP',
                            'risk_level': risk_level,
                            'description_zh': description_zh,
                            'description_en': description_en,
                            'message': '当前版本不存在此配置参数'
                        })
                    else:
                        print(f"[WARN] 基线检查失败: {param_name}, 错误: {e}")
                        results.append({
                            'param_name': param_name,
                            'current_value': 'ERROR',
                            'expected_value': expected_value,
                            'operator': operator,
                            'status': 'ERROR',
                            'risk_level': risk_level,
                            'description_zh': description_zh,
                            'description_en': description_en,
                            'message': error_str
                        })
            
            self.context['baseline_results'] = results
            print(f"[INFO] 基线检查完成，共 {len(results)} 项")
            
        except Exception as e:
            print(f"[ERROR] 基线检查失败: {e}")
            import traceback; traceback.print_exc(file=sys.stdout)
            self.context['baseline_results'] = []

    def _render_markdown_to_doc(self, doc, md_text, chapter_num=None):
        """将 Markdown 文本渲染为 Word 段落（支持标题、粗体、斜体、列表、代码块）

        Args:
            doc: Document 对象
            md_text: Markdown 文本
            chapter_num: 章节号（如 21），用于给 ## 子标题加序号 21.1, 21.2...
        """
        import re
        lines = md_text.split('\n')
        # 去除首尾空行，避免渲染出多余空白段落
        while lines and not lines[0].strip():
            lines.pop(0)
        while lines and not lines[-1].strip():
            lines.pop()

        i = 0
        sub_heading_idx = 0  # 子标题计数器
        prev_blank = False   # 用于合并连续空行

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # 空行 — 连续空行只保留一个
            if not stripped:
                if not prev_blank:
                    doc.add_paragraph()
                prev_blank = True
                i += 1
                continue
            prev_blank = False

            # 标题 # ## ###
            heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
            if heading_match:
                level = min(len(heading_match.group(1)), 6)
                title_text = heading_match.group(2).strip()

                # ## 标题（h2）加上子章节序号：如 21.1 重点关注
                if level == 2 and chapter_num is not None:
                    sub_heading_idx += 1
                    title_text = f'{chapter_num}.{sub_heading_idx} {title_text}'

                p = doc.add_paragraph()
                r = p.add_run(title_text)
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.bold = True
                r.font.size = Pt(max(9, 16 - level * 2))
                r.font.color.rgb = RGBColor(0, 51, 102)
                i += 1
                continue

            # 代码块 ``` ... ```
            if stripped.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    i += 1  # skip closing ```
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1)
                    p.paragraph_format.shading.BackgroundPatternColor = RGBColor(245, 245, 245)
                    r = p.add_run('\n'.join(code_lines).strip())
                    r.font.name = 'Consolas'
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0, 0, 0)
                continue

            # 无序列表项 - xxx / * xxx
            ul_match = re.match(r'^[\-\*]\s+(.*)', stripped)
            if ul_match:
                content = ul_match.group(1)
                p = doc.add_paragraph(style='List Bullet')
                self._add_inline_md_runs(p, content)
                i += 1
                continue

            # 有序列表项 1. xxx
            ol_match = re.match(r'^\d+\.\s+(.*)', stripped)
            if ol_match:
                content = ol_match.group(1)
                p = doc.add_paragraph(style='List Number')
                self._add_inline_md_runs(p, content)
                i += 1
                continue

            # 分割行 ---
            if re.match(r'^(-{3,}|\*{3,}|_{3,})$', stripped):
                doc.add_paragraph()
                i += 1
                continue

            # 普通段落
            p = doc.add_paragraph()
            self._add_inline_md_runs(p, stripped)
            i += 1

    def _add_inline_md_runs(self, paragraph, text):
        """在段落中解析内联 Markdown（粗体、斜体、行内代码）并添加 run"""
        import re

        def make_run(par, txt, bold=False, italic=False, code=False):
            r = par.add_run(txt)
            r.font.name = 'Consolas' if code else '微软雅黑'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r.font.size = Pt(10.5)
            r.font.bold = bold
            r.font.italic = italic
            if code:
                r.font.color.rgb = RGBColor(180, 0, 0)

        # 正则：匹配 **粗体** __粗体__ `行内代码` *斜体* _斜体_
        # 注意：\4 在 lookahead 中引用后续组会报错，拆成两个模式
        bold_pattern = re.compile(r'(\*\*|__)(.*?)\1')
        code_pattern = re.compile(r'`([^`]+)`')
        italic_pattern = re.compile(r'(?<!\*)(\*)(?!\*)(.+?)(?<!\*)\1(?!\*)|(?<![_a-zA-Z0-9])(_)(?!_)(.+?)(?<![_a-zA-Z0-9])\2(?!_)')
        pos = 0
        # 合并所有匹配位置
        matches = []
        for m in bold_pattern.finditer(text):
            matches.append((m.start(), m.end(), 'bold', m.group(2)))
        for m in italic_pattern.finditer(text):
            if m.group(1):  # *italic*
                matches.append((m.start(), m.end(), 'italic', m.group(2)))
            elif m.group(3):  # _italic_
                matches.append((m.start(), m.end(), 'italic', m.group(3)))
        for m in code_pattern.finditer(text):
            matches.append((m.start(), m.end(), 'code', m.group(1)))
        matches.sort(key=lambda x: x[0])
        for m_start, m_end, m_type, m_content in matches:
            if m_start > pos:
                make_run(paragraph, text[pos:m_start])
            if m_type == 'bold':
                make_run(paragraph, m_content, bold=True)
            elif m_type == 'italic':
                make_run(paragraph, m_content, italic=True)
            elif m_type == 'code':
                make_run(paragraph, m_content, code=True)
            pos = m_end

        if pos < len(text):
            make_run(paragraph, text[pos:])

    def _compare_baseline(self, current_value, operator, expected_value, expected_value_min, expected_value_max):
        """对比当前值与期望值，返回状态（PASS/FAIL）"""
        try:
            # 尝试转换为数值
            try:
                current_num = float(current_value)
                expected_num = float(expected_value) if expected_value else None
            except (ValueError, TypeError):
                current_num = None
                expected_num = None
            
            if operator == '=':
                if current_value == expected_value:
                    return 'PASS'
                else:
                    return 'FAIL'
            
            elif operator == '!=':
                if current_value != expected_value:
                    return 'PASS'
                else:
                    return 'FAIL'
            
            elif operator in ['>', '>=', '<', '<=']:
                if current_num is None or expected_num is None:
                    return 'ERROR'
                if operator == '>' and current_num > expected_num:
                    return 'PASS'
                elif operator == '>=' and current_num >= expected_num:
                    return 'PASS'
                elif operator == '<' and current_num < expected_num:
                    return 'PASS'
                elif operator == '<=' and current_num <= expected_num:
                    return 'PASS'
                else:
                    return 'FAIL'
            
            elif operator == 'BETWEEN':
                if current_num is None:
                    return 'ERROR'
                if expected_value_min and expected_value_max:
                    min_num = float(expected_value_min)
                    max_num = float(expected_value_max)
                    if min_num <= current_num <= max_num:
                        return 'PASS'
                    else:
                        return 'FAIL'
                else:
                    return 'ERROR'
            
            elif operator == 'LIKE':
                if expected_value and expected_value in current_value:
                    return 'PASS'
                else:
                    return 'FAIL'
            
            else:
                return 'ERROR'
                
        except Exception as e:
            print(f"[WARN] 基线对比失败: {e}")
            return 'ERROR'

    def _append_chapters(self, output_file):
        """追加章节（通用逻辑）"""
        try:
            doc = Document(output_file)
            
            def _set_cell_bg(cell, hex_color):
                from docx.oxml.ns import nsdecls
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
                cell._tc.get_or_add_tcPr().append(shading)
            
            def _add_heading(text, level=1):
                h = doc.add_heading(text, level=level)
                for run in h.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.color.rgb = RGBColor(0, 51, 102)
                    run.font.size = Pt(14) if level == 1 else Pt(12)
                return h
            
            def _ch_prefix(n, sub=None):
                if self._lang == 'zh':
                    return f"第{n}章" if sub is None else f"{n}.{sub}"
                else:
                    return f"Chapter {n}:" if sub is None else f"{n}.{sub}"

            # ── 动态计算章节起始号 ──────────────────────────
            chapters = self.context.get('_chapters', [])
            max_ch = max((ch.get('chapter_number', 0) for ch in chapters), default=0)
            ch_bl   = max_ch + 1   # 基线配置检查结果
            ch_risk = max_ch + 2   # 风险与建议
            ch_ai   = max_ch + 3   # AI 诊断分析
            print(f"[INFO] _append_chapters: DB max_ch={max_ch}, bl={ch_bl}, risk={ch_risk}, ai={ch_ai}")

            # ── 基线配置检查结果 ───────────────────────────
            baseline_results = self.context.get('baseline_results', [])
            print(f"[INFO] baseline_results count: {len(baseline_results)}")
            if baseline_results:
                _add_heading(f"{_ch_prefix(ch_bl)} {'基线配置检查结果' if self._lang == 'zh' else 'Baseline Configuration Check Results'}")
                tbl = doc.add_table(rows=1+len(baseline_results), cols=7, style='Table Grid')
                bl_headers = ['参数名称', '当前值', '期望值', '运算符', '状态', '风险等级', '描述'] if self._lang == 'zh' else ['Parameter', 'Current Value', 'Expected Value', 'Operator', 'Status', 'Risk Level', 'Description']
                for j, h in enumerate(bl_headers):
                    cell = tbl.rows[0].cells[j]
                    cell.text = h
                    _set_cell_bg(cell, '336699')
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9); run.font.name = '微软雅黑'; run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                for idx, br in enumerate(baseline_results, 1):
                    row = tbl.rows[idx].cells
                    row[0].text = br.get('param_name', '')
                    row[1].text = str(br.get('current_value', ''))
                    row[2].text = str(br.get('expected_value', '')) if br.get('expected_value') else f"{br.get('expected_value_min', '')} - {br.get('expected_value_max', '')}"
                    row[3].text = br.get('operator', '')
                    status = br.get('status', 'ERROR')
                    row[4].text = status
                    if status == 'PASS':
                        _set_cell_bg(row[4], '00B050')
                    elif status == 'FAIL':
                        _set_cell_bg(row[4], 'FF0000')
                    else:
                        _set_cell_bg(row[4], 'FFA500')
                    row[5].text = br.get('risk_level', '')
                    desc = br.get('description_zh' if self._lang == 'zh' else 'description_en', '')
                    row[6].text = desc if desc else br.get('description_en' if self._lang == 'zh' else 'description_zh', '')
                    for j in range(7):
                        for p in row[j].paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9); run.font.name = '微软雅黑'
                doc.add_paragraph()

            # ── 风险与建议 ─────────────────────────────────
            print(f"[INFO] auto_analyze count: {len(self.context.get('auto_analyze', []))}")
            _add_heading(f"{_ch_prefix(ch_risk)} {self._t(f'report.{self.db_type}_ch16', default='风险与建议')}")
            issues = self.context.get("auto_analyze", [])
            if issues:
                _add_heading(self._t(f'report.{self.db_type}_ch16_1', default='智能分析问题明细'), 2)
                tbl = doc.add_table(rows=1+len(issues), cols=7, style='Table Grid')
                hdrs = [
                    self._t(f'report.{self.db_type}_col_seq', default='序号'),
                    self._t(f'report.{self.db_type}_col_item', default='检查项'),
                    self._t(f'report.{self.db_type}_col_risk_level', default='风险等级'),
                    self._t(f'report.{self.db_type}_col_desc', default='描述'),
                    self._t(f'report.{self.db_type}_col_severity', default='严重度'),
                    self._t(f'report.{self.db_type}_col_owner', default='负责人'),
                    self._t(f'report.{self.db_type}_col_fix', default='修复建议'),
                ]
                for j, (cell, ht) in enumerate(zip(tbl.rows[0].cells, hdrs)):
                    cell.text = ht
                    _set_cell_bg(cell, '336699')
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.size = Pt(9); run.font.name = '微软雅黑'; run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                for idx, x in enumerate(issues, 1):
                    row = tbl.rows[idx].cells
                    row[0].text = str(idx)
                    row[1].text = x.get('col1','')
                    row[2].text = x.get('col2','')
                    row[3].text = x.get('col3','')
                    row[4].text = x.get('col4','')
                    row[5].text = x.get('col5','')
                    row[6].text = x.get('fix_sql','')[:200]
                    for j in range(7):
                        for p in row[j].paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9); run.font.name = '微软雅黑'
                doc.add_paragraph()
            else:
                p = doc.add_paragraph(self._t(f'report.{self.db_type}_no_risk_found', default='未发现明显风险项，数据库整体运行状况良好。'))
                for r in p.runs:
                    r.font.size = Pt(10.5); r.font.name = '微软雅黑'

            # ── AI 诊断分析 ────────────────────────────────
            ai_advice = self.context.get('ai_advice', '')
            print(f"[INFO] ai_advice length: {len(ai_advice)}")
            _add_heading(f"{_ch_prefix(ch_ai)} {self._t(f'report.{self.db_type}_ch17', default='AI 诊断分析')}")
            if ai_advice and ai_advice.strip():
                disclaimer = self._t('report.ai_disclaimer')
                dp = doc.add_paragraph()
                dr = dp.add_run(disclaimer)
                dr.font.size = Pt(9.5)
                dr.font.name = '微软雅黑'
                dr.font.italic = True
                dr.font.color.rgb = RGBColor(128, 128, 128)
                doc.add_paragraph()
                self._render_markdown_to_doc(doc, ai_advice.strip(), chapter_num=ch_ai)
            else:
                ap = doc.add_paragraph()
                ar = ap.add_run(self._t('report.ai_disabled'))
                ar.font.size = Pt(10.5)
                ar.font.name = '微软雅黑'
                ar.font.color.rgb = RGBColor(128, 128, 128)

            doc.save(output_file)
            return True
        except Exception as e:
            print(f"[ERROR] 追加章节失败: {e}")
            return False
    
    def _fallback_render(self, output_file, inspector_name="Jack"):
        """动态渲染：纯 python-docx 构建报告（无模板依赖，支持中英文）"""
        try:
            doc = Document()
            
            def _set_cell_bg(cell, hex_color):
                from docx.oxml.ns import nsdecls
                from docx.oxml import parse_xml
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
                cell._tc.get_or_add_tcPr().append(shading)
            
            def _fmt_heading(text, level=1):
                h = doc.add_heading(text, level=level)
                for run in h.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.color.rgb = RGBColor(0, 51, 102)
                    run.font.size = Pt(16) if level == 1 else Pt(12)
                    run.font.bold = True
                return h
            
            def _fmt_para(text, bold=False, size=10.5, color=None):
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(size)
                run.font.bold = bold
                if color:
                    run.font.color.rgb = color
                return p
            
            # ── 封面 ────────────────────────────────────────
            logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'dbcheck_logo.png')
            if os.path.exists(logo_path):
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_para.add_run()
                logo_run.add_picture(logo_path, width=Cm(3.5))
            
            # 封面标题（根据语言）
            is_zh = self._lang == 'zh'
            db_type_display = {
                'dm8': 'DM8', 'mysql': 'MySQL', 'postgresql': 'PostgreSQL',
                'oracle': 'Oracle', 'sqlserver': 'SQL Server',
                'tidb': 'TiDB', 'ivorysql': 'IvorySQL'
            }.get(self.db_type, self.db_type.upper())
            
            title_text = f'{db_type_display} 数据库健康巡检报告' if is_zh else f'{db_type_display} Database Health Inspection Report'
            subtitle_text = 'Database Health Inspection Report' if is_zh else ''
            
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title_text)
            title_run.font.size = Pt(28)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(15, 75, 135)
            title_run.font.name = '微软雅黑'
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            if subtitle_text:
                sub_para = doc.add_paragraph()
                sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub_run = sub_para.add_run(subtitle_text)
                sub_run.font.size = Pt(14)
                sub_run.font.italic = True
                sub_run.font.color.rgb = RGBColor(100, 100, 100)
                sub_run.font.name = 'Times New Roman'
            
            doc.add_paragraph()  # 空行
            
            # 封面信息表格
            report_time = self.context.get('report_time', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            health = self.context.get('health_status', 'Unknown')

            # 获取实例启动时间
            sys_info = self.context.get('system_info', {})
            boot_time = sys_info.get('boot_time', 'Unknown') if isinstance(sys_info, dict) else 'Unknown'

            cover_labels = {
                'zh': ['服务器地址', '实例启动时间', '巡检结果', '巡检人员', '报告生成时间'],
                'en': ['Server Address', 'Instance Start Time', 'Status', 'Inspector', 'Report Time']
            }
            cover_data = [
                f"{self.host}:{self.port}",
                boot_time,
                health,
                inspector_name,
                report_time
            ]
            
            labels = cover_labels['zh'] if is_zh else cover_labels['en']
            tbl = doc.add_table(rows=len(labels), cols=2, style='Table Grid')
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, (label, value) in enumerate(zip(labels, cover_data)):
                tbl.rows[i].cells[0].text = label
                tbl.rows[i].cells[1].text = str(value)
                for cell in tbl.rows[i].cells:
                    cell.paragraphs[0].runs[0].font.name = '微软雅黑'
                    cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                    if cell == tbl.rows[i].cells[0]:
                        _set_cell_bg(cell, '336699')
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                        cell.paragraphs[0].runs[0].font.bold = True
            
            doc.add_page_break()
            
            # ── 章节内容 ────────────────────────────────────
            chapters = self.context.get('_chapters', [])
            if chapters:
                for ch in chapters:
                    ch_num = ch.get('chapter_number', 0)
                    ch_title = ch.get('chapter_title_zh' if is_zh else 'chapter_title_en', '') or ch.get('chapter_title_zh', '') or ch.get('chapter_title_en', '')
                    queries = ch.get('queries', [])
                    
                    if not queries:
                        continue
                    
                    # 章节标题
                    prefix = f'第{ch_num}章' if is_zh else f'Chapter {ch_num}'
                    _fmt_heading(f'{prefix} {ch_title}', level=1)
                    
                    query_idx = 0
                    for q in queries:
                        q_key = q.get('key', '')
                        q_desc = q.get('query_description_zh' if is_zh else 'query_description_en', '') or q.get('query_description_en' if is_zh else 'query_description_zh', '')
                        q_data = self.context.get(q_key, [])
                        query_idx += 1
                        sub_num = f'{ch_num}.{query_idx}'

                        # 子章节标题：二级标题 + 序号
                        if q_desc:
                            _fmt_heading(f'{sub_num} {q_desc}', level=2)
                        else:
                            _fmt_heading(f'{sub_num} {q_key}', level=2)

                        if q_data and isinstance(q_data, list) and len(q_data) > 0:
                            if isinstance(q_data[0], dict):
                                headers = list(q_data[0].keys())
                                # 宽表（>8列）采用卡片式垂直布局
                                if len(headers) > 8:
                                    for ri, row_data in enumerate(q_data):
                                        card_title = f"{'条目' if is_zh else 'Item'} {ri+1}"
                                        cp = doc.add_paragraph()
                                        cr = cp.add_run(card_title)
                                        cr.font.size = Pt(10)
                                        cr.font.name = '微软雅黑'
                                        cr.font.bold = True
                                        cr.font.color.rgb = RGBColor(0, 51, 102)
                                        card = doc.add_table(rows=len(headers), cols=2, style='Table Grid')
                                        for kj, h in enumerate(headers):
                                            card.rows[kj].cells[0].text = str(h)
                                            card.rows[kj].cells[1].text = str(row_data.get(h, ''))[:500] if row_data.get(h) else ''
                                            # Key 列背景+粗体
                                            _set_cell_bg(card.rows[kj].cells[0], 'D9E2F3')
                                            for run in card.rows[kj].cells[0].paragraphs[0].runs:
                                                run.font.size = Pt(9); run.font.name = '微软雅黑'
                                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                                                run.font.bold = True
                                            for run in card.rows[kj].cells[1].paragraphs[0].runs:
                                                run.font.size = Pt(9); run.font.name = '微软雅黑'
                                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                                        doc.add_paragraph()
                                else:
                                    qt = doc.add_table(rows=1+len(q_data), cols=len(headers), style='Table Grid')
                                    for j, h in enumerate(headers):
                                        cell = qt.rows[0].cells[j]
                                        cell.text = self._clean_xml_str(h, max_len=100)
                                        _set_cell_bg(cell, '336699')
                                        for run in cell.paragraphs[0].runs:
                                            run.font.size = Pt(9)
                                            run.font.name = '微软雅黑'
                                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                                            run.font.bold = True
                                            run.font.color.rgb = RGBColor(255, 255, 255)
                                    for i, row_data in enumerate(q_data):
                                        for j, (k, v) in enumerate(row_data.items()):
                                            qt.rows[i+1].cells[j].text = self._clean_xml_str(v)
                                            for run in qt.rows[i+1].cells[j].paragraphs[0].runs:
                                                run.font.size = Pt(9)
                                                run.font.name = '微软雅黑'
                                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                                    doc.add_paragraph()
                        else:
                            # 数据为空：区分"查询失败"和"当前无记录"
                            _err = self._query_errors.get(q_key, '')
                            if _err:
                                # 查询失败：显示友好原因
                                _msg = _err.replace('⚠️ 章节「', '').replace('」：', '：')
                                _p = doc.add_paragraph(_msg)
                            else:
                                # 真正无数据
                                _no_data = '（当前无记录）' if is_zh else '(No data at this time)'
                                _p = doc.add_paragraph(self._t('report.data_missing') + _no_data if is_zh else 'No data ' + _no_data)
                            for _r in _p.runs:
                                _r.font.size = Pt(10); _r.font.name = '微软雅黑'
                                _r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                                _r.font.color.rgb = RGBColor(150, 150, 150)
                            doc.add_paragraph()
            else:
                # 如果没有章节结构，直接渲染所有列表数据
                _fmt_heading(self._t(f'report.{self.db_type}_ch_data', default='巡检数据' if is_zh else 'Inspection Data'), level=1)
                for key in sorted(self.context.keys()):
                    if key.startswith('_') or key in ('auto_analyze', 'system_info', 'health_analysis', 'report_time', 'inspector_name', 'problem_count', 'health_status', 'co_name', 'version'):
                        continue
                    val = self.context.get(key)
                    if val and isinstance(val, list) and len(val) > 0 and isinstance(val[0], dict):
                        _fmt_para(key, bold=True, size=11)
                        headers = list(val[0].keys())
                        qt = doc.add_table(rows=1+len(val), cols=len(headers), style='Table Grid')
                        for j, h in enumerate(headers):
                            cell = qt.rows[0].cells[j]
                            cell.text = self._clean_xml_str(h, max_len=100)
                            _set_cell_bg(cell, '336699')
                            for run in cell.paragraphs[0].runs:
                                run.font.size = Pt(9); run.font.name = '微软雅黑'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                                run.font.bold = True; run.font.color.rgb = RGBColor(255, 255, 255)
                        for i, row_data in enumerate(val):
                            for j, (k, v) in enumerate(row_data.items()):
                                qt.rows[i+1].cells[j].text = self._clean_xml_str(v)
                                for run in qt.rows[i+1].cells[j].paragraphs[0].runs:
                                    run.font.size = Pt(9); run.font.name = '微软雅黑'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                        doc.add_paragraph()
            
            doc.save(output_file)
            return 'fallback'
        except Exception as e:
            print(f"[ERROR] 动态渲染失败: {e}")
            import traceback; traceback.print_exc(file=sys.stdout)
            return False
