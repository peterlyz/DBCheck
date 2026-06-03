# -*- coding: utf-8 -*-
"""
server_inspect.py — 独立服务器巡检模块

从 main_mysql.py 提取并扩展，支持 SSH 远程采集和本地 psutil 采集。
新增检查项：Top 进程、内核版本、系统负载、inode、打开文件描述符。
"""

import os
import json
import platform
import socket
import sqlite3
from datetime import datetime

try:
    import paramiko
except ImportError:
    paramiko = None

try:
    import psutil
except ImportError:
    psutil = None

# ─── i18n ─────────────────────────────────────────────────────────────

_SI_LANG = 'zh'

def _t(key):
    try:
        from i18n import t
        return t(key, _SI_LANG)
    except Exception:
        return key

# ─── ISO/光盘文件系统过滤 ──────────────────────────────────────────────

# 需要过滤的文件系统类型（光盘/ISO/快照）
_ISO_FSTYPES = frozenset(['iso9660', 'udf', 'squashfs'])

# 需要过滤的挂载点关键词（精确匹配路径段）
_ISO_MOUNT_KEYWORDS = ('/mnt/iso', '/iso', '/cdrom', '/media/cdrom', '/snap')


def _is_iso_mount(mountpoint, fstype=''):
    """判断是否为 ISO/光盘挂载点"""
    if fstype.lower() in _ISO_FSTYPES:
        return True
    mp_lower = mountpoint.lower()
    for kw in _ISO_MOUNT_KEYWORDS:
        if kw in mp_lower:
            return True
    return False


# ─── 远程系统信息收集器 ────────────────────────────────────────────────

class RemoteSystemInfoCollector:
    """通过 SSH 连接采集远程主机系统信息"""

    def __init__(self, host, port=22, username='root', password=None, key_file=None):
        self.host = host
        self.port = int(port)
        self.username = username
        self.password = password
        self.key_file = key_file
        self.ssh_client = None

    def connect(self):
        if not paramiko:
            return False
        try:
            self.ssh_client = paramiko.SSHClient()
            self.ssh_client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            if self.key_file:
                private_key = paramiko.RSAKey.from_private_key_file(self.key_file)
                self.ssh_client.connect(hostname=self.host, port=self.port,
                                        username=self.username, pkey=private_key, timeout=15)
            else:
                self.ssh_client.connect(hostname=self.host, port=self.port,
                                        username=self.username, password=self.password,
                                        timeout=15, look_for_keys=False, allow_agent=False)
            return True
        except Exception:
            return False

    def disconnect(self):
        if self.ssh_client:
            try:
                self.ssh_client.close()
            except Exception:
                pass

    def exec_cmd(self, command, timeout=10):
        try:
            stdin, stdout, stderr = self.ssh_client.exec_command(command, timeout=timeout)
            out = stdout.read().decode('utf-8', errors='ignore').strip()
            err = stderr.read().decode('utf-8', errors='ignore').strip()
            return out, err
        except Exception:
            return '', ''

    def get_cpu_info(self):
        try:
            out, _ = self.exec_cmd("top -bn1 | grep 'Cpu(s)' | awk '{print $2}' | cut -d'%' -f1")
            cpu_percent = float(out) if out else 0.0

            out, _ = self.exec_cmd("nproc")
            logical_cores = int(out) if out else 0

            out, _ = self.exec_cmd("lscpu | grep 'Core(s) per socket' | awk '{print $4}'")
            cores_per_socket = int(out) if out else 1

            out, _ = self.exec_cmd("lscpu | grep 'Socket(s)' | awk '{print $2}'")
            sockets = int(out) if out else 1
            physical_cores = cores_per_socket * sockets

            out, _ = self.exec_cmd("lscpu | grep 'CPU MHz' | awk '{print $3}'")
            current_freq = float(out) / 1000 if out else 0.0

            out, _ = self.exec_cmd("lscpu | grep 'CPU max MHz' | awk '{print $4}'")
            max_freq = float(out) / 1000 if out else current_freq

            # 系统负载
            out, _ = self.exec_cmd("cat /proc/loadavg")
            load_1m, load_5m, load_15m = 0.0, 0.0, 0.0
            if out:
                parts = out.split()
                try:
                    load_1m = float(parts[0])
                    load_5m = float(parts[1])
                    load_15m = float(parts[2])
                except (IndexError, ValueError):
                    pass

            return {
                'usage_percent': cpu_percent,
                'physical_cores': physical_cores,
                'logical_cores': logical_cores,
                'current_frequency': round(current_freq, 2),
                'max_frequency': round(max_freq, 2),
                'load_1m': load_1m,
                'load_5m': load_5m,
                'load_15m': load_15m,
            }
        except Exception:
            return {}

    def get_memory_info(self):
        try:
            out, _ = self.exec_cmd("free -b | grep Mem")
            if not out:
                return {}
            parts = out.split()
            total_bytes = int(parts[1])
            used_bytes = int(parts[2])
            available_bytes = int(parts[6]) if len(parts) > 6 else int(parts[3])
            usage_percent = (used_bytes / total_bytes) * 100 if total_bytes > 0 else 0
            mem = {
                'total_gb': round(total_bytes / (1024**3), 2),
                'available_gb': round(available_bytes / (1024**3), 2),
                'used_gb': round(used_bytes / (1024**3), 2),
                'usage_percent': round(usage_percent, 2),
            }
            out, _ = self.exec_cmd("free -b | grep Swap")
            if out:
                parts = out.split()
                swap_total = int(parts[1])
                swap_used = int(parts[2])
                swap_pct = (swap_used / swap_total) * 100 if swap_total > 0 else 0
                mem.update({
                    'swap_total_gb': round(swap_total / (1024**3), 2),
                    'swap_used_gb': round(swap_used / (1024**3), 2),
                    'swap_usage_percent': round(swap_pct, 2),
                })
            return mem
        except Exception:
            return {}

    def get_disk_info(self):
        try:
            # 用 df -T 获取文件系统类型，用 df -i 获取 inode
            out, _ = self.exec_cmd("df -Th | grep -vE 'tmpfs|devtmpfs|overlay|iso9660|udf' | tail -n +2")
            disk_data = []
            if not out:
                return disk_data

            # 同时采集 inode
            inode_out, _ = self.exec_cmd("df -i | grep -vE 'tmpfs|devtmpfs|overlay|iso9660|udf' | tail -n +2")
            inode_map = {}
            if inode_out:
                for line in inode_out.strip().split('\n'):
                    p = line.split()
                    if len(p) >= 6:
                        inode_map[p[5]] = p[4].rstrip('%')

            for line in out.strip().split('\n'):
                parts = line.split()
                if len(parts) >= 7:
                    device = parts[0]
                    fstype = parts[1]
                    mountpoint = parts[6]
                    if any(vfs in device for vfs in ['tmpfs', 'devtmpfs', 'overlay']):
                        continue
                    # 过滤 ISO/光盘挂载点
                    if _is_iso_mount(mountpoint, fstype):
                        continue

                    def _to_gb(s):
                        s = s.strip().upper()
                        if s.endswith('G'): return round(float(s[:-1]), 2)
                        elif s.endswith('M'): return round(float(s[:-1]) / 1024, 2)
                        elif s.endswith('T'): return round(float(s[:-1]) * 1024, 2)
                        elif s.endswith('K'): return round(float(s[:-1]) / (1024**2), 2)
                        elif s.endswith('E'): return round(float(s[:-1]) * 1024 * 1024, 2)
                        else:
                            try: return round(float(s), 2)
                            except: return 0.0

                    total_gb = _to_gb(parts[2])
                    used_gb = _to_gb(parts[3])
                    free_gb = _to_gb(parts[4])
                    try:
                        usage_percent = float(parts[5].rstrip('%'))
                    except:
                        usage_percent = 0.0

                    inode_pct = 0.0
                    try:
                        inode_pct = float(inode_map.get(mountpoint, 0))
                    except (ValueError, TypeError):
                        pass

                    disk_data.append({
                        'device': device,
                        'mountpoint': mountpoint,
                        'fstype': fstype,
                        'total_gb': total_gb,
                        'used_gb': used_gb,
                        'free_gb': free_gb,
                        'usage_percent': usage_percent,
                        'inode_usage_percent': inode_pct,
                    })
            return disk_data
        except Exception:
            return []

    def get_top_processes(self, limit=5):
        """采集 Top N 进程（按 CPU 排序）"""
        try:
            out, _ = self.exec_cmd(
                f"ps aux --sort=-%cpu | head -{limit + 1} | tail -{limit}"
            )
            procs = []
            if out:
                for line in out.strip().split('\n'):
                    parts = line.split(None, 10)
                    if len(parts) >= 11:
                        procs.append({
                            'pid': int(parts[1]),
                            'user': parts[0],
                            'cpu': float(parts[2]),
                            'mem': float(parts[3]),
                            'command': parts[10][:80],
                        })
            return procs
        except Exception:
            return []

    def get_open_files(self):
        """采集打开文件描述符数量"""
        try:
            out, _ = self.exec_cmd("cat /proc/sys/fs/file-nr")
            if out:
                parts = out.split()
                return {'allocated': int(parts[0]), 'max': int(parts[2])}
            return {}
        except Exception:
            return {}

    def get_system_info(self):
        if not self.connect():
            return {}
        try:
            info = {
                'cpu': self.get_cpu_info(),
                'memory': self.get_memory_info(),
                'disk': self.get_disk_info(),
                'top_processes': self.get_top_processes(),
                'open_files': self.get_open_files(),
                'hostname': '',
                'platform': '',
                'kernel': '',
                'boot_time': '',
                'uptime_days': 0,
            }
            out, _ = self.exec_cmd("hostname")
            if out:
                info['hostname'] = out

            out, _ = self.exec_cmd("uname -a")
            if out:
                info['platform'] = out

            out, _ = self.exec_cmd("uname -r")
            if out:
                info['kernel'] = out

            out, _ = self.exec_cmd("who -b | awk '{print $3 \" \" $4}'")
            if out:
                info['boot_time'] = out

            # 计算 uptime 天数
            out, _ = self.exec_cmd("cat /proc/uptime")
            if out:
                try:
                    uptime_sec = float(out.split()[0])
                    info['uptime_days'] = round(uptime_sec / 86400, 1)
                except (ValueError, IndexError):
                    pass

            return info
        finally:
            self.disconnect()


# ─── 本地系统信息收集器 ────────────────────────────────────────────────

class LocalSystemInfoCollector:
    """使用 psutil 采集当前主机系统信息，无需 SSH"""

    def __init__(self):
        pass

    def get_cpu_info(self):
        if not psutil:
            return {}
        try:
            cpu_percent = psutil.cpu_percent(interval=1)
            cpu_count = psutil.cpu_count(logical=False)
            cpu_count_logical = psutil.cpu_count(logical=True)
            cpu_freq = psutil.cpu_freq()

            load_1m, load_5m, load_15m = 0.0, 0.0, 0.0
            try:
                if hasattr(os, 'getloadavg'):
                    load_1m, load_5m, load_15m = os.getloadavg()
            except OSError:
                pass

            return {
                'usage_percent': cpu_percent,
                'physical_cores': cpu_count or 0,
                'logical_cores': cpu_count_logical or 0,
                'current_frequency': round(cpu_freq.current, 2) if cpu_freq else 0,
                'max_frequency': round(cpu_freq.max, 2) if cpu_freq else 0,
                'load_1m': round(load_1m, 2),
                'load_5m': round(load_5m, 2),
                'load_15m': round(load_15m, 2),
            }
        except Exception:
            return {}

    def get_memory_info(self):
        if not psutil:
            return {}
        try:
            mem = psutil.virtual_memory()
            swap = psutil.swap_memory()
            return {
                'total_gb': round(mem.total / (1024**3), 2),
                'available_gb': round(mem.available / (1024**3), 2),
                'used_gb': round(mem.used / (1024**3), 2),
                'usage_percent': mem.percent,
                'swap_total_gb': round(swap.total / (1024**3), 2),
                'swap_used_gb': round(swap.used / (1024**3), 2),
                'swap_usage_percent': swap.percent,
            }
        except Exception:
            return {}

    def get_disk_info(self):
        if not psutil:
            return []
        try:
            disk_list = []
            partitions = psutil.disk_partitions()
            for partition in partitions:
                mp = partition.mountpoint
                fstype = partition.fstype or ''
                # 过滤 loop 设备和 ISO/光盘挂载点
                if fstype and 'loop' not in partition.device and not _is_iso_mount(mp, fstype):
                    try:
                        usage = psutil.disk_usage(mp)
                        disk_list.append({
                            'device': partition.device,
                            'mountpoint': mp,
                            'fstype': partition.fstype,
                            'total_gb': round(usage.total / (1024**3), 2),
                            'used_gb': round(usage.used / (1024**3), 2),
                            'free_gb': round(usage.free / (1024**3), 2),
                            'usage_percent': usage.percent,
                            'inode_usage_percent': 0.0,
                        })
                    except PermissionError:
                        continue
            return disk_list
        except Exception:
            return []

    def get_top_processes(self, limit=5):
        if not psutil:
            return []
        try:
            procs = []
            for proc in psutil.process_iter(['pid', 'username', 'cpu_percent', 'memory_percent', 'name']):
                try:
                    info = proc.info
                    procs.append({
                        'pid': info['pid'],
                        'user': info.get('username', ''),
                        'cpu': info.get('cpu_percent', 0.0),
                        'mem': info.get('memory_percent', 0.0),
                        'command': info.get('name', ''),
                    })
                except (psutil.NoSuchProcess, psutil.AccessDenied):
                    continue
            procs.sort(key=lambda x: x['cpu'], reverse=True)
            return procs[:limit]
        except Exception:
            return []

    def get_open_files(self):
        if not psutil:
            return {}
        try:
            # Windows 不支持此方法
            if platform.system() == 'Windows':
                return {}
            allocated = psutil.os.sysconf('SC_OPEN_MAX') if hasattr(psutil.os, 'sysconf') else 0
            return {'allocated': allocated, 'max': allocated}
        except Exception:
            return {}

    def get_system_info(self):
        hostname = socket.gethostname()
        plat = platform.platform()
        kernel = platform.release()
        boot_time_str = ''
        uptime_days = 0
        if psutil:
            try:
                bt = datetime.fromtimestamp(psutil.boot_time())
                boot_time_str = bt.strftime('%Y-%m-%d %H:%M:%S')
                uptime_days = round((datetime.now() - bt).total_seconds() / 86400, 1)
            except Exception:
                pass

        return {
            'cpu': self.get_cpu_info(),
            'memory': self.get_memory_info(),
            'disk': self.get_disk_info(),
            'top_processes': self.get_top_processes(),
            'open_files': self.get_open_files(),
            'hostname': hostname,
            'platform': plat,
            'kernel': kernel,
            'boot_time': boot_time_str,
            'uptime_days': uptime_days,
        }


# ─── 工厂类 ──────────────────────────────────────────────────────────

class SystemInfoCollector:
    @staticmethod
    def create(host_type='local', **kwargs):
        if host_type == 'remote':
            return RemoteSystemInfoCollector(**kwargs)
        return LocalSystemInfoCollector()


# ─── 健康评分 ─────────────────────────────────────────────────────────

def compute_health_score(info):
    """
    根据系统信息计算健康评分（0-100）和问题列表。
    返回 (score, status, issues)
    """
    score = 100
    issues = []

    # CPU
    cpu = info.get('cpu', {})
    cpu_pct = cpu.get('usage_percent', 0)
    if cpu_pct > 90:
        score -= 20
        issues.append(f"CPU 使用率过高 ({cpu_pct:.1f}% > 90%)")
    elif cpu_pct > 70:
        score -= 10
        issues.append(f"CPU 使用率偏高 ({cpu_pct:.1f}% > 70%)")

    # 内存
    mem = info.get('memory', {})
    mem_pct = mem.get('usage_percent', 0)
    if mem_pct > 90:
        score -= 20
        issues.append(f"内存使用率过高 ({mem_pct:.1f}% > 90%)")
    elif mem_pct > 75:
        score -= 10
        issues.append(f"内存使用率偏高 ({mem_pct:.1f}% > 75%)")

    # Swap
    swap_pct = mem.get('swap_usage_percent', 0)
    if swap_pct > 50:
        score -= 10
        issues.append(f"Swap 使用率过高 ({swap_pct:.1f}% > 50%)")

    # 磁盘（跳过 ISO/光盘挂载）
    for disk in info.get('disk', []):
        mp = disk.get('mountpoint', '/')
        fstype = disk.get('fstype', '')
        if _is_iso_mount(mp, fstype):
            continue  # ISO 挂载不参与评分
        pct = disk.get('usage_percent', 0)
        if pct > 90:
            score -= 15
            issues.append(f"磁盘 {mp} 使用率危险 ({pct:.1f}% > 90%)")
        elif pct > 80:
            score -= 8
            issues.append(f"磁盘 {mp} 使用率偏高 ({pct:.1f}% > 80%)")

    # inode（跳过 ISO/光盘挂载）
    for disk in info.get('disk', []):
        mp = disk.get('mountpoint', '/')
        fstype = disk.get('fstype', '')
        if _is_iso_mount(mp, fstype):
            continue
        inode_pct = disk.get('inode_usage_percent', 0)
        if inode_pct > 80:
            score -= 10
            issues.append(f"inode {mp} 使用率偏高 ({inode_pct:.1f}%)")

    score = max(0, min(100, score))

    if score >= 90:
        status = '优秀'
    elif score >= 75:
        status = '良好'
    elif score >= 60:
        status = '一般'
    else:
        status = '需关注'

    return score, status, issues


# ─── 服务器巡检主入口 ─────────────────────────────────────────────────

def run_server_inspection(ssh_host, ssh_port=22, ssh_user='root',
                          ssh_password='', ssh_key_file=''):
    """
    执行服务器巡检，返回结果字典。

    :param ssh_host: SSH 主机地址
    :param ssh_port: SSH 端口
    :param ssh_user: SSH 用户名
    :param ssh_password: SSH 密码
    :param ssh_key_file: SSH 私钥文件路径
    :return: 巡检结果字典
    """
    collector = RemoteSystemInfoCollector(
        host=ssh_host, port=int(ssh_port),
        username=ssh_user,
        password=ssh_password if ssh_password else None,
        key_file=ssh_key_file if ssh_key_file else None,
    )
    info = collector.get_system_info()
    if not info:
        return {'error': 'SSH 连接失败或采集超时'}

    # 网络连通性检测
    try:
        info['network'] = check_network_connectivity(collector)
    except Exception:
        info['network'] = {}

    # 服务状态检测（需要重新连接）
    try:
        collector2 = RemoteSystemInfoCollector(
            host=ssh_host, port=int(ssh_port),
            username=ssh_user,
            password=ssh_password if ssh_password else None,
            key_file=ssh_key_file if ssh_key_file else None,
        )
        if collector2.connect():
            info['services'] = check_service_status(collector2)
            collector2.disconnect()
        else:
            info['services'] = []
    except Exception:
        info['services'] = []

    score, status, issues = compute_health_score(info)
    info['health_score'] = score
    info['health_status'] = status
    info['issues'] = issues
    info['host'] = ssh_host
    return info


def run_local_inspection():
    """执行本地服务器巡检（无需 SSH）"""
    collector = LocalSystemInfoCollector()
    info = collector.get_system_info()
    if not info:
        return {'error': '本地采集失败'}

    score, status, issues = compute_health_score(info)
    info['health_score'] = score
    info['health_status'] = status
    info['issues'] = issues
    return info


def test_ssh_connection(ssh_host, ssh_port=22, ssh_user='root',
                        ssh_password='', ssh_key_file=''):
    """
    测试 SSH 连接是否成功。
    返回 (ok: bool, msg: str)
    """
    if not paramiko:
        return False, 'paramiko 未安装'
    collector = RemoteSystemInfoCollector(
        host=ssh_host, port=int(ssh_port),
        username=ssh_user,
        password=ssh_password if ssh_password else None,
        key_file=ssh_key_file if ssh_key_file else None,
    )
    try:
        if collector.connect():
            collector.disconnect()
            return True, 'SSH 连接成功'
        return False, 'SSH 连接失败'
    except Exception as e:
        return False, f'SSH 连接失败: {e}'


# ─── Word 报告生成 ────────────────────────────────────────────────────

def generate_server_report(info, output_dir=None):
    """
    生成服务器巡检 Word 报告。
    返回 (ok, filepath_or_error)
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return False, 'python-docx 未安装'

    if not output_dir:
        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'reports')
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # ── 封面 ──
    logo_path = os.path.join(os.path.dirname(__file__), 'dbcheck_logo.png')
    if os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(logo_path, width=Cm(3.5))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('服务器巡检报告')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 75, 135)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run('Server Health Inspection Report')
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(100, 100, 100)
    sr.font.italic = True

    doc.add_paragraph()

    # 信息表
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    info_rows = [
        ('主机名', info.get('hostname', '')),
        ('主机地址', info.get('host', '')),
        ('操作系统', info.get('platform', '')),
        ('内核版本', info.get('kernel', '')),
        ('巡检时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
    ]
    for i, (label, value) in enumerate(info_rows):
        cells = table.rows[i].cells
        cells[0].text = label
        cells[1].text = str(value)

    doc.add_page_break()

    # ── 第1章：健康评分 ──
    doc.add_heading('1. 健康评分', level=1)
    score = info.get('health_score', 0)
    status = info.get('health_status', '未知')
    p = doc.add_paragraph()
    run = p.add_run(f'{score} 分 — {status}')
    run.font.size = Pt(16)
    run.font.bold = True
    if score >= 90:
        run.font.color.rgb = RGBColor(46, 160, 67)
    elif score >= 75:
        run.font.color.rgb = RGBColor(56, 139, 253)
    elif score >= 60:
        run.font.color.rgb = RGBColor(227, 179, 65)
    else:
        run.font.color.rgb = RGBColor(248, 81, 73)

    issues = info.get('issues', [])
    if issues:
        doc.add_paragraph('发现的问题：')
        for issue in issues:
            doc.add_paragraph(f'• {issue}')

    # ── 第2章：系统资源 ──
    doc.add_heading('2. 系统资源检查', level=1)

    # CPU
    doc.add_heading('2.1 CPU 信息', level=2)
    cpu = info.get('cpu', {})
    t = doc.add_table(rows=2, cols=6)
    t.style = 'Table Grid'
    hdrs = ['使用率', '物理核心', '逻辑核心', '当前频率(GHz)', '负载(1m)', '负载(5m)']
    for i, h in enumerate(hdrs):
        t.rows[0].cells[i].text = h
    vals = [
        f"{cpu.get('usage_percent', 0):.1f}%",
        str(cpu.get('physical_cores', '')),
        str(cpu.get('logical_cores', '')),
        f"{cpu.get('current_frequency', 0):.2f}",
        f"{cpu.get('load_1m', 0):.2f}",
        f"{cpu.get('load_5m', 0):.2f}",
    ]
    for i, v in enumerate(vals):
        t.rows[1].cells[i].text = v

    # 内存
    doc.add_heading('2.2 内存信息', level=2)
    mem = info.get('memory', {})
    t = doc.add_table(rows=2, cols=4)
    t.style = 'Table Grid'
    hdrs = ['总内存(GB)', '已使用(GB)', '可用(GB)', '使用率']
    for i, h in enumerate(hdrs):
        t.rows[0].cells[i].text = h
    vals = [
        f"{mem.get('total_gb', 0):.2f}",
        f"{mem.get('used_gb', 0):.2f}",
        f"{mem.get('available_gb', 0):.2f}",
        f"{mem.get('usage_percent', 0):.1f}%",
    ]
    for i, v in enumerate(vals):
        t.rows[1].cells[i].text = v

    # 磁盘
    doc.add_heading('2.3 磁盘信息', level=2)
    disks = info.get('disk', [])
    if disks:
        t = doc.add_table(rows=1 + len(disks), cols=6)
        t.style = 'Table Grid'
        hdrs = ['挂载点', '设备', '文件系统', '总量(GB)', '已用(GB)', '使用率']
        for i, h in enumerate(hdrs):
            t.rows[0].cells[i].text = h
        for ri, d in enumerate(disks, 1):
            t.rows[ri].cells[0].text = d.get('mountpoint', '')
            t.rows[ri].cells[1].text = d.get('device', '')
            t.rows[ri].cells[2].text = d.get('fstype', '')
            t.rows[ri].cells[3].text = f"{d.get('total_gb', 0):.1f}"
            t.rows[ri].cells[4].text = f"{d.get('used_gb', 0):.1f}"
            t.rows[ri].cells[5].text = f"{d.get('usage_percent', 0):.1f}%"

    # ── 第3章：Top 进程 ──
    doc.add_heading('3. Top 进程（按 CPU 排序）', level=1)
    procs = info.get('top_processes', [])
    if procs:
        t = doc.add_table(rows=1 + len(procs), cols=5)
        t.style = 'Table Grid'
        hdrs = ['PID', '用户', 'CPU%', '内存%', '命令']
        for i, h in enumerate(hdrs):
            t.rows[0].cells[i].text = h
        for ri, p in enumerate(procs, 1):
            t.rows[ri].cells[0].text = str(p.get('pid', ''))
            t.rows[ri].cells[1].text = p.get('user', '')
            t.rows[ri].cells[2].text = f"{p.get('cpu', 0):.1f}"
            t.rows[ri].cells[3].text = f"{p.get('mem', 0):.1f}"
            t.rows[ri].cells[4].text = p.get('command', '')

    # ── 第4章：其他信息 ──
    doc.add_heading('4. 其他系统信息', level=1)
    t = doc.add_table(rows=3, cols=2)
    t.style = 'Table Grid'
    rows_data = [
        ('启动时间', info.get('boot_time', '')),
        ('运行天数', f"{info.get('uptime_days', 0)} 天"),
        ('打开文件描述符', str(info.get('open_files', {}))),
    ]
    for i, (label, value) in enumerate(rows_data):
        t.rows[i].cells[0].text = label
        t.rows[i].cells[1].text = value

    # ── 保存 ──
    hostname = info.get('hostname', 'unknown').replace(' ', '_')
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f'服务器巡检_{hostname}_{ts}.docx'
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return True, filepath


# ─── 网络连通性检测 ──────────────────────────────────────────────────

def check_network_connectivity(ssh_collector):
    """
    通过 SSH 检测远程主机的网络连通性。
    返回 {'ping': {...}, 'dns': {...}, 'ports': [...]}
    """
    result = {'ping': {}, 'dns': {}, 'ports': []}

    # ping 检测（外网 + 内网网关）
    for target in ['8.8.8.8', '114.114.114.114']:
        out, _ = ssh_collector.exec_cmd(f'ping -c 2 -W 3 {target} 2>/dev/null | tail -1')
        if out and 'avg' in out:
            try:
                avg_ms = out.split('=')[1].split('/')[1]
                result['ping'][target] = {'ok': True, 'latency_ms': float(avg_ms)}
            except (IndexError, ValueError):
                result['ping'][target] = {'ok': True, 'latency_ms': 0}
        else:
            result['ping'][target] = {'ok': False, 'latency_ms': 0}

    # DNS 解析检测
    for domain in ['baidu.com', 'google.com']:
        out, err = ssh_collector.exec_cmd(f'nslookup {domain} 2>&1 | head -5')
        ok = bool(out and 'Address' in out and 'can\'t find' not in out.lower())
        result['dns'][domain] = {'ok': ok}

    # 常用端口连通性检测（本机监听状态）
    ports_to_check = [
        (22, 'SSH'), (80, 'HTTP'), (443, 'HTTPS'),
        (3306, 'MySQL'), (5432, 'PostgreSQL'), (1521, 'Oracle'),
        (6379, 'Redis'), (27017, 'MongoDB'), (8080, 'HTTP-Alt'),
    ]
    out, _ = ssh_collector.exec_cmd("ss -tlnp 2>/dev/null | tail -n +2 | awk '{print $4}' | grep -oE '[0-9]+$' | sort -un")
    listening_ports = set()
    if out:
        for line in out.strip().split('\n'):
            try:
                listening_ports.add(int(line.strip()))
            except ValueError:
                pass

    for port, name in ports_to_check:
        result['ports'].append({
            'port': port,
            'name': name,
            'listening': port in listening_ports,
        })

    return result


# ─── 服务状态检测 ────────────────────────────────────────────────────

def check_service_status(ssh_collector):
    """
    通过 SSH 检测常见服务的运行状态。
    返回 [{'name': ..., 'status': 'running'|'stopped'|'not_installed'}]
    """
    services = [
        'nginx', 'httpd', 'apache2',
        'mysqld', 'mariadb', 'postgresql',
        'redis', 'redis-server',
        'mongod', 'docker', 'containerd',
        'firewalld', 'iptables',
        'crond', 'cron',
        'sshd', 'ntpd', 'chronyd',
    ]

    results = []
    # 批量检测（减少 SSH 往返）
    svc_list = ' '.join(services)
    out, _ = ssh_collector.exec_cmd(
        f'for s in {svc_list}; do '
        f'if systemctl is-enabled "$s" >/dev/null 2>&1; then '
        f'state=$(systemctl is-active "$s" 2>/dev/null); '
        f'echo "$s:$state"; '
        f'else '
        f'echo "$s:not_installed"; '
        f'fi; done'
    )

    if out:
        for line in out.strip().split('\n'):
            if ':' in line:
                parts = line.strip().split(':', 1)
                name = parts[0].strip()
                raw_status = parts[1].strip()
                if raw_status == 'active':
                    status = 'running'
                elif raw_status in ('inactive', 'failed', 'deactivating'):
                    status = 'stopped'
                else:
                    status = 'not_installed'
                results.append({'name': name, 'status': status})

    return results


# ─── 巡检历史管理 ────────────────────────────────────────────────────

def _get_history_db_path():
    """获取历史数据库路径"""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    data_dir = os.path.join(base_dir, 'data')
    os.makedirs(data_dir, exist_ok=True)
    return os.path.join(data_dir, 'server_history.db')


def _init_history_db():
    """初始化巡检历史数据库表"""
    db_path = _get_history_db_path()
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS server_inspection_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            host TEXT NOT NULL,
            port INTEGER DEFAULT 22,
            hostname TEXT,
            health_score INTEGER,
            health_status TEXT,
            cpu_usage REAL,
            memory_usage REAL,
            disk_usage_max REAL,
            issues_count INTEGER,
            issues_json TEXT,
            report_path TEXT,
            result_json TEXT,
            inspect_time TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    conn.close()


def save_server_inspection(host, port, result, report_path=None):
    """保存巡检历史到 server_history.db"""
    _init_history_db()
    db_path = _get_history_db_path()
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # 计算最大磁盘使用率
    disks = result.get('disk', [])
    disk_max = max((d.get('usage_percent', 0) for d in disks), default=0)

    cursor.execute("""
        INSERT INTO server_inspection_history
        (host, port, hostname, health_score, health_status,
         cpu_usage, memory_usage, disk_usage_max, issues_count,
         issues_json, report_path, result_json, inspect_time)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        host, int(port),
        result.get('hostname', ''),
        result.get('health_score', 0),
        result.get('health_status', ''),
        result.get('cpu', {}).get('usage_percent', 0),
        result.get('memory', {}).get('usage_percent', 0),
        round(disk_max, 1),
        len(result.get('issues', [])),
        json.dumps(result.get('issues', []), ensure_ascii=False),
        report_path or '',
        json.dumps(result, ensure_ascii=False, default=str),
        datetime.now().isoformat(),
    ))

    conn.commit()
    record_id = cursor.lastrowid
    conn.close()
    return record_id


def get_server_inspection_history(host=None, limit=50):
    """获取巡检历史列表"""
    _init_history_db()
    db_path = _get_history_db_path()
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    if host:
        cursor.execute("""
            SELECT id, host, port, hostname, health_score, health_status,
                   cpu_usage, memory_usage, disk_usage_max, issues_count,
                   report_path, inspect_time
            FROM server_inspection_history
            WHERE host = ?
            ORDER BY inspect_time DESC
            LIMIT ?
        """, (host, limit))
    else:
        cursor.execute("""
            SELECT id, host, port, hostname, health_score, health_status,
                   cpu_usage, memory_usage, disk_usage_max, issues_count,
                   report_path, inspect_time
            FROM server_inspection_history
            ORDER BY inspect_time DESC
            LIMIT ?
        """, (limit,))

    rows = cursor.fetchall()
    conn.close()
    return [dict(row) for row in rows]


def get_server_inspection_detail(record_id):
    """获取单条巡检详情（含完整 result_json）"""
    _init_history_db()
    db_path = _get_history_db_path()
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    cursor.execute("SELECT * FROM server_inspection_history WHERE id = ?", (record_id,))
    row = cursor.fetchone()
    conn.close()

    if not row:
        return None

    record = dict(row)
    # 解析 result_json
    if record.get('result_json'):
        try:
            record['result'] = json.loads(record['result_json'])
        except (json.JSONDecodeError, TypeError):
            record['result'] = {}
    # 解析 issues_json
    if record.get('issues_json'):
        try:
            record['issues'] = json.loads(record['issues_json'])
        except (json.JSONDecodeError, TypeError):
            record['issues'] = []

    return record


def delete_server_inspection(record_id):
    """删除指定巡检历史记录，同时删除对应的报告文件"""
    _init_history_db()
    db_path = _get_history_db_path()
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    # 先查 report_path
    cursor.execute("SELECT report_path FROM server_inspection_history WHERE id = ?", (record_id,))
    row = cursor.fetchone()
    report_path = row[0] if row and row[0] else ''
    # 删除 DB 记录
    cursor.execute("DELETE FROM server_inspection_history WHERE id = ?", (record_id,))
    deleted = cursor.rowcount
    conn.commit()
    conn.close()
    # 同步删除报告文件
    if deleted and report_path and os.path.isfile(report_path):
        try:
            os.remove(report_path)
        except Exception:
            pass
    return deleted > 0





def delete_server_inspection_by_filename(filename):
    """按报告文件名删除对应的巡检历史记录"""
    _init_history_db()
    db_path = _get_history_db_path()
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM server_inspection_history WHERE report_path LIKE ?", (f'%{filename}',))
    deleted = cursor.rowcount
    conn.commit()
    conn.close()
    return deleted > 0
# ─── 分享链接功能 ──────────────────────────────────────────────────

def _init_shares_db():
    """初始化分享链接数据库表"""
    db_path = _get_history_db_path()
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS shares (
            id TEXT PRIMARY KEY,
            share_type TEXT NOT NULL,
            title TEXT,
            data_json TEXT NOT NULL,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            access_count INTEGER DEFAULT 0
        )
    """)
    conn.commit()
    conn.close()


def create_share(share_type, title, data):
    """创建分享链接，返回 share_id"""
    import uuid
    _init_shares_db()
    share_id = uuid.uuid4().hex[:12]
    db_path = _get_history_db_path()
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO shares (id, share_type, title, data_json)
        VALUES (?, ?, ?, ?)
    """, (share_id, share_type, title, json.dumps(data, ensure_ascii=False)))
    conn.commit()
    conn.close()
    return share_id


def get_share(share_id):
    """获取分享数据"""
    _init_shares_db()
    db_path = _get_history_db_path()
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("SELECT share_type, title, data_json, created_at FROM shares WHERE id = ?", (share_id,))
    row = cursor.fetchone()
    if row:
        # 更新访问计数
        cursor.execute("UPDATE shares SET access_count = access_count + 1 WHERE id = ?", (share_id,))
        conn.commit()
    conn.close()
    if row:
        return {
            'share_type': row[0],
            'title': row[1],
            'data': json.loads(row[2]),
            'created_at': row[3]
        }
    return None


def delete_share(share_id):
    """删除分享链接"""
    _init_shares_db()
    db_path = _get_history_db_path()
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM shares WHERE id = ?", (share_id,))
    deleted = cursor.rowcount
    conn.commit()
    conn.close()
    return deleted > 0


def list_shares():
    """获取所有分享链接列表"""
    _init_shares_db()
    db_path = _get_history_db_path()
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("SELECT id, share_type, title, created_at, access_count FROM shares ORDER BY created_at DESC")
    rows = cursor.fetchall()
    conn.close()
    return [{'id': r[0], 'share_type': r[1], 'title': r[2], 'created_at': r[3], 'access_count': r[4]} for r in rows]


# ─── 分享 HTML 生成（保留用于下载）──────────────────────────────────────

def generate_server_share_html(result, output_dir=None):
    """
    生成自包含的服务器巡检分享 HTML 页面。
    返回 (ok, filepath_or_error)
    """
    if not output_dir:
        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'reports')
    os.makedirs(output_dir, exist_ok=True)

    score = result.get('health_score', 0)
    status = result.get('health_status', '未知')
    hostname = result.get('hostname', 'unknown')
    host = result.get('host', '')
    issues = result.get('issues', [])
    cpu = result.get('cpu', {})
    mem = result.get('memory', {})
    disks = result.get('disk', [])
    procs = result.get('top_processes', [])
    network = result.get('network', {})
    services = result.get('services', [])
    inspect_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    # 评分颜色
    if score >= 90:
        score_color = '#2ea043'
    elif score >= 75:
        score_color = '#388bfd'
    elif score >= 60:
        score_color = '#e3b341'
    else:
        score_color = '#f85149'

    # 磁盘表格行
    disk_rows = ''
    for d in disks:
        pct = d.get('usage_percent', 0)
        pct_color = '#f85149' if pct > 90 else ('#e3b341' if pct > 80 else '#c9d1d9')
        disk_rows += f'''<tr>
          <td>{d.get('mountpoint', '')}</td><td>{d.get('device', '')}</td>
          <td>{d.get('fstype', '')}</td><td>{d.get('total_gb', 0):.1f} GB</td>
          <td>{d.get('used_gb', 0):.1f} GB</td>
          <td style="color:{pct_color};font-weight:600">{pct:.1f}%</td></tr>'''

    # 进程表格行
    proc_rows = ''
    for p in procs:
        proc_rows += f'''<tr>
          <td>{p.get('pid', '')}</td><td>{esc_html(p.get('user', ''))}</td>
          <td>{p.get('cpu', 0):.1f}</td><td>{p.get('mem', 0):.1f}</td>
          <td style="font-size:12px">{esc_html(p.get('command', '')[:80])}</td></tr>'''

    # 网络检测行
    network_rows = ''
    if network:
        ping_data = network.get('ping', {})
        for target, info in ping_data.items():
            ok = info.get('ok', False)
            icon = '✅' if ok else '❌'
            latency = f"{info.get('latency_ms', 0):.1f} ms" if ok else '超时'
            network_rows += f'<tr><td>Ping {target}</td><td>{icon}</td><td>{latency}</td></tr>'

        dns_data = network.get('dns', {})
        for domain, info in dns_data.items():
            ok = info.get('ok', False)
            icon = '✅' if ok else '❌'
            network_rows += f'<tr><td>DNS {domain}</td><td>{icon}</td><td>{"正常" if ok else "解析失败"}</td></tr>'

        ports_data = network.get('ports', [])
        for p in ports_data:
            listening = p.get('listening', False)
            icon = '✅' if listening else '➖'
            network_rows += f'<tr><td>端口 {p["port"]} ({p["name"]})</td><td>{icon}</td><td>{"监听中" if listening else "未监听"}</td></tr>'

    # 服务状态行
    service_rows = ''
    for s in services:
        st = s.get('status', 'not_installed')
        if st == 'running':
            icon, color, label = '✅', '#2ea043', '运行中'
        elif st == 'stopped':
            icon, color, label = '❌', '#f85149', '已停止'
        else:
            icon, color, label = '➖', '#8b949e', '未安装'
        service_rows += f'<tr><td>{esc_html(s["name"])}</td><td style="color:{color}">{icon} {label}</td></tr>'

    # 问题列表
    issues_html = ''
    if issues:
        items = ''.join(f'<div class="issue-item">{esc_html(i)}</div>' for i in issues)
        issues_html = f'<div class="section"><h2>⚠️ 发现的问题</h2>{items}</div>'

    # 网络/服务区块
    network_section = ''
    if network_rows:
        network_section = f'''<div class="section"><h2>🌐 网络连通性检测</h2>
          <table><tr><th>检测项</th><th>状态</th><th>结果</th></tr>{network_rows}</table></div>'''

    service_section = ''
    if service_rows:
        service_section = f'''<div class="section"><h2>🔧 服务状态</h2>
          <table><tr><th>服务</th><th>状态</th></tr>{service_rows}</table></div>'''

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>服务器巡检报告 - {esc_html(hostname)}</title>
<style>
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ background: #0d1117; color: #c9d1d9; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif; padding: 20px; line-height: 1.6; }}
  .container {{ max-width: 960px; margin: 0 auto; }}
  .header {{ text-align: center; padding: 40px 20px; background: linear-gradient(135deg, #161b22 0%, #0d1117 100%); border: 1px solid #30363d; border-radius: 12px; margin-bottom: 20px; }}
  .header h1 {{ font-size: 24px; color: #f0f6fc; margin-bottom: 8px; }}
  .header .subtitle {{ color: #8b949e; font-size: 14px; }}
  .score-box {{ font-size: 64px; font-weight: 700; color: {score_color}; margin: 16px 0 4px; }}
  .score-label {{ font-size: 18px; color: {score_color}; }}
  .meta {{ color: #8b949e; font-size: 13px; margin-top: 12px; }}
  .section {{ background: #161b22; border: 1px solid #30363d; border-radius: 10px; padding: 20px; margin-bottom: 16px; }}
  .section h2 {{ font-size: 16px; color: #f0f6fc; margin-bottom: 12px; }}
  .grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 16px; }}
  .metric {{ font-size: 13px; line-height: 2; }}
  .metric b {{ color: #f0f6fc; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 13px; }}
  th {{ text-align: left; padding: 8px; border-bottom: 1px solid #30363d; color: #8b949e; font-weight: 500; }}
  td {{ padding: 8px; border-bottom: 1px solid #21262d; }}
  .issue-item {{ padding: 8px 12px; margin-bottom: 6px; background: rgba(248,81,73,.08); border-left: 3px solid #f85149; border-radius: 4px; font-size: 13px; }}
  .footer {{ text-align: center; color: #484f58; font-size: 12px; padding: 20px; }}
  @media (max-width: 600px) {{ .grid {{ grid-template-columns: 1fr; }} }}
</style>
</head>
<body>
<div class="container">
  <div class="header">
    <h1>🖥️ 服务器巡检报告</h1>
    <div class="subtitle">{esc_html(hostname)} {f'({esc_html(host)})' if host else ''}</div>
    <div class="score-box">{score}</div>
    <div class="score-label">{esc_html(status)}</div>
    <div class="meta">巡检时间: {inspect_time} | 内核: {esc_html(result.get('kernel', '-'))} | 运行: {result.get('uptime_days', 0)} 天</div>
  </div>

  {issues_html}

  <div class="section">
    <h2>📊 系统资源</h2>
    <div class="grid">
      <div class="metric">
        <b>🔲 CPU</b><br/>
        使用率: <b>{cpu.get('usage_percent', 0):.1f}%</b><br/>
        物理核心: {cpu.get('physical_cores', '-')} | 逻辑核心: {cpu.get('logical_cores', '-')}<br/>
        频率: {cpu.get('current_frequency', '-')} / {cpu.get('max_frequency', '-')} GHz<br/>
        负载: {cpu.get('load_1m', 0):.2f} / {cpu.get('load_5m', 0):.2f} / {cpu.get('load_15m', 0):.2f}
      </div>
      <div class="metric">
        <b>💾 内存</b><br/>
        总内存: <b>{mem.get('total_gb', 0):.2f} GB</b><br/>
        已使用: {mem.get('used_gb', 0):.2f} GB ({mem.get('usage_percent', 0):.1f}%)<br/>
        可用: {mem.get('available_gb', 0):.2f} GB<br/>
        Swap: {mem.get('swap_used_gb', 0):.2f} / {mem.get('swap_total_gb', 0):.2f} GB
      </div>
    </div>
  </div>

  <div class="section">
    <h2>💿 磁盘</h2>
    <table>
      <tr><th>挂载点</th><th>设备</th><th>文件系统</th><th>总量</th><th>已用</th><th>使用率</th></tr>
      {disk_rows if disk_rows else '<tr><td colspan="6" style="color:#8b949e">无磁盘信息</td></tr>'}
    </table>
  </div>

  {network_section}
  {service_section}

  <div class="section">
    <h2>⚙️ Top 进程</h2>
    <table>
      <tr><th>PID</th><th>用户</th><th>CPU%</th><th>内存%</th><th>命令</th></tr>
      {proc_rows if proc_rows else '<tr><td colspan="5" style="color:#8b949e">无进程信息</td></tr>'}
    </table>
  </div>

  <div class="footer">
    Generated by DBCheck — Database Intelligent Inspection Tool<br/>
    {inspect_time}
  </div>
</div>
</body>
</html>'''

    hostname_safe = hostname.replace(' ', '_').replace('/', '_')
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f'服务器巡检_{hostname_safe}_{ts}.html'
    filepath = os.path.join(output_dir, filename)

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(html)

    return True, filepath


def esc_html(text):
    """HTML 转义"""
    if not text:
        return ''
    return str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')
