# -*- coding: utf-8 -*-
"""
DBCheck - Oracle AWR 报告 Word 报告生成器
==========================================
将 awr_parser.py 解析的结构化数据生成格式化的 Word 分析报告。
复用 DBCheck 现有报告风格（蓝色表头、章节编号、封面）。
"""

import os
import re
import datetime
from typing import Dict, Any, Optional, List, Tuple

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

# ── 辅助函数 ──────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def _docx_table(doc, headers, rows, header_bg='336699', max_rows=100):
    """生成 Word 表格（带表头背景色），最多渲染 max_rows 行"""
    if not headers:
        return None
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            _set_cell_bg(hdr_cells[i], header_bg)
        except Exception:
            pass
    # 数据行
    display_rows = rows[:max_rows] if max_rows else rows
    for row in display_rows:
        cells = tbl.add_row().cells
        for ci in range(len(headers)):
            cell_val = row[ci] if ci < len(row) else ''
            cells[ci].text = str(cell_val) if cell_val is not None else ''
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
    if len(rows) > max_rows:
        p = doc.add_paragraph(f'（共 {len(rows)} 行，仅展示前 {max_rows} 行）')
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.italic = True
    return tbl

def _add_heading(doc, text, level=1, color=None, size=None):
    """添加带样式的标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = size
    return h

def _add_text(doc, text, size=None, color=None, bold=False):
    """添加普通文本段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size if size else Pt(10.5)
    if color:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    return p

# ── 章节渲染函数 ──────────────────────────────────────────────────────────

def _render_cover(doc, awr_data):
    """渲染封面"""
    meta = awr_data.get('metadata', {})

    # Logo
    logo_path = os.path.join(os.path.dirname(__file__), 'dbcheck_logo.png')
    if os.path.exists(logo_path):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture(logo_path, width=Cm(3.5))

    # 标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('Oracle AWR 分析报告')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 75, 135)
    title_run.font.name = '微软雅黑'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run('AWR Report Analysis by DBCheck')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    subtitle_run.font.italic = True
    subtitle_run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # 封面信息
    report_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    cover_labels = ['数据库', '实例', '快照范围', '分析时段', 'AWR 来源文件', '报告生成时间']
    cover_data = [
        meta.get('db_name', 'N/A'),
        meta.get('instance', 'N/A'),
        meta.get('snap_range', 'N/A'),
        meta.get('elapsed', 'N/A'),
        meta.get('source_file', os.path.basename(__file__)),
        report_time
    ]
    tbl = doc.add_table(rows=len(cover_labels), cols=2, style='Table Grid')
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(zip(cover_labels, cover_data)):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = str(value)
        for cell in tbl.rows[i].cells:
            cell.paragraphs[0].runs[0].font.name = '微软雅黑'
            cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
            if cell == tbl.rows[i].cells[0]:
                _set_cell_bg(cell, '336699')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

def _render_table_data(doc, tables_list, subtitle=''):
    """渲染一个或多个表格数据"""
    if not tables_list:
        p = doc.add_paragraph('（无数据）')
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(128, 128, 128)
        return

    if subtitle:
        h = doc.add_heading(subtitle, level=3)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.font.size = Pt(11)

    for tdata in tables_list:
        headers = tdata.get('headers', [])
        rows = tdata.get('rows', [])
        if headers:
            _docx_table(doc, headers, rows)
            doc.add_paragraph()

def _render_kv_pairs(doc, tables_list):
    """渲染键值对表格"""
    for tdata in tables_list:
        headers = tdata.get('headers', [])
        rows = tdata.get('rows', [])
        if headers and rows:
            _docx_table(doc, headers, rows)
            doc.add_paragraph()

# ── 诊断分析函数 ──────────────────────────────────────────────────────────

def _parse_float(s):
    """安全解析浮点数"""
    try:
        return float(str(s).replace(',', '').replace('%', '').strip())
    except (ValueError, TypeError):
        return None

def _diagnose_wait_events(doc, awr_data):
    """分析等待事件，生成诊断"""
    diagnostics = []

    # 从 fg_wait_events 中提取
    fg_events = awr_data.get('fg_wait_events', [])
    top_waits = []
    for tdata in fg_events:
        headers = tdata.get('headers', [])
        rows = tdata.get('rows', [])
        if not headers or not rows:
            continue
        # 查找 Event 和 Time% 列
        event_idx = None
        time_pct_idx = None
        class_idx = None
        for i, h in enumerate(headers):
            hl = h.lower().strip()
            if 'event' in hl and event_idx is None:
                event_idx = i
            if 'time' in hl and ('%' in hl or 'pct' in hl):
                time_pct_idx = i
            if 'class' in hl:
                class_idx = i

        if event_idx is not None:
            for row in rows[:15]:
                if event_idx < len(row):
                    evt_name = row[event_idx]
                    pct = _parse_float(row[time_pct_idx]) if time_pct_idx is not None and time_pct_idx < len(row) else None
                    cls = row[class_idx] if class_idx is not None and class_idx < len(row) else ''
                    top_waits.append((evt_name, pct, cls))

    if not top_waits:
        # 尝试从 wait_class 表格提取
        for tdata in fg_events:
            headers = tdata.get('headers', [])
            rows = tdata.get('rows', [])
            if not headers or not rows:
                continue
            for i, h in enumerate(headers):
                if 'class' in h.lower():
                    class_idx = i
                    for row in rows:
                        if class_idx < len(row):
                            cls_name = row[class_idx]
                            pct = None
                            for ci, cell in enumerate(row):
                                if ci != class_idx:
                                    pct = _parse_float(cell)
                                    break
                            top_waits.append((cls_name, pct, ''))
                    break

    # 生成诊断建议
    if top_waits:
        for evt_name, pct, cls in top_waits[:5]:
            if pct is not None and pct > 10:
                advice = _wait_event_advice(evt_name, pct, cls)
                if advice:
                    diagnostics.append(advice)

    return diagnostics

def _wait_event_advice(event_name, pct, wait_class):
    """根据等待事件名称返回诊断建议"""
    en = event_name.lower()

    if 'cpu' in en or 'idle' in en:
        return None  # CPU 和 idle 不算问题

    if 'db file sequential read' in en:
        if pct > 30:
            return ('高危', f'单块读等待 {event_name} 占比 {pct}%：可能存在全表扫描或缺失索引，建议检查 Top SQL 的执行计划')
        return ('中', f'单块读等待 {event_name} 占比 {pct}%：关注高频单块读 SQL')

    if 'db file scattered read' in en:
        return ('中', f'多块读等待 {event_name} 占比 {pct}%：可能存在全表扫描，建议优化 SQL 或添加索引')

    if 'direct path read' in en or 'direct path' in en:
        return ('低', f'直接路径读 {event_name} 占比 {pct}%：通常与排序、哈希操作或并行查询相关')

    if 'log file sync' in en or 'log file parallel write' in en:
        return ('中', f'日志写入等待 {event_name} 占比 {pct}%：commit 频繁或redo日志I/O瓶颈，检查redo log文件大小和磁盘性能')

    if 'enq:' in en or 'library cache lock' in en or 'library cache pin' in en:
        return ('中', f'锁等待 {event_name} 占比 {pct}%：存在锁竞争，检查DDL操作或并发DML')

    if 'latch' in en:
        return ('中', f'闩锁等待 {event_name} 占比 {pct}%：存在闩锁争用，检查并发访问热点')

    if 'gc cr' in en or 'gc buffer' in en:
        return ('中', f'RAC全局缓存等待 {event_name} 占比 {pct}%：检查跨节点数据访问模式')

    if 'buffer busy' in en:
        return ('中', f'缓冲块忙等待 {event_name} 占比 {pct}%：热点块争用，检查高并发访问的表')

    if 'sort' in en:
        return ('低', f'排序等待 {event_name} 占比 {pct}%：PGA可能不足，检查 sort_area_size 或 workarea_size_policy')

    if pct > 15:
        return ('中', f'等待事件 {event_name} 占比 {pct}%（类别: {wait_class}），需要关注')

    return None

def _diagnose_sql_performance(doc, awr_data):
    """分析 Top SQL 性能"""
    diagnostics = []
    top_sql = awr_data.get('top_sql', {})

    # 分析 elapsed time 排序的 SQL
    elapsed_tables = top_sql.get('elapsed', [])
    for tdata in elapsed_tables:
        headers = tdata.get('headers', [])
        rows = tdata.get('rows', [])
        if not headers or not rows:
            continue
        # 找 SQL_ID, Elapsed Time, Executions, Elaps/Exec 列
        sql_idx = elapsed_idx = exec_idx = per_exec_idx = None
        for i, h in enumerate(headers):
            hl = h.lower().strip()
            if 'sql_id' in hl:
                sql_idx = i
            elif 'elapsed' in hl and 'time' in hl and per_exec_idx is None:
                elapsed_idx = i
            elif 'executions' in hl or 'execs' in hl:
                exec_idx = i
            elif 'elaps' in hl and ('/exec' in hl or 'per' in hl):
                per_exec_idx = i

        for row in rows[:5]:
            sql_id = row[sql_idx] if sql_idx is not None and sql_idx < len(row) else 'unknown'
            per_exec = _parse_float(row[per_exec_idx]) if per_exec_idx is not None and per_exec_idx < len(row) else None
            execs = _parse_float(row[exec_idx]) if exec_idx is not None and exec_idx < len(row) else None

            if per_exec is not None and per_exec > 1000:
                diagnostics.append(('高危', f'SQL {sql_id} 单次执行耗时 {per_exec:.0f}ms，建议分析执行计划'))
            elif execs is not None and execs > 10000 and per_exec is not None and per_exec > 100:
                diagnostics.append(('中', f'SQL {sql_id} 执行 {execs:.0f} 次，单次 {per_exec:.0f}ms，高频执行需关注'))

    # 分析逻辑读
    lr_tables = top_sql.get('logical_reads', [])
    for tdata in lr_tables:
        headers = tdata.get('headers', [])
        rows = tdata.get('rows', [])
        if not headers or not rows:
            continue
        sql_idx = None
        gets_idx = None
        for i, h in enumerate(headers):
            hl = h.lower().strip()
            if 'sql_id' in hl:
                sql_idx = i
            elif 'buffer' in hl or 'gets' in hl:
                gets_idx = i

        for row in rows[:3]:
            sql_id = row[sql_idx] if sql_idx is not None and sql_idx < len(row) else 'unknown'
            gets = _parse_float(row[gets_idx]) if gets_idx is not None and gets_idx < len(row) else None
            if gets is not None and gets > 1000000:
                diagnostics.append(('中', f'SQL {sql_id} 逻辑读 {gets/1000000:.1f}M，可能存在全表扫描'))

    return diagnostics

def _diagnose_io(doc, awr_data):
    """分析 I/O 性能"""
    diagnostics = []
    file_io = awr_data.get('file_io', [])

    for tdata in file_io:
        headers = tdata.get('headers', [])
        rows = tdata.get('rows', [])
        if not headers or not rows:
            continue
        # 查找 Avg svc 或 lat 列（平均服务时间）
        lat_idx = None
        file_idx = None
        for i, h in enumerate(headers):
            hl = h.lower().strip()
            if ('svc' in hl or 'lat' in hl or 'latency' in hl) and 'avg' in hl:
                lat_idx = i
            if 'file' in hl or 'name' in hl:
                file_idx = i

        if lat_idx is not None:
            for row in rows:
                lat = _parse_float(row[lat_idx]) if lat_idx < len(row) else None
                file_name = row[file_idx] if file_idx is not None and file_idx < len(row) else ''
                if lat is not None and lat > 50:
                    diagnostics.append(('中', f'数据文件 {file_name} 平均 I/O 延迟 {lat:.0f}ms，超过50ms阈值'))

    return diagnostics

def _preprocess_ai_markdown(text):
    """预处理 AI 生成的 Markdown 文本，返回清理后的标准 Markdown。

    处理规则：
    1. 删除整行 **xxx**（无数字编号，如 **Oracle数据库性能优化建议（按优先级排序）**）
    2. **N. xxx** → ## 16.N xxx（二级标题，自动编号）
    3. 其他行内 **xxx** → 去掉 ** 标记
    4. 合并连续空行
    5. 去掉末尾多余空行
    """
    lines = text.split('\n')
    result = []
    heading_counter = 0
    prev_blank = False

    for line in lines:
        stripped = line.strip()

        # 规则4：跳过空行（合并连续空行）
        if not stripped:
            if not prev_blank and result:
                result.append('')
                prev_blank = True
            continue
        prev_blank = False

        # 规则1+2：检查是否整行 **xxx**
        if stripped.startswith('**') and stripped.endswith('**'):
            inner = stripped[2:-2].strip()
            # 提取开头的数字（支持 1. 或 1、）
            num_str = ''
            title_start = 0
            for j, ch in enumerate(inner):
                if ch.isdigit():
                    num_str += ch
                elif ch in '.、':
                    title_start = j + 1
                    break
                else:
                    break
            if num_str:
                # 有数字编号 → 规则2：转换成 ## N. xxx（二级标题，自动编号）
                heading_counter += 1
                title = inner[title_start:].strip()
                result.append(f'## {heading_counter}. {title}')
                continue
            else:
                # 无数字编号
                # 子标题关键词（应保留并转成 ### xxx，三级标题）
                sub_keywords = ['问题描述', '具体操作', '根因分析', '影响范围',
                                '解决方案', '建议措施', '优化步骤', '执行步骤']
                is_sub_heading = any(kw in inner for kw in sub_keywords)
                # 纯总标题关键词（应删除）
                generic_keywords = ['优化建议', '性能优化', '优化方案',
                                    '诊断建议', '性能建议', '优化措施']
                is_generic_title = any(kw in inner for kw in generic_keywords)
                if is_generic_title and not is_sub_heading:
                    continue
                # 子标题 → 转成 ### xxx（三级标题）
                if is_sub_heading:
                    result.append(f'### {inner}')
                else:
                    # 其他无编号标题 → 转成 ## xxx（二级标题）
                    result.append(f'## {inner}')
                continue

        # 规则3：去掉行内 ** 标记
        cleaned = stripped.replace('**', '')
        result.append(cleaned)

    # 规则5：去掉末尾空行
    while result and not result[-1].strip():
        result.pop()

    return '\n'.join(result)

def _render_markdown_to_word(doc, md_text):
    """将标准 Markdown 文本渲染为 Word 段落。

    只处理标准 Markdown（# ## ### 标题、有序/无序列表、
    引用、代码块、行内 **bold** *italic* `code`）。
    ** 整行标题由 _preprocess_ai_markdown() 预处理转换，此处不再处理。
    """
    lines = md_text.split('\n')
    in_list = False
    in_code_block = False
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块 ``` 开关
        if stripped.startswith('```'):
            if in_code_block:
                in_code_block = False
                doc.add_paragraph()
            else:
                in_code_block = True
                in_list = False
                doc.add_paragraph()
            i += 1
            continue

        # 代码块内内容
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Pt(24)
            run = p.add_run(stripped)
            run.font.size = Pt(9.5)
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue

        # 引用块 > xxx
        if stripped.startswith('>'):
            if in_list:
                in_list = False
                doc.add_paragraph()
            quote_text = re.sub(r'^>\s*', '', stripped)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Pt(24)
            p.paragraph_format.right_indent = Pt(12)
            _add_inline_runs(p, quote_text)
            if p.runs:
                p.runs[0].italic = True
                p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            i += 1
            continue

        # 表格行 | col1 | col2 |
        if stripped.startswith('|') and stripped.endswith('|'):
            if in_list:
                in_list = False
                doc.add_paragraph()
            is_separator = all(re.match(r'^[\s\-:|]+$', c) for c in stripped.split('|')[1:-1])
            if is_separator:
                i += 1
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1] if c.strip()]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Pt(12)
            _add_inline_runs(p, '  |  '.join(cells))
            i += 1
            continue

        # 分割线 ---
        if re.match(r'^(---+|___+|\*\*\*+)\s*$', stripped):
            if in_list:
                in_list = False
                doc.add_paragraph()
            doc.add_paragraph()
            i += 1
            continue

        # 空行 → 合并连续空行
        if not stripped:
            if in_list:
                in_list = False
                doc.add_paragraph()
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines):
                next_line = lines[j].strip()
                # 如果下一段是标题，不添加空段落（标题自带间距）
                is_next_heading = bool(re.match(r'^(#{1,4})\s+', next_line))
                if not is_next_heading:
                    doc.add_paragraph()
            i = j
            continue

        # 标题 # ## ###
        heading_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if heading_match:
            if in_list:
                in_list = False
                doc.add_paragraph()
            level = len(heading_match.group(1))
            title_text = heading_match.group(2).strip()
            sizes = {1: Pt(14), 2: Pt(12), 3: Pt(11), 4: Pt(11)}
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(title_text)
            run.font.size = sizes.get(level, Pt(11))
            run.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            i += 1
            # 跳过紧跟的空行（标题自带间距，不需要额外空段落）
            while i < len(lines) and not lines[i].strip():
                i += 1
            continue

        # 有序列表 1. xxx
        ordered_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if ordered_match:
            item_text = ordered_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            num_run = p.add_run(f"{ordered_match.group(1)}. ")
            num_run.font.size = Pt(10.5)
            num_run.bold = True
            num_run.font.color.rgb = RGBColor(0, 102, 204)
            _add_inline_runs(p, item_text)
            in_list = True
            i += 1
            continue

        # 无序列表 - * •
        ul_match = re.match(r'^[-*•]\s+(.*)', stripped)
        if ul_match:
            item_text = ul_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            _add_inline_runs(p, item_text)
            in_list = True
            i += 1
            continue

        # 普通段落
        if in_list:
            in_list = False
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _add_inline_runs(p, stripped)
        i += 1


def _add_inline_runs(p, text):
    """在段落中渲染行内 Markdown：加粗、斜体、代码"""
    markers = []

    # 1) 收集 **bold** 的位置
    for m in re.finditer(r'\*\*([^*]+)\*\*', text):
        markers.append((m.start(), m.end(), 'bold', m.group(1)))
    # __bold__
    for m in re.finditer(r'__([^_]+)__', text):
        markers.append((m.start(), m.end(), 'bold', m.group(1)))

    # 2) 收集 *italic*（排除已在 **/__ 内部的）
    bold_ranges = [(m.start(), m.end()) for m in re.finditer(r'\*\*[^*]+\*\*|__[^_]+__', text)]
    for m in re.finditer(r'(?<!\*)\*(?!\*)([^*]+)(?<!\*)\*(?!\*)', text):
        if not any(s <= m.start() < e for s, e in bold_ranges):
            markers.append((m.start(), m.end(), 'italic', m.group(1)))
    for m in re.finditer(r'(?<!_)_(?!_)([^_]+)(?<!_)_(?!_)', text):
        if not any(s <= m.start() < e for s, e in bold_ranges):
            markers.append((m.start(), m.end(), 'italic', m.group(1)))

    # 3) 收集 `code`
    for m in re.finditer(r'`([^`]+)`', text):
        markers.append((m.start(), m.end(), 'code', m.group(1)))

    if not markers:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        return

    # 按起始位置排序，去除重叠（保留先出现的）
    markers.sort(key=lambda x: x[0])
    filtered = []
    last_end = -1
    for start, end, style, content in markers:
        if start >= last_end:
            filtered.append((start, end, style, content))
            last_end = end

    # 按区间渲染
    pos = 0
    for start, end, style, content in filtered:
        if pos < start:
            run = p.add_run(text[pos:start])
            run.font.size = Pt(10.5)
        run = p.add_run(content)
        run.font.size = Pt(10.5)
        if style == 'bold':
            run.bold = True
        elif style == 'italic':
            run.italic = True
        elif style == 'code':
            run.font.size = Pt(10)
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(139, 0, 0)
        pos = end

    if pos < len(text):
        run = p.add_run(text[pos:])
        run.font.size = Pt(10.5)

def _diagnose_cache(doc, awr_data):
    """分析缓存命中率"""
    diagnostics = []

    # 从 instance_efficiency 获取
    eff = awr_data.get('instance_efficiency', [])
    if isinstance(eff, list):
        for tdata in eff:
            if not isinstance(tdata, dict):
                continue
            headers = tdata.get('headers', [])
            rows = tdata.get('rows', [])
            if not headers or not rows:
                continue
            for i, h in enumerate(headers):
                hl = h.lower().strip()
                if 'hit' in hl or 'ratio' in hl:
                    for row in rows:
                        name = row[0] if len(row) > 0 else ''
                        val = _parse_float(row[i]) if i < len(row) else None
                        if val is not None and 'library' in name.lower() and val < 95:
                            diagnostics.append(('中', f'库缓存命中率 {val:.1f}%，低于95%建议增加 shared_pool_size'))
                        if val is not None and 'buffer' in name.lower() and 'hit' in name.lower() and val < 90:
                            diagnostics.append(('中', f'Buffer Cache 命中率 {val:.1f}%，低于90%建议增加 db_cache_size'))
    elif isinstance(eff, dict):
        for key, tdata in eff.items():
            if isinstance(tdata, dict) and 'headers' in tdata:
                headers = tdata.get('headers', [])
                rows = tdata.get('rows', [])
                for i, h in enumerate(headers):
                    hl = h.lower().strip()
                    if 'hit' in hl or 'ratio' in hl:
                        for row in rows:
                            name = row[0] if len(row) > 0 else ''
                            val = _parse_float(row[i]) if i < len(row) else None
                            if val is not None and 'library' in name.lower() and val < 95:
                                diagnostics.append(('中', f'库缓存命中率 {val:.1f}%，低于95%建议增加 shared_pool_size'))
                            if val is not None and 'buffer' in name.lower() and 'hit' in name.lower() and val < 90:
                                diagnostics.append(('中', f'Buffer Cache 命中率 {val:.1f}%，低于90%建议增加 db_cache_size'))

    return diagnostics

# ── 主报告生成函数 ────────────────────────────────────────────────────────

def build_awr_word_report(awr_data: Dict[str, Any], output_path: str = None, source_filename: str = ''):
    """
    根据 AWR 解析数据生成 Word 分析报告。

    Args:
        awr_data: awr_parser.parse_awr_report() 返回的结构化数据
        output_path: 输出 .docx 文件路径，默认为 AWR 同目录的 awr_analysis.docx
        source_filename: 原始 AWR HTML 文件名（用于封面显示）

    Returns:
        生成的 .docx 文件绝对路径
    """
    if not _HAS_DOCX:
        raise ImportError('python-docx 未安装，无法生成 Word 报告')

    meta = awr_data.get('metadata', {})
    if source_filename:
        meta['source_file'] = source_filename

    # 默认输出路径
    if output_path is None:
        db_name = meta.get('db_name', 'awr')
        snap = meta.get('snap_range', 'unknown')
        output_path = os.path.join(os.getcwd(), f'DBCheck_AWR_Analysis_{db_name}_{snap}.docx'.replace('/', '_').replace(' ', '_'))

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ── 封面 ──
    _render_cover(doc, awr_data)

    # ── 章节序号 ──
    _section_num = 0
    _subsection_num = 0

    def _add_section(title):
        nonlocal _section_num, _subsection_num
        _section_num += 1
        _subsection_num = 0
        h = doc.add_heading(f"{_section_num} {title}", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.font.size = Pt(14)

    def _add_subsection(title):
        nonlocal _subsection_num
        _subsection_num += 1
        h = doc.add_heading(f"{_section_num}.{_subsection_num} {title}", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.font.size = Pt(12)

    # ═══════════════════════════════════════════════════════════════════
    # 第1章：AWR 报告概要
    # ═══════════════════════════════════════════════════════════════════
    _add_section('AWR 报告概要')
    summary_kv = [
        ('数据库名称', meta.get('db_name', 'N/A')),
        ('实例名', meta.get('instance', 'N/A')),
        ('快照范围', meta.get('snap_range', 'N/A')),
        ('分析时段', meta.get('elapsed', 'N/A')),
        ('数据库版本', meta.get('db_version', 'N/A')),
        ('RAC', meta.get('rac', 'N/A')),
        ('AWR 来源文件', meta.get('source_file', 'N/A')),
    ]
    _docx_table(doc, ['项目', '值'], [[k, v] for k, v in summary_kv])
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 第2章：负载概况
    # ═══════════════════════════════════════════════════════════════════
    _add_section('负载概况')
    lp = awr_data.get('load_profile', [])
    if isinstance(lp, list) and lp:
        _add_subsection('负载概况')
        _render_table_data(doc, lp)
    elif isinstance(lp, dict):
        _add_subsection('每秒负载')
        _render_table_data(doc, lp.get('per_second', []))
        _add_subsection('每活跃会话负载')
        _render_table_data(doc, lp.get('per_active_session', []))
    else:
        _add_text(doc, '（无负载概况数据）')
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 第3章：实例效率
    # ═══════════════════════════════════════════════════════════════════
    _add_section('实例效率')
    eff = awr_data.get('instance_efficiency', [])
    if isinstance(eff, list) and eff:
        _render_table_data(doc, eff)
    elif isinstance(eff, dict):
        for key, tdata in eff.items():
            if isinstance(tdata, dict) and 'headers' in tdata:
                _render_table_data(doc, [tdata])
    else:
        _add_text(doc, '（无实例效率数据）')
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 第4章：时间模型
    # ═══════════════════════════════════════════════════════════════════
    _add_section('DB Time 模型')
    _render_table_data(doc, awr_data.get('time_model', []))

    # ═══════════════════════════════════════════════════════════════════
    # 第5章：等待事件分析
    # ═══════════════════════════════════════════════════════════════════
    _add_section('等待事件分析')
    _add_subsection('前台等待事件')
    _render_table_data(doc, awr_data.get('fg_wait_events', []))
    _add_subsection('后台等待事件')
    _render_table_data(doc, awr_data.get('bg_wait_events', []))

    # ═══════════════════════════════════════════════════════════════════
    # 第6章：系统统计
    # ═══════════════════════════════════════════════════════════════════
    _add_section('系统统计')
    _render_table_data(doc, awr_data.get('system_stats', []))

    # ═══════════════════════════════════════════════════════════════════
    # 第7章：闩锁统计
    # ═══════════════════════════════════════════════════════════════════
    _add_section('闩锁统计')
    _render_table_data(doc, awr_data.get('latch_stats', []))

    # ═══════════════════════════════════════════════════════════════════
    # 第8章：I/O 统计
    # ═══════════════════════════════════════════════════════════════════
    _add_section('I/O 统计')
    _add_subsection('文件 I/O')
    _render_table_data(doc, awr_data.get('file_io', []))
    _add_subsection('表空间 I/O')
    _render_table_data(doc, awr_data.get('tablespace_io', []))

    # ═══════════════════════════════════════════════════════════════════
    # 第9章：服务统计
    # ═══════════════════════════════════════════════════════════════════
    _add_section('服务统计')
    _render_table_data(doc, awr_data.get('service_stats', []))

    # ═══════════════════════════════════════════════════════════════════
    # 第10章：内存与 Undo/Redo
    # ═══════════════════════════════════════════════════════════════════
    _add_section('内存与 Undo/Redo')
    _add_subsection('SGA 内存')
    _render_table_data(doc, awr_data.get('sga_memory', []))
    _add_subsection('Undo 统计')
    _render_table_data(doc, awr_data.get('undo_stats', []))
    _add_subsection('Redo 统计')
    _render_table_data(doc, awr_data.get('redo_stats', []))

    # ═══════════════════════════════════════════════════════════════════
    # 第11章：Top SQL 分析
    # ═══════════════════════════════════════════════════════════════════
    _add_section('Top SQL 分析')
    sql_sorts = [
        ('elapsed', 'Elapsed Time'),
        ('cpu', 'CPU Time'),
        ('executions', 'Executions'),
        ('logical_reads', 'Logical Reads'),
        ('physical_reads', 'Physical Reads'),
        ('parse_calls', 'Parse Calls'),
        ('buffer_gets', 'Buffer Gets'),
        ('sharable_mem', 'Sharable Mem'),
    ]
    top_sql = awr_data.get('top_sql', {})
    for key, label in sql_sorts:
        tables = top_sql.get(key, [])
        if tables:
            _add_subsection(f'按 {label} 排序')
            _render_table_data(doc, tables)
    if not any(top_sql.get(k) for k, _ in sql_sorts):
        _add_text(doc, '（无 Top SQL 数据）')
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 第12章：段与对象统计
    # ═══════════════════════════════════════════════════════════════════
    _add_section('段与对象统计')
    seg = awr_data.get('segment_stats', {})
    seg_sorts = [
        ('buffer_gets', 'Buffer Gets'),
        ('physical_reads', 'Physical Reads'),
        ('logical_reads', 'Logical Reads'),
    ]
    for key, label in seg_sorts:
        tables = seg.get(key, [])
        if tables:
            _add_subsection(f'按 {label} 排序')
            _render_table_data(doc, tables)
    _add_subsection('对象统计')
    _render_table_data(doc, awr_data.get('object_stats', []))

    # ═══════════════════════════════════════════════════════════════════
    # 第13章：SQL 执行计划变更
    # ═══════════════════════════════════════════════════════════════════
    _add_section('SQL 执行计划变更')
    _render_table_data(doc, awr_data.get('sql_plan_changes', []))

    # ═══════════════════════════════════════════════════════════════════
    # 第14章：ASH（活跃会话历史）
    # ═══════════════════════════════════════════════════════════════════
    _add_section('ASH 活跃会话历史')
    _render_table_data(doc, awr_data.get('ash', []))

    # ═══════════════════════════════════════════════════════════════════
    # 附录：Cache Fusion（RAC 环境）
    # ═══════════════════════════════════════════════════════════════════
    cf = awr_data.get('cache_fusion', [])
    if cf:
        _add_section('Cache Fusion 统计')
        _render_table_data(doc, cf)

    # ═══════════════════════════════════════════════════════════════════
    # 诊断分析总结（DBCheck 自动生成）
    # ═══════════════════════════════════════════════════════════════════
    _add_section('DBCheck AWR 诊断分析')

    _add_text(doc, '以下诊断由 DBCheck 基于 AWR 数据自动生成，仅供参考。', size=Pt(10.5), color=RGBColor(128, 128, 128))
    doc.add_paragraph()

    all_diagnostics = []
    all_diagnostics.extend(_diagnose_wait_events(doc, awr_data))
    all_diagnostics.extend(_diagnose_sql_performance(doc, awr_data))
    all_diagnostics.extend(_diagnose_io(doc, awr_data))
    all_diagnostics.extend(_diagnose_cache(doc, awr_data))

    if all_diagnostics:
        # 按风险等级排序
        level_order = {'高危': 0, '中': 1, '低': 2}
        all_diagnostics.sort(key=lambda x: level_order.get(x[0], 99))

        _add_subsection('发现的风险与建议')
        level_colors = {
            '高危': RGBColor(204, 0, 0),
            '中': RGBColor(204, 102, 0),
            '低': RGBColor(0, 128, 0),
        }
        for level, msg in all_diagnostics:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run_level = p.add_run(f'[{level}] ')
            run_level.font.bold = True
            run_level.font.size = Pt(10.5)
            run_level.font.color.rgb = level_colors.get(level, RGBColor(0, 0, 0))
            run_msg = p.add_run(msg)
            run_msg.font.size = Pt(10.5)
    else:
        _add_text(doc, '未检测到明显性能问题。', size=Pt(11), color=RGBColor(0, 128, 0), bold=True)

    # ═══════════════════════════════════════════════════════════════════
    # AI 智能诊断章节（如果启用了 AI）
    # ═══════════════════════════════════════════════════════════════════
    ai_result = awr_data.get('ai_diagnosis', '')
    if ai_result and ai_result.strip():
        _add_section('AI 智能诊断建议')
        _add_text(doc, '以下建议由 AI 基于 AWR 数据自动生成，仅供参考，请结合实际业务场景判断.',
                  size=Pt(10), color=RGBColor(128, 128, 128))

        # 预处理：去掉不需要的行、转换 **N. xxx** 为 ## 标题、清理空行
        cleaned = _preprocess_ai_markdown(ai_result.strip())
        _render_markdown_to_word(doc, cleaned)

    doc.add_paragraph()
    _add_text(doc, f'报告由 DBCheck 自动生成 | 生成时间: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}',
              size=Pt(9), color=RGBColor(160, 160, 160))

    # 保存
    doc.save(output_path)
    return os.path.abspath(output_path)

if __name__ == '__main__':
    import sys
    import json

    if len(sys.argv) < 2:
        print("Usage: python build_awr_word_report.py <awr_parsed.json> [output.docx]")
        sys.exit(1)

    json_path = sys.argv[1]
    output = sys.argv[2] if len(sys.argv) > 2 else None

    with open(json_path, 'r', encoding='utf-8') as f:
        awr_data = json.load(f)

    result = build_awr_word_report(awr_data, output)
    print(f"AWR analysis report generated: {result}")
